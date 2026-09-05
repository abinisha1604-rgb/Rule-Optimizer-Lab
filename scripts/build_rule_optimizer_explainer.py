from __future__ import annotations

import math
import os
import textwrap
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "output" / "documents"
BUILD_DIR = ROOT / "tmp" / "docx_build" / "rule_optimizer_explainer"
OUTPUT_PATH = OUTPUT_DIR / "EFRM_Rule_Optimization_Agent_Refined_Enterprise_Architecture_and_Orchestration.docx"

OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
BUILD_DIR.mkdir(parents=True, exist_ok=True)


# ---------------------------------------------------------------------------
# Design system
# ---------------------------------------------------------------------------

NAVY = "#12314A"
BLUE = "#2E74B5"
TEAL = "#177E89"
PALE_TEAL = "#E7F4F5"
PALE_BLUE = "#EAF2F8"
PALE_ORANGE = "#FFF1DF"
ORANGE = "#E58B2A"
PALE_GREEN = "#EAF5EC"
GREEN = "#3F7D44"
PALE_RED = "#FBEAEA"
RED = "#B44949"
INK = "#24323D"
MID = "#5F6B73"
LIGHT = "#F4F7F9"
BORDER = "#CCD6DD"
WHITE = "#FFFFFF"

FONT_REGULAR = Path(r"C:\Windows\Fonts\calibri.ttf")
FONT_BOLD = Path(r"C:\Windows\Fonts\calibrib.ttf")


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value.lstrip("#"))


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill.lstrip("#"))


def set_cell_border(cell, **edges) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in edges:
            continue
        edge_data = edges[edge]
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key in ("val", "sz", "space", "color"):
            if key in edge_data:
                value = str(edge_data[key])
                if key == "color":
                    value = value.lstrip("#")
                element.set(qn("w:{}".format(key)), value)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_width(table, width_twips=9360) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_twips))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_table_grid(table, widths) -> None:
    """Make tblW, tblGrid, and every tcW agree in fixed DXA geometry."""
    set_table_width(table, sum(widths))
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths):
                set_cell_width(cell, widths[idx])


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_run_font(run, name="Calibri", size=11, color=INK, bold=False, italic=False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.font.bold = bold
    run.font.italic = italic


def set_keep_with_next(paragraph, value=True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep_next = p_pr.find(qn("w:keepNext"))
    if keep_next is None:
        keep_next = OxmlElement("w:keepNext")
        p_pr.append(keep_next)
    keep_next.set(qn("w:val"), "1" if value else "0")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MID)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def set_section_header_footer(section, title: str) -> None:
    header = section.header
    hp = header.paragraphs[0]
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    run = hp.add_run(title)
    set_run_font(run, size=8.5, color=MID, bold=True)
    line = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), BORDER)
    line.append(bottom)
    hp._p.get_or_add_pPr().append(line)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.clear()
    add_page_number(fp)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    # Design preset: standard_business_brief.
    # Named overrides: editorial cover title block, dense technical tables (8.0-9.2 pt),
    # and full-page architecture figures sized to the 6.5-inch content width.
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    set_section_header_footer(section, "EFRM Rule Optimization Agent | Enterprise Architecture")
    section.different_first_page_header_footer = True
    section.first_page_header.paragraphs[0].clear()
    section.first_page_footer.paragraphs[0].clear()

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Title", 30, NAVY, 0, 12),
        ("Subtitle", 14, MID, 0, 10),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, "#1F4D78", 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(11)
        style.font.color.rgb = rgb(INK)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    doc.core_properties.title = "EFRM Rule Optimization Agent - Complete Enterprise Architecture and Orchestration"
    doc.core_properties.subject = "Enterprise EFRM Rule Optimization Agent architecture, orchestration, implementation, and controls"
    doc.core_properties.author = "Rule Optimization Architecture Team"
    doc.core_properties.keywords = "EFRM, rule optimization, false positive, Drools, backtesting, LLM"


# ---------------------------------------------------------------------------
# Document content helpers
# ---------------------------------------------------------------------------


def add_para(doc, text="", *, bold_lead=None, style=None, align=None, color=None, size=None,
             italic=False, keep=False):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, size=size or 11, color=color or INK, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2, size=size or 11, color=color or INK, italic=italic)
    else:
        r = p.add_run(text)
        set_run_font(r, size=size or 11, color=color or INK, italic=italic)
    if keep:
        set_keep_with_next(p)
    return p


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.5 + level * 0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        if isinstance(item, tuple):
            lead, rest = item
            r1 = p.add_run(lead)
            set_run_font(r1, bold=True)
            r2 = p.add_run(rest)
            set_run_font(r2)
        else:
            r = p.add_run(item)
            set_run_font(r)


def add_numbered(doc, items):
    for index, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.34)
        p.paragraph_format.first_line_indent = Inches(-0.24)
        p.paragraph_format.space_after = Pt(3)
        number = p.add_run(f"{index}. ")
        set_run_font(number, bold=True, color=NAVY)
        r = p.add_run(item)
        set_run_font(r)


def add_callout(doc, title, body, *, fill=PALE_BLUE, accent=BLUE, icon=None):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_grid(table, [180, 9180])
    left, right = table.rows[0].cells
    set_cell_width(left, 180)
    set_cell_width(right, 9180)
    shade_cell(left, accent)
    shade_cell(right, fill)
    for cell in (left, right):
        set_cell_margins(cell, top=110, bottom=110)
        set_cell_border(
            cell,
            top={"val": "single", "sz": "6", "color": BORDER},
            bottom={"val": "single", "sz": "6", "color": BORDER},
            left={"val": "single", "sz": "6", "color": BORDER},
            right={"val": "single", "sz": "6", "color": BORDER},
        )
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    left.text = ""
    p = right.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=NAVY, bold=True)
    p2 = right.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=9.7, color=INK)
    prevent_row_split(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_table(doc, headers, rows, *, widths=None, header_fill=NAVY, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    resolved_widths = widths or [int(9360 / len(headers))] * len(headers)
    resolved_widths[-1] += 9360 - sum(resolved_widths)
    set_table_grid(table, resolved_widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, heading in enumerate(headers):
        cell = header.cells[i]
        shade_cell(cell, header_fill)
        set_cell_margins(cell, top=100, bottom=100)
        set_cell_width(cell, resolved_widths[i])
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(heading))
        set_run_font(r, size=font_size, color=WHITE, bold=True)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    prevent_row_split(header)

    for ridx, row_values in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        fill = WHITE if ridx % 2 == 0 else LIGHT
        for i, value in enumerate(row_values):
            cell = row.cells[i]
            shade_cell(cell, fill)
            set_cell_margins(cell)
            set_cell_width(cell, resolved_widths[i])
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if isinstance(value, tuple):
                lead, rest = value
                r1 = p.add_run(str(lead))
                set_run_font(r1, size=font_size, color=INK, bold=True)
                r2 = p.add_run(str(rest))
                set_run_font(r2, size=font_size, color=INK)
            else:
                r = p.add_run(str(value))
                set_run_font(r, size=font_size, color=INK)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(
                cell,
                top={"val": "single", "sz": "4", "color": BORDER},
                bottom={"val": "single", "sz": "4", "color": BORDER},
                left={"val": "single", "sz": "4", "color": BORDER},
                right={"val": "single", "sz": "4", "color": BORDER},
            )
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_phase_card(doc, phase_no, title, purpose, component, input_text, sources, processing,
                   tech, llm, other_model, output, transition):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    set_keep_with_next(p)
    badge = p.add_run(f"PHASE {phase_no}  ")
    badge.font.highlight_color = None
    set_run_font(badge, size=10, color=TEAL, bold=True)
    title_run = p.add_run(title)
    set_run_font(title_run, size=15, color=NAVY, bold=True)

    add_callout(doc, "What this phase achieves", purpose, fill=PALE_TEAL, accent=TEAL)

    rows = [
        (("Main component: ", component),),
        (("Input: ", input_text),),
        (("Data comes from: ", sources),),
        (("Processing method: ", processing),),
        (("Concrete implementation: ", tech),),
        (("LLM used? ", llm),),
        (("Other AI/ML model? ", other_model),),
        (("Output: ", output),),
        (("How the next phase starts: ", transition),),
    ]
    table = doc.add_table(rows=0, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_grid(table, [9360])
    fills = [PALE_BLUE, WHITE, WHITE, LIGHT, WHITE, PALE_ORANGE, WHITE, PALE_GREEN, PALE_BLUE]
    for idx, rowdata in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        cell = row.cells[0]
        shade_cell(cell, fills[idx])
        set_cell_margins(cell, top=95, bottom=95, start=140, end=140)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        lead, rest = rowdata[0]
        r1 = p.add_run(lead)
        set_run_font(r1, size=9.6, color=NAVY, bold=True)
        r2 = p.add_run(rest)
        set_run_font(r2, size=9.6, color=INK)
        set_cell_border(
            cell,
            top={"val": "single", "sz": "4", "color": BORDER},
            bottom={"val": "single", "sz": "4", "color": BORDER},
            left={"val": "single", "sz": "4", "color": BORDER},
            right={"val": "single", "sz": "4", "color": BORDER},
        )
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code_block(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_grid(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, "16222C")
    set_cell_margins(cell, top=130, bottom=130, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(text.splitlines()):
        if idx:
            p.add_run().add_break()
        r = p.add_run(line)
        set_run_font(r, name="Consolas", size=8.7, color="E9F2F7")
    prevent_row_split(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_section_intro(doc, number, title, text):
    doc.add_heading(f"{number}. {title}", level=1)
    add_para(doc, text, size=11)


def add_part_banner(doc, label, title, subtitle, *, accent=TEAL, fill=PALE_TEAL):
    add_page_break(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(78)
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(label.upper())
    set_run_font(r, size=11, color=accent, bold=True)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(14)
    r = title_p.add_run(title)
    set_run_font(r, size=25, color=NAVY, bold=True)
    add_callout(doc, "What this part gives you", subtitle, fill=fill, accent=accent)


def add_toc(doc):
    doc.add_heading("Contents", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "Open this document in Microsoft Word and update fields to refresh the contents."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(display)
    run._r.append(end)
    set_run_font(run, size=10, color=MID)
    add_callout(
        doc,
        "Two reading paths",
        "Management path: read the Executive Snapshot and Sections 1-9. "
        "Technical approval path: continue through Sections 10-31 and the implementation appendices.",
        fill=PALE_BLUE,
        accent=BLUE,
    )


def add_figure(doc, image_path, caption, *, width=6.35):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    p.add_run().add_picture(str(image_path), width=Inches(width))
    cap = add_para(
        doc,
        caption,
        size=8.8,
        color=MID,
        italic=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        keep=True,
    )
    cap.paragraph_format.space_after = Pt(8)


# ---------------------------------------------------------------------------
# Diagram generation (Pillow only)
# ---------------------------------------------------------------------------


def pil_font(size, bold=False):
    path = FONT_BOLD if bold else FONT_REGULAR
    return ImageFont.truetype(str(path), size=size)


def draw_wrapped(draw, box, text, font, fill, *, align="center", spacing=6):
    x0, y0, x1, y1 = box
    max_width = max(20, x1 - x0 - 24)
    words = text.split()
    lines = []
    line = ""
    for word in words:
        trial = (line + " " + word).strip()
        if draw.textbbox((0, 0), trial, font=font)[2] <= max_width:
            line = trial
        else:
            if line:
                lines.append(line)
            line = word
    if line:
        lines.append(line)
    line_h = font.size + spacing
    total_h = len(lines) * line_h - spacing
    y = y0 + max(0, (y1 - y0 - total_h) / 2)
    for ln in lines:
        bbox = draw.textbbox((0, 0), ln, font=font)
        if align == "left":
            x = x0 + 14
        else:
            x = x0 + (x1 - x0 - (bbox[2] - bbox[0])) / 2
        draw.text((x, y), ln, font=font, fill=fill)
        y += line_h


def rounded_box(draw, box, fill, outline, title, subtitle=None, *, title_color=INK,
                subtitle_color=MID, radius=22):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=3)
    x0, y0, x1, y1 = box
    if subtitle:
        draw_wrapped(draw, (x0 + 10, y0 + 8, x1 - 10, y0 + 58), title,
                     pil_font(25, True), title_color)
        draw_wrapped(draw, (x0 + 14, y0 + 55, x1 - 14, y1 - 8), subtitle,
                     pil_font(20), subtitle_color)
    else:
        draw_wrapped(draw, box, title, pil_font(24, True), title_color)


def arrow(draw, start, end, color=BLUE, width=5):
    x0, y0 = start
    x1, y1 = end
    draw.line((x0, y0, x1, y1), fill=color, width=width)
    angle = math.atan2(y1 - y0, x1 - x0)
    head = 16
    for delta in (2.55, -2.55):
        hx = x1 + head * math.cos(angle + delta)
        hy = y1 + head * math.sin(angle + delta)
        draw.line((x1, y1, hx, hy), fill=color, width=width)


def make_architecture_diagram(path: Path):
    img = Image.new("RGB", (1800, 1180), WHITE)
    d = ImageDraw.Draw(img)
    d.text((70, 35), "Enterprise Rule Optimizer - logical architecture",
           font=pil_font(38, True), fill=NAVY)
    d.text((70, 88), "Existing EFRM remains authoritative. The optimizer is a separate decision-support boundary.",
           font=pil_font(23), fill=MID)

    # Existing EFRM lane
    d.rounded_rectangle((55, 150, 1745, 365), radius=28, fill=PALE_BLUE, outline=BLUE, width=3)
    d.text((85, 166), "EXISTING EFRM", font=pil_font(22, True), fill=BLUE)
    existing = [
        ((90, 220, 385, 325), "Case Service", "Final case outcome"),
        ((455, 220, 750, 325), "Admin / Config", "Rules, versions, policies"),
        ((820, 220, 1115, 325), "PostgreSQL", "Events, alerts, matches, cases"),
        ((1185, 220, 1480, 325), "Rule Engine", "Drools / DRL"),
    ]
    for box, title, sub in existing:
        rounded_box(d, box, WHITE, BORDER, title, sub, title_color=NAVY)
    rounded_box(d, (1520, 220, 1705, 325), WHITE, BORDER, "Compiler", "JSON to DRL", title_color=NAVY)

    # Optimizer lane
    d.rounded_rectangle((55, 405, 1745, 935), radius=28, fill="#F9FBFC", outline=TEAL, width=4)
    d.text((85, 425), "RULE OPTIMIZATION BOUNDARY", font=pil_font(22, True), fill=TEAL)
    top = [
        ((90, 490, 370, 605), "Trigger Gateway", "Event, API, recovery poll"),
        ((435, 490, 715, 605), "Orchestrator", "Durable 8-phase workflow"),
        ((780, 490, 1060, 605), "Data Resolver", "Lineage, config, snapshot"),
        ((1125, 490, 1405, 605), "Analytics Engine", "SQL + deterministic Python"),
        ((1470, 490, 1705, 605), "Evidence Store", "Immutable artifacts"),
    ]
    for box, title, sub in top:
        rounded_box(d, box, WHITE, BORDER, title, sub, title_color=NAVY)
    for idx in range(len(top) - 1):
        arrow(d, (top[idx][0][2] + 6, 548), (top[idx + 1][0][0] - 6, 548), TEAL, 5)

    bottom = [
        ((250, 690, 565, 825), "Optional LLM Gateway", "Evidence-grounded hypotheses and explanation"),
        ((665, 690, 980, 825), "Candidate + Test", "JSON candidate, compiler, isolated Drools replay"),
        ((1080, 690, 1395, 825), "Recommendation Only", "Structured advice; no executable candidate"),
    ]
    for box, title, sub in bottom:
        rounded_box(d, box, PALE_ORANGE if "LLM" in title else PALE_GREEN,
                    ORANGE if "LLM" in title else GREEN, title, sub, title_color=NAVY)
    arrow(d, (1265, 605), (410, 690), ORANGE, 4)
    arrow(d, (1265, 605), (820, 690), GREEN, 4)
    arrow(d, (1265, 605), (1235, 690), GREEN, 4)
    arrow(d, (565, 758), (665, 758), ORANGE, 4)

    rounded_box(
        d, (510, 995, 1290, 1115), PALE_TEAL, TEAL,
        "Human review and existing rule governance",
        "Only people and the existing governance process can change, retire, or activate a production rule.",
        title_color=NAVY
    )
    arrow(d, (820, 825), (760, 995), TEAL, 5)
    arrow(d, (1235, 825), (1040, 995), TEAL, 5)

    # Integration arrows from EFRM
    arrow(d, (235, 325), (235, 490), BLUE, 4)
    arrow(d, (600, 325), (920, 490), BLUE, 4)
    arrow(d, (967, 325), (970, 490), BLUE, 4)
    arrow(d, (1332, 325), (820, 690), BLUE, 4)
    arrow(d, (1612, 325), (880, 690), BLUE, 4)

    img.save(path, quality=95)


def make_workflow_diagram(path: Path):
    img = Image.new("RGB", (1800, 2140), WHITE)
    d = ImageDraw.Draw(img)
    d.text((70, 35), "Complete 8-phase workflow",
           font=pil_font(40, True), fill=NAVY)
    d.text((70, 90), "The false-positive case selects a rule. Full historical evidence determines what to recommend.",
           font=pil_font(24), fill=MID)

    phases = [
        ("1", "Trigger and confirm", "Receive the final case event. Verify the institution-specific outcome is false positive."),
        ("2", "Trace and attribute", "Follow case to alerts, results, matches, and rules. Identify real contributors."),
        ("3", "Resolve exact configuration", "Load historical and current rule versions, group context, metrics, and policy."),
        ("4", "Build trusted evidence", "Create a frozen historical snapshot and run data-quality/replayability gates."),
        ("5", "Analyze the rule", "Calculate false-positive burden, threshold patterns, overlap, stability, and leakage signals."),
        ("6", "Reason about changes", "Deterministic candidate search plus optional bounded LLM hypotheses and explanation."),
        ("7", "Follow configured mode", "Candidate + Test performs compile and replay. Recommendation Only skips executable candidate."),
        ("8", "Recommend and govern", "Publish evidence-linked recommendation. Human review controls every production action."),
    ]
    y = 165
    boxes = []
    for i, (num, title, desc) in enumerate(phases):
        fill = PALE_TEAL if i in (0, 7) else (PALE_ORANGE if i == 5 else PALE_BLUE)
        outline = TEAL if i in (0, 7) else (ORANGE if i == 5 else BLUE)
        box = (190, y, 1610, y + 155)
        boxes.append(box)
        d.rounded_rectangle(box, radius=26, fill=fill, outline=outline, width=4)
        d.ellipse((95, y + 35, 175, y + 115), fill=outline)
        num_box = d.textbbox((0, 0), num, font=pil_font(34, True))
        d.text((135 - (num_box[2] - num_box[0]) / 2, y + 53),
               num, font=pil_font(34, True), fill=WHITE)
        d.text((225, y + 24), title, font=pil_font(29, True), fill=NAVY)
        draw_wrapped(d, (225, y + 65, 1575, y + 140), desc, pil_font(22), INK, align="left")
        if i < len(phases) - 1:
            arrow(d, (900, y + 155), (900, y + 195), TEAL, 5)
        y += 195

    # Mode branch callout
    branch_y = 1745
    d.rounded_rectangle((250, branch_y, 860, branch_y + 280), radius=26,
                        fill=PALE_GREEN, outline=GREEN, width=4)
    d.text((285, branch_y + 24), "Mode A - Candidate + Test",
           font=pil_font(27, True), fill=NAVY)
    draw_wrapped(
        d, (285, branch_y + 70, 825, branch_y + 255),
        "Create a candidate in the optimizer store, validate JSON, call the existing compiler, replay the current baseline, replay the candidate, compare results, and recommend.",
        pil_font(22), INK, align="left"
    )

    d.rounded_rectangle((940, branch_y, 1550, branch_y + 280), radius=26,
                        fill=PALE_ORANGE, outline=ORANGE, width=4)
    d.text((975, branch_y + 24), "Mode B - Recommendation Only",
           font=pil_font(27, True), fill=NAVY)
    draw_wrapped(
        d, (975, branch_y + 70, 1515, branch_y + 255),
        "Do not create or compile an executable candidate. Produce a structured change suggestion from verified analysis and clearly state that no candidate backtest was performed.",
        pil_font(22), INK, align="left"
    )

    d.text((70, 2070), "Both modes end in human review. Neither mode writes to the production Rule Engine.",
           font=pil_font(25, True), fill=TEAL)
    img.save(path, quality=95)


def make_responsibility_diagram(path: Path):
    img = Image.new("RGB", (1800, 1050), WHITE)
    d = ImageDraw.Draw(img)
    d.text((70, 35), "Which technology does which job?",
           font=pil_font(40, True), fill=NAVY)
    d.text((70, 92), "Consistency comes from giving exact work to exact tools.",
           font=pil_font(24), fill=MID)

    columns = [
        (70, 380, "Backend / SQL", BLUE, PALE_BLUE,
         ["Authenticate and scope", "Read exact rows", "Trace lineage", "Resolve versions", "Store audit evidence"]),
        (505, 815, "Python analytics", TEAL, PALE_TEAL,
         ["Calculate metrics", "Find threshold patterns", "Measure trends and drift", "Estimate uncertainty", "Compare cohorts"]),
        (940, 1250, "Drools replay", GREEN, PALE_GREEN,
         ["Execute formal rule logic", "Replay current baseline", "Replay candidate", "Measure match changes", "Prove simulator parity"]),
        (1375, 1705, "Optional LLM", ORANGE, PALE_ORANGE,
         ["Interpret approved evidence", "Suggest bounded hypotheses", "Explain technical findings", "Draft readable recommendation", "Never invent metrics"]),
    ]
    for x0, x1, title, outline, fill, items in columns:
        d.rounded_rectangle((x0, 165, x1, 920), radius=28, fill=fill, outline=outline, width=4)
        draw_wrapped(d, (x0 + 10, 190, x1 - 10, 260), title, pil_font(28, True), NAVY)
        y = 310
        for item in items:
            d.ellipse((x0 + 28, y + 8, x0 + 44, y + 24), fill=outline)
            draw_wrapped(d, (x0 + 55, y, x1 - 22, y + 76), item, pil_font(22), INK, align="left")
            y += 105

    d.rounded_rectangle((260, 955, 1540, 1025), radius=18, fill=NAVY, outline=NAVY)
    draw_wrapped(
        d, (280, 960, 1520, 1020),
        "The orchestrator moves the job between these tools and records every input, result, and decision.",
        pil_font(24, True), WHITE
    )
    img.save(path, quality=95)


def make_trigger_diagram(path: Path):
    img = Image.new("RGB", (1800, 1120), WHITE)
    d = ImageDraw.Draw(img)
    d.text((70, 35), "Three trigger paths, one trusted orchestration entry",
           font=pil_font(39, True), fill=NAVY)
    d.text((70, 90), "All integrations normalize into the same authenticated, idempotent trigger contract.",
           font=pil_font(23), fill=MID)

    sources = [
        ((80, 175, 500, 330), "Required durable event path", "Case commit + outbox row -> publisher -> optional broker consumer"),
        ((690, 175, 1110, 330), "Supported API path", "Case Service calls a secured idempotent trigger endpoint after commit"),
        ((1300, 175, 1720, 330), "Recovery polling", "High-water mark + overlap window finds missed finalized decisions"),
    ]
    for box, title, subtitle in sources:
        rounded_box(d, box, PALE_BLUE, BLUE, title, subtitle, title_color=NAVY)

    rounded_box(
        d, (450, 435, 1350, 585), PALE_TEAL, TEAL,
        "Trigger Gateway",
        "Normalize schema, authenticate service identity, authorize institution, validate freshness, and assign correlation IDs.",
        title_color=NAVY
    )
    for box, _, _ in sources:
        arrow(d, ((box[0] + box[2]) // 2, box[3]), (900, 435), TEAL, 5)

    controls = [
        ((90, 700, 500, 840), "Transactional Inbox", "Persist before acknowledgement; deduplicate event/API retries"),
        ((695, 700, 1105, 840), "Policy + Outcome Resolver", "Load effective mapping; final case decision takes precedence"),
        ((1300, 700, 1710, 840), "Case Job Creator", "Create one case-analysis job or link to an eligible rule-analysis run"),
    ]
    for box, title, subtitle in controls:
        rounded_box(d, box, WHITE, BORDER, title, subtitle, title_color=NAVY)
    arrow(d, (900, 585), (295, 700), TEAL, 5)
    arrow(d, (500, 770), (695, 770), TEAL, 5)
    arrow(d, (1105, 770), (1300, 770), TEAL, 5)

    rounded_box(
        d, (495, 950, 1305, 1070), NAVY, NAVY,
        "Durable Orchestrator",
        "Starts the governed state machine; duplicate triggers return the existing identity.",
        title_color=WHITE, subtitle_color=WHITE
    )
    arrow(d, (1505, 840), (1050, 950), TEAL, 5)
    d.text((70, 1075), "Transactional outbox and inbox are required controls. Kafka and a separate workflow product are optional scale-out choices.",
           font=pil_font(20, True), fill=RED)
    img.save(path, quality=95)


def make_lineage_diagram(path: Path):
    img = Image.new("RGB", (1800, 1340), WHITE)
    d = ImageDraw.Draw(img)
    d.text((70, 35), "Case-to-rule lineage: the exact paths the resolver must follow",
           font=pil_font(38, True), fill=NAVY)
    d.text((70, 90), "The case-alert mapping is polymorphic. Alert, result, and match IDs are not trusted without institution-scoped validation.",
           font=pil_font(22), fill=MID)

    rounded_box(d, (650, 155, 1150, 260), PALE_TEAL, TEAL, "case_master", "Authoritative final decision + institution scope", title_color=NAVY)
    rounded_box(d, (590, 335, 1210, 455), PALE_BLUE, BLUE, "case_alert_mapping", "Route using alert_id + alert_type + alert_source_table", title_color=NAVY)
    arrow(d, (900, 260), (900, 335), TEAL, 5)

    # Transaction branch
    rounded_box(d, (90, 555, 500, 675), WHITE, BLUE, "transaction_alert", "Resolved through the discriminator", title_color=NAVY)
    rounded_box(d, (90, 755, 500, 875), WHITE, BLUE, "transaction_result", "Parent of alert and match", title_color=NAVY)
    rounded_box(d, (35, 970, 360, 1090), PALE_BLUE, BLUE, "transaction_match", "Stored rule reference/version", title_color=NAVY)
    rounded_box(d, (430, 970, 755, 1090), PALE_BLUE, BLUE, "transaction_master -> request", "Event, source, channel, institution cross-check", title_color=NAVY)
    arrow(d, (700, 455), (295, 555), BLUE, 5)
    arrow(d, (295, 675), (295, 755), BLUE, 5)
    arrow(d, (295, 875), (195, 970), BLUE, 5)
    arrow(d, (380, 875), (592, 970), BLUE, 5)

    # Device branch
    rounded_box(d, (1300, 555, 1710, 675), WHITE, GREEN, "device_alert", "Resolved through the discriminator", title_color=NAVY)
    rounded_box(d, (1300, 755, 1710, 875), WHITE, GREEN, "device_result", "Parent of alert and match", title_color=NAVY)
    rounded_box(d, (1045, 970, 1370, 1090), PALE_GREEN, GREEN, "device_match", "Stored rule reference/version", title_color=NAVY)
    rounded_box(d, (1440, 970, 1765, 1090), PALE_GREEN, GREEN, "device_master -> request", "Event, source, channel, institution cross-check", title_color=NAVY)
    arrow(d, (1100, 455), (1505, 555), GREEN, 5)
    arrow(d, (1505, 675), (1505, 755), GREEN, 5)
    arrow(d, (1420, 875), (1207, 970), GREEN, 5)
    arrow(d, (1590, 875), (1602, 970), GREEN, 5)

    rounded_box(
        d, (470, 1190, 1330, 1305), PALE_ORANGE, ORANGE,
        "Effective Configuration Resolver",
        "Resolve the historical rule reference, then load historical and current bundles. Deleted or ambiguous history stops candidate claims.",
        title_color=NAVY
    )
    arrow(d, (195, 1090), (670, 1190), ORANGE, 5)
    arrow(d, (1207, 1090), (1130, 1190), ORANGE, 5)
    d.text((70, 1125), "Attribution must reproduce rule-group and decision context. case_alert_mapping.is_primary_flag means primary alert, not primary rule.",
           font=pil_font(19, True), fill=RED)
    img.save(path, quality=95)


def make_state_machine_diagram(path: Path):
    img = Image.new("RGB", (1800, 1700), WHITE)
    d = ImageDraw.Draw(img)
    d.text((70, 35), "Durable orchestration state machine",
           font=pil_font(40, True), fill=NAVY)
    d.text((70, 92), "Every transition has a guard, idempotency key, persisted result, retry policy, and audit event.",
           font=pil_font(23), fill=MID)

    common = [
        "TRIGGER_RECEIVED", "TRIGGER_VALIDATED", "POLICY_RESOLVED",
        "OUTCOME_VERIFIED", "CASE_JOB_CREATED", "LINEAGE_RESOLVING",
        "CONFIG_RESOLVING", "ATTRIBUTING", "RULE_RUN_AGGREGATING",
        "SNAPSHOT_BUILDING", "METRICS_RECONSTRUCTING", "DATA_VALIDATING",
        "ANALYZING", "LLM_REASONING (optional)", "MODE_SELECTED",
    ]
    boxes = []
    x_positions = [85, 620, 1155]
    y = 175
    for idx, state in enumerate(common):
        col = idx % 3
        row = idx // 3
        x0 = x_positions[col]
        y0 = y + row * 150
        box = (x0, y0, x0 + 455, y0 + 92)
        boxes.append(box)
        fill = PALE_ORANGE if "LLM" in state else PALE_BLUE
        outline = ORANGE if "LLM" in state else BLUE
        rounded_box(d, box, fill, outline, state, title_color=NAVY, radius=16)
        if idx:
            prior = boxes[idx - 1]
            if col:
                arrow(d, (prior[2], (prior[1] + prior[3]) // 2), (box[0], (box[1] + box[3]) // 2), TEAL, 4)
            else:
                arrow(d, ((prior[0] + prior[2]) // 2, prior[3]), ((box[0] + box[2]) // 2, box[1]), TEAL, 4)

    branch_y = 980
    rounded_box(d, (95, branch_y, 820, branch_y + 275), PALE_ORANGE, ORANGE,
                "RECOMMENDATION ONLY",
                "PROPOSAL_BUILDING -> RECOMMENDATION_VALIDATING\nNo executable candidate, compiler call, or replay claim.",
                title_color=NAVY)
    rounded_box(d, (980, branch_y, 1705, branch_y + 275), PALE_GREEN, GREEN,
                "CANDIDATE + TEST",
                "CANDIDATE_GENERATING -> VALIDATING -> HISTORICAL_PARITY -> CURRENT_BASELINE -> CANDIDATE_BACKTEST -> EVALUATING",
                title_color=NAVY)
    arrow(d, (1382, 867), (457, branch_y), ORANGE, 5)
    arrow(d, (1382, 867), (1342, branch_y), GREEN, 5)

    rounded_box(d, (475, 1335, 1325, 1455), PALE_TEAL, TEAL,
                "STALE_STATE_CHECK -> REVIEW_REQUIRED",
                "Then human governance records APPROVED, REJECTED, MORE_ANALYSIS_REQUESTED, or CLOSED.",
                title_color=NAVY)
    arrow(d, (457, 1255), (700, 1335), TEAL, 5)
    arrow(d, (1342, 1255), (1100, 1335), TEAL, 5)

    d.rounded_rectangle((70, 1520, 1730, 1655), radius=22, fill=PALE_RED, outline=RED, width=3)
    draw_wrapped(
        d, (95, 1535, 1705, 1640),
        "Explicit terminal or limited states: DUPLICATE_TRIGGER | IGNORED_NOT_FALSE_POSITIVE | NO_OPTIMIZABLE_RULE | "
        "UNRESOLVED_ATTRIBUTION | CONFIGURATION_AMBIGUITY | ELIGIBLE_POPULATION_NOT_RECONSTRUCTABLE | "
        "METRIC_NOT_REPRODUCIBLE | INSUFFICIENT_EVIDENCE | CANDIDATE_INVALID | REPLAY_MISMATCH | STALE | SUPERSEDED | CANCELLED | FAILED",
        pil_font(19, True), RED, align="left"
    )
    img.save(path, quality=95)


def make_evidence_pipeline_diagram(path: Path):
    img = Image.new("RGB", (1800, 1120), WHITE)
    d = ImageDraw.Draw(img)
    d.text((70, 35), "Historical evidence pipeline",
           font=pil_font(40, True), fill=NAVY)
    d.text((70, 92), "False-positive cases select the rule; a full frozen population evaluates it.",
           font=pil_font(24), fill=MID)

    stages = [
        ((60, 250, 310, 430), "EFRM read boundary", "Approved APIs, replica, or analytical store"),
        ((355, 250, 605, 430), "Reviewed queries", "Typed templates + institution and time parameters"),
        ((650, 250, 900, 430), "Eligibility builder", "Fired + non-fired events; source/channel/group context"),
        ((945, 250, 1195, 430), "Metric SQL executor", "Point-in-time Custom SQL under strict controls"),
        ((1240, 250, 1490, 430), "Frozen snapshot", "Parquet partitions + manifest + checksums"),
        ((1535, 250, 1785, 430), "Data-quality gate", "Ready, limited, or explicit stop state"),
    ]
    for box, title, subtitle in stages:
        rounded_box(d, box, WHITE, BORDER, title, subtitle, title_color=NAVY)
    for idx in range(len(stages) - 1):
        arrow(d, (stages[idx][0][2], 340), (stages[idx + 1][0][0], 340), TEAL, 4)

    controls = [
        ((100, 560, 500, 730), "Scope controls", "Institution authorization, source/channel/fact, analysis cutoff, timezone normalization"),
        ((700, 560, 1100, 730), "Time-travel controls", "Execution manifest, immutable rule/metric/policy versions, event-time inputs, label maturity"),
        ((1300, 560, 1700, 730), "Workload controls", "Read-only role, allowlisted schemas/functions, timeouts, row/memory/plan limits"),
    ]
    for box, title, subtitle in controls:
        rounded_box(d, box, PALE_BLUE, BLUE, title, subtitle, title_color=NAVY)

    rounded_box(d, (350, 875, 1450, 1035), PALE_GREEN, GREEN,
                "Deterministic Analysis Evidence",
                "Counts, denominators, rates, threshold bands, drift, overlap, alert-volume impact, uncertainty, and limitations.",
                title_color=NAVY)
    arrow(d, (1660, 430), (1100, 875), GREEN, 5)
    d.text((70, 1065), "No LLM writes SQL, chooses joins, or changes the frozen cohort.",
           font=pil_font(22, True), fill=TEAL)
    img.save(path, quality=95)


def make_llm_loop_diagram(path: Path):
    img = Image.new("RGB", (1800, 1260), WHITE)
    d = ImageDraw.Draw(img)
    d.text((70, 35), "Bounded LLM reasoning loop",
           font=pil_font(40, True), fill=NAVY)
    d.text((70, 92), "The LLM may interpret and ask for approved calculations; it never becomes the source of facts.",
           font=pil_font(23), fill=MID)

    top = [
        ((70, 195, 410, 350), "Evidence Package", "Masked aggregates, candidate schema, evidence IDs, limitations"),
        ((515, 195, 855, 350), "Prompt Builder", "Policy, allowed actions, iteration/token budget, output schema"),
        ((960, 195, 1300, 350), "Model Gateway", "Local model or approved external API; provider-neutral"),
        ((1405, 195, 1730, 350), "LLM Reasoner", "Hypothesis, explanation, or typed tool request"),
    ]
    for box, title, subtitle in top:
        rounded_box(d, box, PALE_ORANGE if "LLM" in title or "Model" in title else PALE_BLUE,
                    ORANGE if "LLM" in title or "Model" in title else BLUE,
                    title, subtitle, title_color=NAVY)
    for idx in range(len(top) - 1):
        arrow(d, (top[idx][0][2], 273), (top[idx + 1][0][0], 273), ORANGE, 5)

    bottom = [
        ((1130, 520, 1670, 690), "Approved Tool Broker", "Accept only typed, allowlisted requests; attach caller, scope, timeout, and audit ID"),
        ((630, 520, 1030, 690), "Deterministic Tools", "Cohort slice, threshold sensitivity, overlap, statistics; no arbitrary SQL"),
        ((130, 520, 530, 690), "Exact Tool Result", "Structured numbers, denominators, uncertainty, evidence references"),
    ]
    for box, title, subtitle in bottom:
        rounded_box(d, box, PALE_BLUE, BLUE, title, subtitle, title_color=NAVY)
    arrow(d, (1567, 350), (1400, 520), ORANGE, 5)
    arrow(d, (1130, 605), (1030, 605), BLUE, 5)
    arrow(d, (630, 605), (530, 605), BLUE, 5)
    arrow(d, (330, 520), (1450, 350), TEAL, 4)

    rounded_box(d, (310, 820, 1490, 980), PALE_GREEN, GREEN,
                "Structured Reasoning Result -> Claim Validator",
                "Unknown fields, unsupported numbers, missing evidence IDs, forbidden operations, and ungrounded causal claims are rejected. "
                "A deterministic template is used if the LLM fails.",
                title_color=NAVY)
    arrow(d, (1500, 350), (1100, 820), GREEN, 5)

    d.rounded_rectangle((90, 1065, 1710, 1195), radius=22, fill=PALE_RED, outline=RED, width=3)
    draw_wrapped(
        d, (120, 1075, 1680, 1185),
        "Blocked: raw database credentials | arbitrary SQL | production-rule write tools | secrets | raw customer/account/card/device data | "
        "unapproved documents | self-reported confidence as evidence",
        pil_font(21, True), RED, align="center"
    )
    img.save(path, quality=95)


def make_replay_diagram(path: Path):
    img = Image.new("RGB", (1800, 1220), WHITE)
    d = ImageDraw.Draw(img)
    d.text((70, 35), "The correct three-part replay and backtesting model",
           font=pil_font(39, True), fill=NAVY)
    d.text((70, 92), "Parity proves the runner; the current baseline and candidate then use the same frozen cohort and context.",
           font=pil_font(22), fill=MID)

    panels = [
        ((65, 185, 565, 785), "A. Historical parity replay", BLUE, PALE_BLUE,
         ["Use configuration effective for each event", "Preserve event order and state", "Compare stored match/signal result", "Fail with REPLAY_MISMATCH if outside approved tolerance"]),
        ((650, 185, 1150, 785), "B. Current baseline", TEAL, PALE_TEAL,
         ["Use today's current target configuration", "Run on frozen evaluation cohort", "Execute full group/decision context when available", "Store baseline metrics and evidence"]),
        ((1235, 185, 1735, 785), "C. Candidate simulation", GREEN, PALE_GREEN,
         ["Use the identical cohort", "Replace only the selected target rule", "Preserve all other group/policy context", "Store candidate metrics, diff, and limitations"]),
    ]
    for box, title, outline, fill, items in panels:
        d.rounded_rectangle(box, radius=28, fill=fill, outline=outline, width=4)
        draw_wrapped(d, (box[0] + 15, box[1] + 20, box[2] - 15, box[1] + 95),
                     title, pil_font(26, True), NAVY)
        y = box[1] + 135
        for item in items:
            d.ellipse((box[0] + 28, y + 8, box[0] + 46, y + 26), fill=outline)
            draw_wrapped(d, (box[0] + 60, y, box[2] - 22, y + 88),
                         item, pil_font(21), INK, align="left")
            y += 105

    arrow(d, (900, 785), (740, 910), TEAL, 5)
    arrow(d, (1485, 785), (1060, 910), GREEN, 5)
    rounded_box(d, (430, 910, 1370, 1065), WHITE, BORDER,
                "Deterministic Comparator",
                "Compare alert volume, mature internal outcomes, overlap, leakage indicators, uncertainty, holdout performance, and constraint status.",
                title_color=NAVY)
    d.text((70, 1110), "Claim rule: selected-rule-only replay supports match-level claims. Alert/case claims require full downstream context.",
           font=pil_font(22, True), fill=RED)
    img.save(path, quality=95)


def make_security_diagram(path: Path):
    img = Image.new("RGB", (1800, 1260), WHITE)
    d = ImageDraw.Draw(img)
    d.text((70, 35), "Security and deployment trust zones",
           font=pil_font(40, True), fill=NAVY)
    d.text((70, 92), "Product names are proposed reference choices; the trust boundaries and permissions are mandatory design requirements.",
           font=pil_font(22), fill=MID)

    zones = [
        ((55, 175, 520, 585), "Zone 1 - Existing EFRM", BLUE, PALE_BLUE,
         ["Case/Admin/Rule services", "Operational PostgreSQL", "Production Rule Engine", "No optimizer write credentials"]),
        ((590, 175, 1210, 585), "Zone 2 - Optimizer control plane", TEAL, PALE_TEAL,
         ["Trigger API/consumer", "Durable orchestrator", "Policy/config adapters", "Recommendation API", "Institution-scoped authorization"]),
        ((1280, 175, 1745, 585), "Zone 3 - Isolated execution", GREEN, PALE_GREEN,
         ["Python analytics workers", "Compiler adapter", "Java/Drools replay workers", "No production activation path"]),
        ((90, 730, 600, 1115), "Zone 4 - Evidence storage", NAVY, "#EDF2F6",
         ["Optimizer PostgreSQL metadata", "Immutable Parquet/object storage", "Encryption + prefix isolation", "Checksums, retention, legal hold"]),
        ((685, 730, 1195, 1115), "Zone 5 - Model boundary", ORANGE, PALE_ORANGE,
         ["Provider-neutral gateway", "Local or approved external LLM", "Masking/DLP/egress controls", "No training/retention terms", "Emergency disable switch"]),
        ((1280, 730, 1710, 1115), "Zone 6 - Human governance", RED, PALE_RED,
         ["Reviewer RBAC", "Four-eyes approval", "Stale-state check", "Existing governed production change"]),
    ]
    for box, title, outline, fill, items in zones:
        d.rounded_rectangle(box, radius=26, fill=fill, outline=outline, width=4)
        draw_wrapped(d, (box[0] + 10, box[1] + 18, box[2] - 10, box[1] + 82),
                     title, pil_font(25, True), NAVY)
        y = box[1] + 105
        for item in items:
            d.ellipse((box[0] + 25, y + 7, box[0] + 41, y + 23), fill=outline)
            draw_wrapped(d, (box[0] + 52, y, box[2] - 18, y + 62),
                         item, pil_font(19), INK, align="left")
            y += 64

    arrow(d, (520, 380), (590, 380), TEAL, 5)
    arrow(d, (1210, 380), (1280, 380), GREEN, 5)
    arrow(d, (900, 585), (350, 730), NAVY, 4)
    arrow(d, (900, 585), (940, 730), ORANGE, 4)
    arrow(d, (1500, 585), (1495, 730), RED, 4)
    d.text((70, 1185), "Mandatory across every zone: encryption, service identity, least privilege, institution scope, audit correlation, environment separation, and signed artifacts.",
           font=pil_font(20, True), fill=TEAL)
    img.save(path, quality=95)


def make_recommendation_ui_diagram(path: Path):
    img = Image.new("RGB", (1800, 1260), "#F3F6F8")
    d = ImageDraw.Draw(img)
    d.rounded_rectangle((55, 45, 1745, 1215), radius=26, fill=WHITE, outline=BORDER, width=3)
    d.rounded_rectangle((55, 45, 1745, 155), radius=26, fill=NAVY, outline=NAVY, width=3)
    d.text((95, 75), "Rule Optimization Recommendation", font=pil_font(32, True), fill=WHITE)
    d.rounded_rectangle((1430, 73, 1700, 132), radius=16, fill=PALE_ORANGE, outline=ORANGE, width=2)
    draw_wrapped(d, (1440, 76, 1690, 130), "HUMAN REVIEW REQUIRED", pil_font(18, True), ORANGE)

    rounded_box(d, (90, 205, 725, 405), PALE_BLUE, BLUE,
                "Rule R-204 | Current version 6",
                "Issue: false-positive concentration among established customers using trusted beneficiaries during device days 4-7.",
                title_color=NAVY)
    rounded_box(d, (770, 205, 1710, 405), PALE_GREEN, GREEN,
                "Recommended action",
                "Review reducing the new-device window from 7 days to 3 days. Candidate + Test evidence is shown only when replay parity and full-context simulation pass.",
                title_color=NAVY)

    d.text((95, 465), "Evidence snapshot", font=pil_font(25, True), fill=NAVY)
    cards = [
        ((95, 515, 465, 675), "Mode", "CANDIDATE + TEST", TEAL, PALE_TEAL),
        ((505, 515, 875, 675), "Evidence strength", "SUFFICIENT FOR REVIEW", BLUE, PALE_BLUE),
        ((915, 515, 1285, 675), "Projected alerts", "-28 on frozen cohort", GREEN, PALE_GREEN),
        ((1325, 515, 1695, 675), "Coverage change", "-1 mature internal outcome", ORANGE, PALE_ORANGE),
    ]
    for box, label, value, outline, fill in cards:
        d.rounded_rectangle(box, radius=20, fill=fill, outline=outline, width=3)
        draw_wrapped(d, (box[0] + 8, box[1] + 15, box[2] - 8, box[1] + 58),
                     label.upper(), pil_font(17, True), outline)
        draw_wrapped(d, (box[0] + 12, box[1] + 62, box[2] - 12, box[3] - 12),
                     value, pil_font(23, True), NAVY)

    rounded_box(d, (95, 740, 1110, 950), WHITE, BORDER,
                "Why this is recommended",
                "Three evidence-linked findings explain the pattern. Each number opens the underlying metric, cohort definition, snapshot hash, and query/code version.",
                title_color=NAVY)
    rounded_box(d, (1150, 740, 1695, 950), PALE_RED, RED,
                "Limitations",
                "Hard optimization constraints are incomplete. Example values are fictional. Production action is not permitted from this screen.",
                title_color=NAVY)

    d.rounded_rectangle((95, 1030, 450, 1110), radius=15, fill=WHITE, outline=BLUE, width=3)
    draw_wrapped(d, (105, 1035, 440, 1105), "VIEW EVIDENCE", pil_font(20, True), BLUE)
    d.rounded_rectangle((490, 1030, 880, 1110), radius=15, fill=WHITE, outline=ORANGE, width=3)
    draw_wrapped(d, (500, 1035, 870, 1105), "REQUEST MORE ANALYSIS", pil_font(20, True), ORANGE)
    d.rounded_rectangle((920, 1030, 1265, 1110), radius=15, fill=WHITE, outline=RED, width=3)
    draw_wrapped(d, (930, 1035, 1255, 1105), "REJECT", pil_font(20, True), RED)
    d.rounded_rectangle((1305, 1030, 1695, 1110), radius=15, fill=TEAL, outline=TEAL, width=3)
    draw_wrapped(d, (1315, 1035, 1685, 1105), "SEND TO GOVERNANCE", pil_font(20, True), WHITE)

    d.text((95, 1155), "Illustrative reviewer contract. Reported portal routes require owner verification and secured optimizer APIs.",
           font=pil_font(20, True), fill=MID)
    img.save(path, quality=95)


# ---------------------------------------------------------------------------
# Document build
# ---------------------------------------------------------------------------


def build_document():
    architecture_png = BUILD_DIR / "architecture.png"
    workflow_png = BUILD_DIR / "workflow.png"
    responsibility_png = BUILD_DIR / "responsibility.png"
    trigger_png = BUILD_DIR / "trigger_convergence.png"
    lineage_png = BUILD_DIR / "case_rule_lineage.png"
    state_machine_png = BUILD_DIR / "orchestration_state_machine.png"
    evidence_pipeline_png = BUILD_DIR / "historical_evidence_pipeline.png"
    llm_loop_png = BUILD_DIR / "bounded_llm_loop.png"
    replay_png = BUILD_DIR / "replay_model.png"
    security_png = BUILD_DIR / "security_deployment_zones.png"
    recommendation_ui_png = BUILD_DIR / "recommendation_ui_concept.png"
    make_architecture_diagram(architecture_png)
    make_workflow_diagram(workflow_png)
    make_responsibility_diagram(responsibility_png)
    make_trigger_diagram(trigger_png)
    make_lineage_diagram(lineage_png)
    make_state_machine_diagram(state_machine_png)
    make_evidence_pipeline_diagram(evidence_pipeline_png)
    make_llm_loop_diagram(llm_loop_png)
    make_replay_diagram(replay_png)
    make_security_diagram(security_png)
    make_recommendation_ui_diagram(recommendation_ui_png)

    doc = Document()
    configure_document(doc)

    # Cover
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(58)
    p.paragraph_format.space_after = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AGENTIC RULE OPTIMIZER")
    set_run_font(r, size=12, color=TEAL, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    r = title.add_run("EFRM Rule Optimization Agent\nComplete Enterprise Architecture and Orchestration")
    set_run_font(r, size=27, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    r = subtitle.add_run(
        "A management-friendly explanation and detailed technical design for false-positive-triggered "
        "rule analysis, optional candidate backtesting, and human-governed recommendations"
    )
    set_run_font(r, size=13, color=MID)

    cover_table = doc.add_table(rows=1, cols=3)
    cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cover_table.autofit = False
    set_table_grid(cover_table, [3120, 3120, 3120])
    labels = [
        ("TRIGGER", "Final false-positive case"),
        ("DECISION SUPPORT", "Evidence-based recommendation"),
        ("CONTROL", "Human governance only"),
    ]
    for i, (label, value) in enumerate(labels):
        cell = cover_table.cell(0, i)
        set_cell_width(cell, 3120)
        shade_cell(cell, PALE_TEAL if i != 1 else PALE_BLUE)
        set_cell_margins(cell, top=170, bottom=170, start=110, end=110)
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_after = Pt(5)
        r1 = p1.add_run(label)
        set_run_font(r1, size=9, color=TEAL if i != 1 else BLUE, bold=True)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(value)
        set_run_font(r2, size=10, color=NAVY, bold=True)
        set_cell_border(
            cell,
            top={"val": "single", "sz": "6", "color": BORDER},
            bottom={"val": "single", "sz": "6", "color": BORDER},
            left={"val": "single", "sz": "6", "color": BORDER},
            right={"val": "single", "sz": "6", "color": BORDER},
        )

    add_para(
        doc,
        "Audience: stakeholders, managers, product owners, business analysts, fraud teams, architects, "
        "and new team members.",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=10.5,
        color=MID,
        italic=True,
    )
    add_para(
        doc,
        "Architecture status: refined enterprise logical design and proposed reference implementation; not yet implemented.",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=9.5,
        color=MID,
    )
    add_para(
        doc,
        "All example names, IDs, thresholds, and results in this document are fictional. "
        "They do not represent the supplied sample dump or production behavior.",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=9,
        color=RED,
        italic=True,
    )

    add_page_break(doc)

    # Executive snapshot
    doc.add_heading("Executive Snapshot", level=1)
    add_callout(
        doc,
        "Short answer for non-technical stakeholders",
        "When a finalized case is mapped to false positive, the Rule Optimization Agent finds the rule that contributed, "
        "studies the rule against full historical evidence, optionally tests a safe candidate outside production, and gives "
        "a human reviewer one evidence-linked recommendation. It never changes a production rule by itself.",
        fill=PALE_TEAL,
        accent=TEAL,
    )
    add_table(
        doc,
        ["Management question", "Simple answer"],
        [
            ("What starts it?", "A final case decision that the institution's configured mapping classifies as FALSE_POSITIVE."),
            ("Does it analyze every rule?", "No. The case selects possible contributing rules; full eligible history evaluates only those rules."),
            ("Who calculates the facts?", "Governed backend queries, deterministic Python/statistics, and isolated Drools replay."),
            ("Where can an LLM help?", "Optional hypothesis formation and plain-English explanation after facts have been calculated."),
            ("What does it produce?", "A recommendation containing the issue, proposed human action, evidence, test status, and limitations."),
            ("Can it change production?", "Never. Existing human governance owns every change, retirement, or deletion."),
        ],
        widths=[3000, 6360],
        font_size=8.9,
    )
    doc.add_heading("The complete behavior in 8 simple steps", level=2)
    add_numbered(doc, [
        "A case is finalized, and the final decision is interpreted using the institution's versioned outcome mapping.",
        "The Agent is notified by an event or API call, then verifies the authoritative case state and prevents duplicate processing.",
        "It traces the case to alerts, results, matches, and exact rule references, and attributes only rules that actually contributed.",
        "It rebuilds the historical and current rule configuration with version identifiers, hashes, timestamps, and provenance.",
        "It gathers the full eligible historical population and checks data quality, label maturity, reproducibility, and replayability.",
        "Deterministic analytics calculates the evidence; an optional LLM may form bounded hypotheses and explain verified findings.",
        "The institution's configured mode selects either Recommendation Only or Candidate + Test, including historical simulation where permitted.",
        "The Agent validates and publishes one evidence-linked recommendation; a human and the existing governance process control every production action.",
    ])
    doc.add_heading("What management is being asked to approve", level=2)
    add_bullets(doc, [
        "False-positive case closure is the normal selection trigger; alert-only triggering remains out of scope unless separately approved.",
        "Full historical evidence, not one case, is the basis for analysis.",
        "Historical reproducibility is a Phase 0 prerequisite: case outbox, canonical outcome versions, immutable execution manifests, "
        "versioned metric and decision-policy definitions, retention, and tenant controls must exist before historical replay claims are allowed.",
        "Both operating modes are configurable per institution.",
        "The optimizer has read-only access to existing EFRM production data and rules.",
        "Authoritative calculations remain deterministic; LLM reasoning is optional and bounded.",
        "Candidate artifacts and simulations remain outside production.",
        "Human governance is mandatory for every production action.",
    ])

    add_page_break(doc)
    add_toc(doc)
    add_part_banner(
        doc,
        "Part I",
        "Stakeholder Guide",
        "A simple explanation of the problem, the running example, the complete workflow, the two modes, "
        "and exactly where deterministic logic, Drools, the LLM, and people are used.",
        accent=TEAL,
        fill=PALE_TEAL,
    )

    # How to use
    add_section_intro(
        doc, "0", "How to read this document",
        "This guide explains the system from the outside in. It first explains the business problem, "
        "then follows one fictional case through the complete architecture. Every phase states the input, "
        "component, processing method, technology, LLM use, output, and hand-off."
    )
    add_callout(
        doc,
        "The one idea to remember",
        "A false-positive case tells the optimizer which rule deserves attention. It does not prove the rule is bad. "
        "The optimizer uses the full eligible historical population to decide whether a change is supported.",
        fill=PALE_TEAL,
        accent=TEAL,
    )
    doc.add_heading("What this document covers", level=2)
    add_table(
        doc,
        ["Part", "Question answered"],
        [
            ("1-3", "What problem are we solving, and what does 'agentic' mean?"),
            ("4-5", "What is the architecture, and which component owns each job?"),
            ("6", "What happens in each of the eight phases?"),
            ("7-9", "How do the two modes, backtesting, LLM, and RAG work?"),
            ("10-13", "What is the output, what data is used, and how is the system governed?"),
            ("14-15", "What framework is confirmed, what is proposed, and how do we implement it?"),
            ("16-30", "What are the detailed orchestration, contracts, analytics, security, operations, and tests?"),
            ("31", "How does one complete example move through the entire system?"),
        ],
        widths=[1500, 7860],
    )
    add_page_break(doc)
    doc.add_heading("Words used in this guide", level=2)
    add_table(
        doc,
        ["Term", "Simple meaning"],
        [
            ("Rule", "A formal if/then condition used by the existing Rule Engine."),
            ("False positive", "A case whose final institution-mapped outcome says the alert was not useful."),
            ("Agent", "The complete controlled workflow that can gather evidence, choose permitted next steps, and produce a recommendation."),
            ("LLM", "A language model used only for bounded reasoning and clear explanation."),
            ("Backtest", "Running current or proposed rule logic on frozen historical inputs."),
            ("Candidate", "A proposed rule change stored outside production for testing."),
            ("Evidence", "Traceable rows, configuration snapshots, metrics, test results, and limitations that support a recommendation."),
        ],
        widths=[1900, 7460],
    )

    add_page_break(doc)

    # Problem
    add_section_intro(
        doc, "1", "The problem in one minute",
        "The existing EFRM Rule Engine already evaluates transaction and device events and creates alerts. "
        "The Rule Optimizer does not replace that engine and does not generate live alerts. Its job starts after "
        "an investigator has completed a case and the final outcome is classified as false positive."
    )
    doc.add_heading("Binary rule behavior does not mean binary rule quality", level=2)
    add_para(
        doc,
        "At runtime, a rule has a binary result: it matches or it does not match. But historical rule quality is a "
        "business and statistical question. A matching rule can still create too many false-positive cases, overlap "
        "with other rules, depend on poor data, or lose useful coverage after a threshold change."
    )
    add_table(
        doc,
        ["Live Rule Engine question", "Rule Optimizer question"],
        [
            ("Did the rule match this event?", "Across a reliable historical population, is this rule producing useful and stable results?"),
            ("Should the signal contribute to an alert?", "Which matched rule actually contributed to the false-positive case?"),
            ("What is the current rule logic?", "Would a controlled change reduce burden without unacceptable loss of mature non-false-positive outcomes?"),
        ],
        widths=[3900, 5460],
    )
    doc.add_heading("What the Agent will do", level=2)
    add_bullets(doc, [
        "Start when a finalized case maps to the canonical outcome FALSE_POSITIVE.",
        "Trace the case to alerts, results, rule matches, and exact rule versions.",
        "Select only rules that were primary or supporting contributors.",
        "Evaluate those selected rules using full eligible historical evidence.",
        "Diagnose threshold, overlap, data-quality, stability, and outcome patterns.",
        "Use the configured operating mode to test a candidate or provide recommendation only.",
        "Publish one evidence-linked recommendation for human review.",
    ])
    doc.add_heading("What the Agent will never do", level=2)
    add_bullets(doc, [
        "Generate live transaction alerts.",
        "Scan and optimize every rule merely because a schedule ran.",
        "Treat one false-positive case as proof that a rule is defective.",
        "Let an LLM directly query the database or invent metrics.",
        "Write a candidate into production rule tables.",
        "Modify, activate, retire, or delete a rule.",
        "Replace the existing approval and governance process.",
    ])

    # Running example
    add_page_break(doc)
    add_section_intro(
        doc, "2", "Running example used throughout",
        "The same fictional example is followed through every phase so the reader can see how the technical pieces connect."
    )
    add_callout(
        doc,
        "Fictional case FP-1042",
        "Example Bank has a mobile-transfer rule R-204 called 'Large transfer from a new device'. "
        "The current example logic is: amount >= 100,000 AND device age <= 7 days. "
        "A long-standing customer changes phone and sends money to a trusted beneficiary. The rule contributes "
        "to an alert, a case is created, and the investigator closes the final case as false positive.",
        fill=PALE_ORANGE,
        accent=ORANGE,
    )
    add_table(
        doc,
        ["Example item", "Illustrative value"],
        [
            ("Institution", "Example Bank"),
            ("Channel", "Mobile transfer"),
            ("Case", "FP-1042"),
            ("Alert", "A-781"),
            ("Matched rule", "R-204, historical version 5"),
            ("Current condition", "amount >= 100,000 AND device_age_days <= 7"),
            ("Final outcome", "Institution decision code mapped to canonical FALSE_POSITIVE"),
            ("Possible hypothesis", "The seven-day new-device window may be too broad for trusted beneficiaries."),
        ],
        widths=[2600, 6760],
    )
    add_para(
        doc,
        "Important: the hypothesis is not yet a recommendation. The system must first prove the lineage, "
        "load the correct configuration, analyze the full historical cohort, and check whether the issue still exists "
        "in the current active rule."
    )

    add_page_break(doc)

    # Agentic
    add_section_intro(
        doc, "3", "What makes this an agent",
        "The word 'agent' describes the complete system behavior, not only the presence of an LLM. "
        "A durable orchestrator receives a goal, calls approved tools, checks results, chooses the next allowed action, "
        "stops when evidence is unsafe, and records an auditable outcome."
    )
    add_code_block(
        doc,
        "FALSE-POSITIVE TRIGGER\n"
        "    -> gather exact evidence\n"
        "    -> check lineage and sufficiency\n"
        "    -> calculate deterministic findings\n"
        "    -> choose the configured mode\n"
        "    -> test a candidate OR prepare advice\n"
        "    -> validate every claim\n"
        "    -> publish recommendation for human review"
    )
    doc.add_heading("The Agent is a team of controlled components", level=2)
    add_table(
        doc,
        ["Component", "Job"],
        [
            ("Orchestrator", "Controls the eight phases, retries, stop gates, mode branch, and audit state."),
            ("Data services", "Retrieve exact case, alert, match, rule, configuration, and historical records."),
            ("Analytics engine", "Uses SQL and deterministic Python to calculate trusted evidence."),
            ("Candidate and replay services", "Create formal candidates and execute isolated Drools backtests in Candidate + Test mode."),
            ("Optional LLM", "Interprets approved evidence, proposes bounded hypotheses, and explains the result in plain English."),
            ("Validator", "Checks that every number and claim can be traced to evidence."),
            ("Human governance", "Accepts, rejects, or changes the recommendation and performs any production action."),
        ],
        widths=[2600, 6760],
    )
    add_callout(
        doc,
        "Why this is more consistent than asking an LLM to do everything",
        "Databases, SQL, statistics, and the Drools engine give repeatable results from the same inputs. "
        "The LLM adds language reasoning where it is useful, but it is prevented from becoming the source of truth.",
        fill=PALE_GREEN,
        accent=GREEN,
    )

    # Architecture
    add_page_break(doc)
    add_section_intro(
        doc, "4", "Complete architecture",
        "The optimizer is deployed beside the existing EFRM platform, not inside the live transaction path. "
        "It reads governed information, stores its own immutable evidence, and sends recommendations to human review."
    )
    add_figure(
        doc,
        architecture_png,
        "Figure 1. Logical architecture. Boxes represent responsibilities, not mandatory microservices. "
        "Product names such as Kafka and Temporal are optional reference choices, not confirmed current EFRM infrastructure.",
        width=6.35,
    )
    doc.add_heading("The key boundaries", level=2)
    add_bullets(doc, [
        ("Existing EFRM stays authoritative. ", "Cases, rules, matches, and production execution remain owned by existing services."),
        ("The service is institution- and channel-agnostic. ", "Outcome mappings, thresholds, windows, modes, constraints, and analysis policy come from versioned configuration rather than hardcoded client logic."),
        ("The optimizer has read-only EFRM access. ", "Bulk history should come from approved APIs, a read replica, or an analytical store."),
        ("Candidates stay outside production. ", "Even in Candidate + Test mode, the proposed rule is an optimizer artifact."),
        ("The LLM is behind a gateway. ", "The gateway masks data, limits tools and tokens, and supports either a local model or an approved external API."),
        ("The recommendation goes to people. ", "Only the existing governance process can create or promote a production version."),
    ])

    # Technology responsibilities
    add_section_intro(
        doc, "5", "Which technology performs the analysis?",
        "No single technology is good at every job. The design assigns each kind of processing to the component that "
        "can do it most accurately and repeatedly."
    )
    add_figure(
        doc,
        responsibility_png,
        "Figure 2. Division of responsibility between normal backend logic, analytics, rule replay, and the optional LLM.",
        width=6.35,
    )
    add_table(
        doc,
        ["Question", "Authoritative component", "Why"],
        [
            ("Which case outcome is final?", "Backend lookup + institution mapping", "The answer must match exact governed records."),
            ("Which rule contributed?", "Lineage and attribution logic", "Relationships, versions, and decision policy must be repeatable."),
            ("What are the metrics?", "SQL + deterministic Python", "Counts, rates, trends, and intervals must reproduce exactly."),
            ("What would the rule match?", "Isolated Drools replay", "The same formal rule semantics are required."),
            ("What might explain the pattern?", "Optional LLM over verified evidence", "Language reasoning can connect patterns and produce understandable hypotheses."),
            ("Should production change?", "Human governance", "This is a controlled business and risk decision."),
        ],
        widths=[2800, 3000, 3560],
        font_size=8.7,
    )

    # Workflow overview
    add_page_break(doc)
    add_section_intro(
        doc, "6", "The complete flow from Phase 1 to Phase 8",
        "The diagram shows the full runtime flow. The detailed cards that follow explain exactly what happens in every phase."
    )
    add_figure(
        doc,
        workflow_png,
        "Figure 3. Eight-phase workflow with the configurable operating-mode branch.",
        width=6.25,
    )

    add_page_break(doc)
    add_phase_card(
        doc, 1, "Receive, authenticate, and confirm the false-positive trigger",
        "Create one reliable, idempotent analysis request only when the authoritative final case outcome maps to FALSE_POSITIVE.",
        "Existing Case Service plus required transactional Case Outbox, and the Rule Optimization Service's Trigger Gateway, Trigger Inbox, and Case Outcome Resolver.",
        "Institution ID, case ID, decision/action identity or stable hash, final decision code, finalization time, event ID, and correlation ID.",
        "The Case Service sends identifiers through an event or API. The Outcome Resolver re-reads the authoritative case through a Case Service API or approved read-only query.",
        "Normal backend logic: authenticate the service, authorize the institution, validate the event schema, deduplicate, load the effective outcome mapping, and verify the case is still final. Final case decision overrides alert-level decisions.",
        "Initial deployment: one modular Spring Boot Rule Optimization Service with Spring Security, PostgreSQL job/inbox/outbox state, a secured REST adapter, and a reconciliation poller. "
        "The case-side transactional outbox is required. A Kafka adapter is added only when approved EFRM event infrastructure exists.",
        "No. An LLM must not interpret institution decision codes.",
        "No.",
        "Accepted canonical false-positive trigger plus an immutable policy snapshot; or an explicit ignored/rejected reason.",
        "The Orchestrator creates a case-analysis job and schedules lineage resolution.",
    )
    doc.add_heading("Event, API, outbox, and polling in simple English", level=2)
    add_table(
        doc,
        ["Method", "What happens", "Recommended role"],
        [
            ("Event", "Case Service records that a case was finalized and continues. A consumer processes it asynchronously.", "Primary target because case closure does not depend on the optimizer being online."),
            ("API call", "Case Service directly calls the optimizer and receives an HTTP result.", "Supported integration when event infrastructure is unavailable."),
            ("Transactional outbox", "The case update and a small event row are committed in the same database transaction.", "Prevents a closed case from losing its trigger event."),
            ("Polling", "A worker checks new finalized decisions from a saved high-water mark.", "Recovery/reconciliation, not repeated full-table scanning."),
        ],
        widths=[1800, 4300, 3260],
        font_size=8.9,
    )
    add_callout(
        doc,
        "External engineering review - current-state verification required",
        "A colleague review reports that case closure currently persists the decision and emits CASE_CLOSED, but did not identify a transactional "
        "case-finalized outbox or a case decision-version field. The owning team must verify this against the official repository baseline. "
        "If confirmed, the closure transaction must also write a versioned case-finalized outbox event. Deduplication should use the authoritative "
        "decision action/audit identity; a governed stable hash is only the fallback.",
        fill=PALE_RED,
        accent=RED,
    )
    add_callout(
        doc,
        "Running example",
        "Case FP-1042 is finalized. The institution mapping says its decision code means FALSE_POSITIVE. "
        "The Trigger Inbox has not seen this decision identity before, so a new case-analysis job is created.",
        fill=PALE_ORANGE,
        accent=ORANGE,
    )

    add_phase_card(
        doc, 2, "Trace the case to alerts and contributing rules",
        "Find the exact rule matches connected to the false-positive case and avoid blaming unrelated rules.",
        "Case Lineage Resolver followed by the Rule Attribution Engine.",
        "Accepted case-analysis job, institution ID, case ID, and case-decision identity.",
        "Case and alert read APIs or approved parameterized reads of case_master, polymorphic case_alert_mapping, transaction/device alert, result, match, master, and request records.",
        "A canonical false-positive case makes the case eligible for investigation; it does not label every alert, match, or rule as false positive. "
        "The resolver processes each mapped alert separately and routes case_alert_mapping by alert_id + alert_type + alert_source_table. "
        "An alert and its matches are sibling children of the same transaction_result or device_result. "
        "Institution scope is inherited from case and transaction/device master/request records and cross-checked at every hop. "
        "The final case decision controls case eligibility, while alert decisions and reproducible contribution determine which rules may be analyzed. "
        "Attribution uses reproducible rule-group order, signals, weights, decision policy, upgrades, and alert logic; otherwise the result is UNRESOLVED. Missing or deleted rule references are recorded, never guessed.",
        "Proposed: Spring Boot worker, JDBC/jOOQ-style parameterized repositories, read-only credentials, immutable lineage JSON with evidence IDs and checksums.",
        "No. Relational identity and rule contribution must be exact.",
        "No.",
        "CaseRuleLineage plus preliminary attributed rule targets and data-quality flags.",
        "Only primary/supporting targets move to configuration resolution. Unresolved attribution can produce manual review but not an automatic candidate.",
    )
    doc.add_heading("Transaction and device paths", level=2)
    add_code_block(
        doc,
        "TRANSACTION PATH\n"
        "case_master -> case_alert_mapping -(discriminator)-> transaction_alert\n"
        "transaction_alert -> transaction_result\n"
        "transaction_result -> transaction_match\n"
        "transaction_result -> transaction_master -> transaction_request\n\n"
        "DEVICE PATH\n"
        "case_master -> case_alert_mapping -(discriminator)-> device_alert\n"
        "device_alert -> device_result\n"
        "device_result -> device_match\n"
        "device_result -> device_master -> device_request"
    )
    add_callout(
        doc,
        "Important attribution rule",
        "case_alert_mapping.is_primary_flag identifies a primary alert, not a primary rule. "
        "A rule is called a contributor only when the Agent can reproduce how its match affected the alert decision in the effective group and policy context.",
        fill=PALE_RED,
        accent=RED,
    )
    add_callout(
        doc,
        "Running example",
        "The resolver follows FP-1042 to alert A-781 and its stored matches. R-204 is a primary contributor. "
        "Another rule appeared in the same case but did not influence this alert, so it is marked coincidental and is not analyzed.",
        fill=PALE_ORANGE,
        accent=ORANGE,
    )

    add_phase_card(
        doc, 3, "Resolve exact historical and current rule configuration",
        "Reconstruct what actually ran, then check whether the same problem still exists in the current active rule.",
        "Effective Configuration Resolver and Rule Identity Resolver.",
        "Attributed rule code/reference, event timestamps, institution/source/channel scope, policy snapshot, and configuration evidence from Phase 2.",
        "Preferred source: an immutable rule_execution_manifest created when the original evaluation ran and linked from the result/evaluation record. "
        "For legacy events, use an authoritative historical Admin/Rule API, audit history, archived configuration bundle, or immutable pre-change snapshot. "
        "Fallback read views can use rule_master, rule_version, rule_drl_context, group/version maps, source bindings, decision policy/upgrades, required data, metric dependencies, and source/engine attribute maps only when time travel is provable.",
        "Resolve database record IDs and business version numbers. Build a canonical point-in-time bundle containing structured JSON logic, DRL and checksums, "
        "rule_drl_context, ordered group/version membership, source binding, immutable decision-policy and metric-definition versions, required data, "
        "rule_metric_dependency, source/engine attribute definitions and mappings, reference-data versions, compiler/runtime identity, event/fact schema version, "
        "and effective institution/outcome policy. Compare the historical bundle with the current bundle. "
        "If complete validity history is unavailable, return CONFIGURATION_AMBIGUITY.",
        "Initial deployment: configuration and identity modules inside the modular Rule Optimization Service, behind typed ports such as RuleConfigurationPort and "
        "RuleCompilerPort; approved read-only API/replica access; canonical JSON serialization and SHA-256 hashes.",
        "No. Configuration resolution is a deterministic identity and versioning problem.",
        "No.",
        "ExecutionManifestReference, immutable HistoricalRuleBundle, CurrentRuleBundle, current-version relevance result, policy/version IDs, and checksums.",
        "If the referenced rule was deleted and no immutable history exists, or if the issue no longer applies to the current rule, the workflow stops or publishes a limited review result. Otherwise it starts historical snapshot building.",
    )
    doc.add_heading("Who provides the configuration?", level=2)
    add_para(
        doc,
        "The Agent does not ask the LLM to reconstruct a rule. A dedicated Configuration Resolver owns this task. "
        "It calls a governed Admin/Rule read service when one is available. If no sufficient API exists, an approved "
        "read-only adapter uses parameterized SQL against a replica. The existing compiler exposure is not yet known, "
        "so the design hides it behind RuleCompilerPort until the owning team confirms whether it is a library, service, or other controlled interface. "
        "Because the supplied schema does not contain complete effective-date history for every configuration object, immutable pre-change bundles or an authoritative history service are a production requirement."
    )
    add_callout(
        doc,
        "P0 immutable execution-manifest prerequisite",
        "The source platform should create one immutable manifest per rule evaluation/result rather than copying a full bundle into every match row. "
        "At minimum it pins manifest ID/checksum, institution, event/evaluation identity and time, ordered rule-group membership, rule IDs/versions/checksums, "
        "logic and DRL hashes, metric-definition and decision-policy version IDs/checksums, source/attribute/reference-data versions, compiler/Drools/JDK/application "
        "versions, event/fact schema version, and the effective outcome-policy version. Legacy rows without this complete context may support limited descriptive "
        "analysis, but not trusted historical parity or candidate-impact claims.",
        fill=PALE_RED,
        accent=RED,
    )
    add_callout(
        doc,
        "Running example",
        "The historical match points to R-204 version 5. The resolver loads version 5 and the current version 6. "
        "Version 6 still uses the seven-day device window, so the issue is relevant. Both bundles receive immutable hashes.",
        fill=PALE_ORANGE,
        accent=ORANGE,
    )

    add_page_break(doc)
    add_phase_card(
        doc, 4, "Build a frozen historical evidence snapshot",
        "Create a trustworthy evaluation population before calculating any performance finding.",
        "Historical Snapshot Builder, Metric SQL Executor, and Data Quality Gate.",
        "Selected current rule bundle, configured historical windows, institution/source/channel scope, required fields and metrics, and effective outcome mapping.",
        "Approved bulk APIs, PostgreSQL read replica, or controlled analytical store. The builder reads eligible transaction/device events, results, matches, alerts, cases, and point-in-time metric inputs.",
        "Engineers provide reviewed query templates. The builder retrieves events where the rule fired and did not fire, attaches mature outcomes, reconstructs event-time Custom-SQL metrics, freezes the cohort, and writes a manifest. It checks lineage, tenant scope, label maturity, missingness, version consistency, and whether the eligible population can be reconstructed.",
        "Proposed: asynchronous snapshot worker; parameterized SQL; read-only PostgreSQL; columnar Parquet snapshot in institution-isolated object storage; metadata/checksums in an optimizer PostgreSQL schema. Custom SQL must use an approved parameter contract and restricted execution boundary.",
        "No. The LLM neither writes nor executes SQL.",
        "No model is required. Data-quality rules and point-in-time reconstruction are deterministic.",
        "Immutable HistoricalSnapshot, manifest, row/exclusion counts, configuration hashes, and readiness status.",
        "READY moves to analytics. Failures such as ELIGIBLE_POPULATION_NOT_RECONSTRUCTABLE, METRIC_NOT_REPRODUCIBLE, or INSUFFICIENT_EVIDENCE stop unsafe optimization and become explicit limitations.",
    )
    doc.add_heading("Why the snapshot includes more than false-positive rows", level=2)
    add_para(
        doc,
        "If the Agent looked only at false-positive cases, every row would already show a problem and the analysis would be biased. "
        "It needs the full eligible population to compare fired versus non-fired events, false-positive versus mature non-false-positive "
        "outcomes, overlap with other rules, and the effect of a proposed change."
    )
    add_callout(
        doc,
        "Running example",
        "The snapshot contains every eligible mobile-transfer event in the configured window, not only FP-1042. "
        "It includes events that did not match R-204 and mature outcomes from other cases. The data-quality gate confirms which conclusions are safe.",
        fill=PALE_ORANGE,
        accent=ORANGE,
    )

    add_page_break(doc)
    add_phase_card(
        doc, 5, "Calculate rule findings with deterministic analytics",
        "Measure what is happening, how stable it is, and where the false-positive pattern is concentrated.",
        "Rule Analytics Engine with health, threshold, decay/stability, overlap, and internal-leakage modules.",
        "Frozen historical snapshot, rule bundles, institution analysis policy, outcome mapping, and alert-volume constraint.",
        "Only the immutable snapshot and its manifest are used. No worker silently re-queries changing live data during a run.",
        "SQL performs scoped extraction and aggregation. Deterministic Python calculates counts, rates, threshold-distance distributions, time buckets, cohort comparisons, confidence intervals, EWMA/CUSUM or change-point checks where data supports them, overlap/unique coverage, and sensitivity to unlabeled outcomes.",
        "Proposed: Python analytics worker with Polars or pandas, NumPy, and SciPy; optional scikit-learn for approved clustering/change-point methods. Functions are versioned, unit-tested, seeded where randomness is used, and return structured JSON with evidence IDs.",
        "No LLM is needed for the authoritative calculations.",
        "Optional statistical/ML methods may discover cohorts or change points, but they do not approve a change. The initial release can operate without a trained ML model.",
        "RuleFindingSet: verified metrics, pattern cohorts, decay result, threshold sensitivity, overlap, internal leakage signals, uncertainty, and limitations.",
        "The Orchestrator checks evidence strength. Insufficient evidence becomes a monitor/review recommendation. Sufficient evidence moves to bounded reasoning and candidate strategy.",
    )
    doc.add_heading("Examples of measurements", level=2)
    add_table(
        doc,
        ["Area", "Example measurements", "What it tells us"],
        [
            ("False-positive burden", "False-positive count/rate among mature final outcomes; repeat-entity burden; alert-volume change", "Whether selected rule alerts create avoidable work and how they affect the current alert limit."),
            ("Threshold behavior", "Distance from threshold; outcome rate in threshold bands; sensitivity curve", "Whether small threshold changes may separate useful and unnecessary matches."),
            ("Stability and decay", "Fire-rate drift; outcome-yield drift; missingness; persistent time-bucket change", "Whether performance is getting worse rather than showing normal variation."),
            ("Overlap", "Matches shared with other rules; unique coverage", "Whether the rule adds distinct value or mostly duplicates another control."),
            ("Internal leakage", "Mature internally positive outcomes not captured by the selected rule; candidate coverage loss", "Where available EFRM evidence may be missed. It is not proof of unknown fraud."),
        ],
        widths=[1900, 3900, 3560],
        font_size=8.3,
    )
    add_callout(
        doc,
        "Running example",
        "The analytics engine finds that false-positive outcomes are concentrated among established customers using "
        "trusted beneficiaries during device days 4-7. It also measures how many mature non-false-positive outcomes occur in that band. "
        "These are verified findings, not an LLM opinion.",
        fill=PALE_ORANGE,
        accent=ORANGE,
    )

    add_phase_card(
        doc, 6, "Form bounded hypotheses and change strategies",
        "Turn verified patterns into safe, testable ideas without allowing free-form changes to the rule language.",
        "Candidate Strategy Service, deterministic mutation/search engine, and optional LLM Reasoning Worker behind the Model Gateway.",
        "RuleFindingSet, structured current rule JSON, allowed mutation types, threshold ranges, alert-volume constraint, institution policy, and optional approved documentation context.",
        "Verified evidence comes from Phase 5. Optional RAG can retrieve stable rule descriptions, glossaries, SOPs, and configuration documentation. It does not retrieve live transaction rows or authoritative current configuration.",
        "The deterministic generator enumerates allowed threshold/window variants and rule-template changes. If enabled, the LLM receives only masked aggregates, candidate schema, evidence IDs, and allowed actions. It may explain patterns or suggest a bounded scenario hypothesis. Every idea is converted into a typed proposal and validated against policy.",
        "Proposed: Java strategy service plus Python search functions; provider-neutral Model Gateway using a local-model adapter or approved external-API adapter; JSON Schema/structured output; evidence-grounded prompt; token/data limits; deterministic fallback templates.",
        "Optional, and this is the first phase where it can add real value. It helps connect evidence into understandable hypotheses and drafts scenario-level ideas. It cannot calculate authoritative metrics, access raw data, write SQL/DRL, or decide approval.",
        "No additional trained model is required. Candidate search is deterministic or optimization-based.",
        "CandidateStrategySet. In Candidate + Test mode this contains typed candidate values to test. In Recommendation Only mode it contains structured human change proposals and explanations.",
        "The Orchestrator applies the institution operating mode and sends the correct artifacts to Phase 7.",
    )
    doc.add_heading("What happens if the LLM is unavailable or wrong?", level=2)
    add_bullets(doc, [
        "The workflow still has all verified analytics and can produce a deterministic template recommendation.",
        "Structured output validation rejects unknown fields and forbidden operations.",
        "A claim checker rejects numbers or evidence references that are not present in the verified input.",
        "Confidence comes from deterministic evidence strength, not from an LLM saying it is confident.",
        "The job records the model, prompt version, provider request ID, and whether fallback was used.",
    ])
    add_callout(
        doc,
        "Running example",
        "The bounded strategy service proposes testing a three-day new-device window and, if the formal rule model supports it, "
        "a trusted-beneficiary condition. The LLM may explain why these ideas match the pattern, but it does not decide that either idea is safe.",
        fill=PALE_ORANGE,
        accent=ORANGE,
    )

    add_phase_card(
        doc, 7, "Execute the institution's configured operating mode",
        "Either create and test an executable candidate or produce analysis-only advice, according to client access and policy.",
        "Mode Router. Candidate + Test uses Candidate Validator, RuleCompilerPort, Isolated Drools Replay Runner, and Result Evaluator. Recommendation Only uses the Structured Recommendation Builder.",
        "CandidateStrategySet, frozen snapshot, current rule/group context, compiler/replay availability, and configured operating mode.",
        "All inputs are immutable artifacts created in earlier phases. Candidate + Test also calls the existing JSON-to-DRL compiler through its future confirmed interface.",
        "Candidate + Test: mutate structured rule JSON in the optimizer store; validate schema and dependencies; compile; run historical parity replay; run current baseline simulation; replace only the selected rule in the same group context; run candidate simulation; compare outcomes. Recommendation Only: skip executable candidate, compiler, and candidate replay; convert findings into a clear proposed human change with limitations.",
        "Proposed: Java candidate validator; RuleCompilerPort adapter; isolated Java/Drools replay worker compatible with production semantics; versioned containers/build artifacts; Python result comparison. Recommendation Only is normal backend logic plus optional evidence-grounded wording.",
        "Optional for explanation only. The LLM does not compile or execute the rule.",
        "Drools is the formal rule execution engine, not an AI model. No trained ML model is required.",
        "Candidate + Test outputs executable candidate artifact plus simulation evidence. Recommendation Only outputs analysis evidence plus a proposed human change and the statement 'no executable candidate backtest performed'.",
        "Validated result artifacts move to recommendation composition. Replay mismatch or unsupported full-context simulation produces a limited result rather than overstated impact.",
    )
    doc.add_heading("The three different replay questions", level=2)
    add_table(
        doc,
        ["Replay", "Question", "Why it matters"],
        [
            ("Historical parity replay", "Can the isolated runner reproduce the stored historical result using the rule/configuration effective at each event?", "Proves the simulator is trustworthy."),
            ("Current baseline simulation", "How does today's current rule behave on the frozen evaluation cohort?", "Creates the fair baseline for a current recommendation."),
            ("Candidate simulation", "How does the candidate behave on the exact same cohort and full rule-group context?", "Measures the candidate delta."),
        ],
        widths=[2300, 3900, 3160],
        font_size=8.7,
    )
    add_callout(
        doc,
        "Claim boundary",
        "If the runner can replay only the selected rule, the output may claim match-level changes only. "
        "It may claim alerts or cases prevented only when the full group, decision policy, and downstream alert/case logic are reproduced.",
        fill=PALE_RED,
        accent=RED,
    )

    doc.add_heading("Illustrative simulation for R-204", level=2)
    add_table(
        doc,
        ["Same frozen cohort", "Current baseline", "Example candidate", "Illustrative change"],
        [
            ("Total alerts", "100", "72", "-28"),
            ("False-positive final outcomes retained", "60", "35", "-25"),
            ("Mature non-false-positive outcomes retained", "20", "19", "-1"),
            ("Unresolved/unlabeled alerts retained", "20", "18", "-2"),
        ],
        widths=[3000, 2000, 2100, 2260],
        font_size=8.8,
    )
    add_para(
        doc,
        "These numbers are fictional teaching values. They show the comparison format only. They are not a result from the sample database or a production performance claim.",
        size=8.7,
        color=RED,
        italic=True,
    )

    add_page_break(doc)
    add_phase_card(
        doc, 8, "Publish one evidence-linked recommendation and wait for human governance",
        "Give reviewers a concise answer: which rule is the issue, what change is proposed, why, what evidence supports it, and what is still uncertain.",
        "Recommendation Composer, Claim Validator, Recommendation API/Event Publisher, future UI adapter, and existing Human Governance process.",
        "Verified findings, mode-specific artifacts, simulation results if available, policy/version provenance, evidence references, and mandatory limitations.",
        "The optimizer Evidence Store supplies immutable artifacts. Optional LLM wording is accepted only after structured validation and claim checking.",
        "Build canonical recommendation JSON first. Check every number against evidence; attach IDs and checksums; calculate evidence-strength status; mark stale/superseded recommendations; create a plain-English view. Publish a RECOMMENDATION_READY event or expose a read API. Humans review and make any production change through the existing process.",
        "Proposed: Spring Boot recommendation service and OpenAPI; PostgreSQL evidence metadata; object storage for large artifacts; optional LLM Model Gateway; deterministic template fallback; existing governance adapter to be confirmed.",
        "Optional. It can produce the readable 'why', summary, and limitations from approved structured evidence. It cannot add unsupported facts or decide whether production changes.",
        "No.",
        "One immutable recommendation object with status, rule/version, issue, proposed change, evidence, simulation or no-test statement, limitations, provenance, and HUMAN_REVIEW_REQUIRED.",
        "The workflow ends in REVIEW_REQUIRED. A human accepts, rejects, asks for more evidence, or creates a governed draft through existing EFRM processes. The Agent never activates it.",
    )
    add_callout(
        doc,
        "Running example - Candidate + Test wording",
        "Review R-204. The seven-day new-device condition is associated with a concentrated false-positive pattern among established customers using trusted beneficiaries. "
        "The tested three-day candidate reduced projected alerts in the frozen cohort while retaining the stated mature non-false-positive outcomes. "
        "Human review is required because hard optimization constraints are not yet fully defined.",
        fill=PALE_GREEN,
        accent=GREEN,
    )
    add_callout(
        doc,
        "Running example - Recommendation Only wording",
        "Review reducing the new-device window or adding an approved trusted-beneficiary condition. "
        "The evidence shows concentration during device days 4-7. No executable candidate was created or backtested, so projected alert reduction is not claimed.",
        fill=PALE_ORANGE,
        accent=ORANGE,
    )

    # Modes
    add_section_intro(
        doc, "7", "The two configurable operating modes",
        "The mode is resolved from the effective institution policy. It is not chosen casually by the LLM or by a developer at runtime."
    )
    add_table(
        doc,
        ["Decision point", "Mode A - Candidate + Test", "Mode B - Recommendation Only"],
        [
            ("When to use", "Client permits candidate compilation and isolated testing.", "Client does not permit executable candidate creation or compiler access."),
            ("Candidate artifact", "Yes, stored only in the optimizer boundary.", "No executable candidate."),
            ("Existing compiler", "Called through RuleCompilerPort.", "Not called."),
            ("Drools backtest", "Yes, after replay parity is proven.", "Normally no candidate backtest."),
            ("Output", "Recommendation plus candidate diff and simulation evidence.", "Recommendation plus analysis evidence and explicit no-backtest limitation."),
            ("Production write", "Never.", "Never."),
            ("Human approval", "Required.", "Required."),
        ],
        widths=[2200, 3580, 3580],
        font_size=8.5,
    )
    add_para(
        doc,
        "Both modes may still recommend 'no change', 'monitor', 'review the rule', or 'review for retirement'. "
        "A retirement recommendation is advice only; the Agent cannot retire or delete the rule."
    )

    # LLM section
    add_section_intro(
        doc, "8", "Exactly where the LLM is used, and why",
        "An LLM can make the optimizer easier to understand and better at proposing explainable hypotheses. "
        "It should not be used for work that requires exact records or repeatable calculations."
    )
    add_table(
        doc,
        ["Phase", "LLM use", "Reason"],
        [
            ("1-4", "No", "Security, decision mapping, lineage, configuration, SQL retrieval, and data quality must be exact."),
            ("5", "No for metrics; optional after metrics", "Authoritative calculations are deterministic. The LLM may receive only the verified result."),
            ("6", "Optional and useful", "Connect verified patterns into bounded hypotheses and explain why a change might help."),
            ("7", "Explanation only", "Formal JSON validation, compilation, and Drools replay are deterministic."),
            ("8", "Optional and useful", "Turn structured evidence into simple English and summarize limitations."),
        ],
        widths=[900, 2600, 5860],
        font_size=8.9,
    )
    doc.add_heading("Recommended LLM contract", level=2)
    add_bullets(doc, [
        "Input only approved aggregates, masked cohort descriptions, candidate diffs, limitations, and evidence IDs.",
        "Require structured JSON output with allowed fields and allowed recommendation types.",
        "Block raw customer, account, card, device, investigator note, secret, and unrestricted SQL data.",
        "Validate every numeric claim and evidence reference after generation.",
        "Use a local hosted model or an approved external API through the same Model Gateway.",
        "Fall back to deterministic templates when the model is unavailable or disallowed.",
    ])
    add_callout(
        doc,
        "Would using an LLM make the rule optimization better?",
        "It can improve hypothesis wording, scenario explanation, and reviewer usability. It does not make counts, joins, thresholds, or backtests more correct. "
        "The best design combines deterministic evidence with bounded LLM reasoning.",
        fill=PALE_TEAL,
        accent=TEAL,
    )

    # RAG
    add_section_intro(
        doc, "9", "Is RAG useful?",
        "RAG means the system searches an indexed collection of documents and gives only the relevant passages to the LLM. "
        "It is useful for slow-changing knowledge, but it is the wrong source for exact live rule or transaction data."
    )
    add_table(
        doc,
        ["Information", "Use RAG?", "Correct source"],
        [
            ("Rule workflow documentation", "Yes, optional", "Versioned document knowledge base."),
            ("Business glossary and decision meaning documentation", "Yes, optional", "Versioned, effective-dated approved documents."),
            ("Rule-language and compiler documentation", "Yes, optional", "Approved technical documents."),
            ("Investigation SOP and governance policy", "Yes, optional", "Approved policy documents."),
            ("Current exact rule JSON/DRL", "No", "Admin/Rule read service or read-only configuration query."),
            ("Case, alert, match, and transaction rows", "No", "Approved operational read APIs or read replica."),
            ("Authoritative metrics and simulation results", "No", "Deterministic analytics and replay artifacts."),
        ],
        widths=[3300, 1600, 4460],
        font_size=8.6,
    )
    add_para(
        doc,
        "Recommendation: do not make RAG a dependency for the first working release. Build the event, lineage, configuration, snapshot, analytics, and recommendation flow first. "
        "Add RAG later if reviewers need richer explanations of rule intent, policies, or terms. Every retrieved passage should carry document version and source metadata."
    )

    # Data access
    add_section_intro(
        doc, "10", "How data and configuration retrieval actually work",
        "The Agent's LLM does not freely browse tables. Purpose-built services and reviewed query components retrieve the exact data required by each phase."
    )
    add_code_block(
        doc,
        "ORCHESTRATOR REQUESTS A TOOL\n"
        "    -> tool authenticates institution scope\n"
        "    -> approved service API or parameterized SQL reads exact data\n"
        "    -> result is validated and given an evidence ID\n"
        "    -> immutable artifact and checksum are stored\n"
        "    -> orchestrator passes only the artifact reference to the next phase"
    )
    add_table(
        doc,
        ["Need", "Preferred provider", "Fallback", "LLM involvement"],
        [
            ("Final case outcome", "Case Service read API", "Approved read-only case view/query", "None"),
            ("Case-alert-rule lineage", "Case/Rule read APIs", "Reviewed parameterized joins on read replica", "None"),
            ("Exact rule configuration", "Admin/Rule configuration API", "Approved read-only configuration adapter", "None"),
            ("Bulk historical population", "Analytical/read service", "Read replica or controlled historical store", "None"),
            ("Metric Definition Custom SQL", "Governed Metric SQL Executor", "No unsafe fallback", "None"),
            ("Stable explanatory documents", "Optional RAG retriever", "Direct approved document lookup", "LLM may consume retrieved text"),
        ],
        widths=[2200, 2800, 2600, 1760],
        font_size=8.3,
    )
    add_callout(
        doc,
        "Custom SQL safety",
        "The current product has Metric Definition Custom SQL but no dedicated optimizer controls. "
        "The design does not use aggregated_metric; derived and aggregate values come from governed Metric Definition Custom SQL. "
        "The production optimizer must require approved placeholders/parameters, read-only schemas, tenant enforcement, statement timeouts, row limits, and reproducibility checks. "
        "It must not try to inject filters into arbitrary SQL text.",
        fill=PALE_RED,
        accent=RED,
    )

    # Recommendation
    add_page_break(doc)
    add_section_intro(
        doc, "11", "What the Agent produces",
        "The main user-facing output is a concise recommendation, not an unstructured report. Full evidence remains stored for audit and drill-down."
    )
    add_code_block(
        doc,
        '{\n'
        '  "recommendation_id": "REC-...",\n'
        '  "operating_mode": "CANDIDATE_AND_TEST | RECOMMENDATION_ONLY",\n'
        '  "institution_id": "scope",\n'
        '  "rule": {"rule_code": "R-204", "historical_version": 5, "current_version": 6},\n'
        '  "status": "CANDIDATE_FOR_REVIEW",\n'
        '  "issue": "Evidence-linked description of the problem",\n'
        '  "recommended_change": {"type": "...", "current": "...", "proposed": "..."},\n'
        '  "why": ["Plain-English explanation tied to evidence IDs"],\n'
        '  "evidence": {"analysis_run_id": "...", "snapshot_id": "...", "backtest_id": "..."},\n'
        '  "limitations": ["Missing constraints, unlabeled data, replay scope, or other limits"],\n'
        '  "action": "HUMAN_REVIEW_REQUIRED"\n'
        '}'
    )
    doc.add_heading("Minimum provenance stored with every recommendation", level=2)
    add_bullets(doc, [
        "Trigger and analysis job IDs; triggering case-decision identities.",
        "Institution, source/channel/fact scope, and operating/reasoning mode.",
        "Source rule-execution manifest ID/hash and verification status; legacy-limit reason when unavailable.",
        "Rule master/version IDs, business version numbers, historical/current checksums, group and binding context.",
        "Policy, outcome-mapping, and metric-definition versions.",
        "Snapshot ID/hash, extraction/query hashes, code versions, and evidence IDs.",
        "Attribution result, data-quality gates, replay parity, compiler/engine versions, and simulation scope.",
        "Limitations, stale/superseded state, creation time, and audit correlation ID.",
    ])
    add_para(
        doc,
        "Until business owners define the full optimization constraints, the safe status for a tested idea is "
        "CANDIDATE_FOR_REVIEW, not a claim that the change is policy compliant."
    )

    # Risk leakage and limitations
    add_section_intro(
        doc, "12", "Risk leakage in basic English",
        "Risk leakage asks: if the current or candidate rule does not match, what important internally observed outcomes might no longer be covered?"
    )
    add_para(
        doc,
        "Because no independent fraud truth source is required for the current scope, the system must use careful wording. "
        "It can measure missed or delayed coverage against mature EFRM investigation outcomes. It cannot estimate fraud that no EFRM control ever observed."
    )
    add_table(
        doc,
        ["Safe claim", "Unsafe claim"],
        [
            ("The candidate did not retain 1 of 20 mature non-false-positive EFRM outcomes in the frozen cohort.", "The candidate missed one real fraud."),
            ("The selected rule has little unique coverage because the same mature outcomes were captured by other rules.", "This rule has no risk value."),
        ],
        widths=[4680, 4680],
        font_size=8.8,
    )

    # Safety
    add_section_intro(
        doc, "13", "Enterprise safety, consistency, and failure behavior",
        "The workflow is designed to fail clearly instead of filling gaps with guesses."
    )
    add_table(
        doc,
        ["Control", "Technical behavior"],
        [
            ("Institution isolation", "Every event, query, cache key, snapshot partition, artifact, and authorization decision includes institution scope."),
            ("Idempotency", "Duplicate events or API retries do not create duplicate case-analysis jobs."),
            ("Immutable evidence", "Each phase stores a content hash and never silently changes the inputs of an active run."),
            ("Case correction", "A reopened or corrected case immediately cancels or supersedes active jobs and recommendations tied to the old decision."),
            ("Read-only EFRM access", "Optimizer credentials cannot update production case, alert, match, or rule tables."),
            ("Data sufficiency gates", "Missing lineage, configuration ambiguity, unreproducible metrics, or weak sample evidence produces an explicit stop status."),
            ("Replay parity gate", "Candidate results are not trusted until the isolated runner can reproduce approved historical behavior."),
            ("LLM guardrails", "Masked input, allow-listed tools, schema output, claim validation, timeouts, cost limits, and deterministic fallback."),
            ("Human control", "No endpoint or worker can activate, retire, delete, or directly modify a production rule."),
        ],
        widths=[2300, 7060],
        font_size=8.6,
    )
    doc.add_heading("Typical stop states", level=2)
    add_table(
        doc,
        ["State", "Meaning"],
        [
            ("IGNORED_NOT_FALSE_POSITIVE", "The final effective case outcome does not qualify."),
            ("LINEAGE_FAILED", "The case cannot be reliably traced to contributing rules."),
            ("CONFIGURATION_AMBIGUITY", "The exact rule/group/binding cannot be resolved."),
            ("METRIC_NOT_REPRODUCIBLE", "A required point-in-time Custom-SQL metric cannot be safely rebuilt."),
            ("ELIGIBLE_POPULATION_NOT_RECONSTRUCTABLE", "The full rule-eligible historical population cannot be identified."),
            ("INSUFFICIENT_EVIDENCE", "There is not enough mature, reliable data for a change recommendation."),
            ("REPLAY_MISMATCH", "The isolated runner cannot reproduce the approved baseline."),
            ("REVIEW_REQUIRED", "A complete recommendation is ready for a human."),
        ],
        widths=[3600, 5760],
        font_size=8.7,
    )

    # Framework
    add_section_intro(
        doc, "14", "Application framework: confirmed facts and proposed implementation",
        "The Rule Optimization Agent has not yet been built. The original architecture is based on the supplied PRD, database material, and user-confirmed product behavior. "
        "A later colleague review reports additional observations from Rule Engine, Admin Service, replay-controller, and Agentic portal code, but those repositories and versions "
        "are not available in this workspace. They are therefore recorded as reviewer-reported and pending owner verification, not as independently confirmed facts."
    )
    doc.add_heading("Confirmed from the supplied EFRM material", level=2)
    add_table(
        doc,
        ["Confirmed item", "What is not yet confirmed"],
        [
            ("PostgreSQL efrm schema and extensive JSONB usage", "Production PostgreSQL version and final optimizer storage topology."),
            ("Spring-based controllers/services in the existing Rule Engine family", "Exact Java, Spring, and Spring Boot versions."),
            ("Drools/DRL rule execution", "Exact Drools version and deployment interface."),
            ("Structured rule logic in rule_version.logic and generated DRL", "Exact compiler exposure: library, REST, gRPC, or build service."),
            ("Metric Definition Custom SQL", "Production governance and historical execution controls."),
            ("Transaction/device result, match, alert, and case lineage", "Final read APIs available to the optimizer."),
            ("Institution ID scoping and configurable decision meanings", "First pilot institution or channel; none is selected."),
        ],
        widths=[4700, 4660],
        font_size=8.4,
    )
    add_callout(
        doc,
        "Reviewer-reported implementation evidence - pending owner verification",
        "The external review reports: rule logic/checksum versioning exists; match rows do not retain a complete execution bundle; Metric SQL and decision-policy history are "
        "not immutable enough for parity; case closure emits CASE_CLOSED without a durable finalization outbox; bulk/replay endpoints and Agentic portal routes already exist. "
        "These observations refine the target interfaces, but the owning teams must verify the exact repository revision before low-level design approval.",
        fill=PALE_ORANGE,
        accent=ORANGE,
    )
    doc.add_heading("Concrete reference implementation proposed for approval", level=2)
    add_table(
        doc,
        ["Layer", "Proposed technology", "Why"],
        [
            ("Initial deployment unit", "One modular Java/Spring Boot Rule Optimization Service", "Packages control, policy, lineage, orchestration, evidence, and recommendation modules without premature microservices."),
            ("Durable orchestration", "PostgreSQL state machine, jobs, leases, inbox, outbox, and work queue", "Provides retries, idempotency, resume, cancellation, and audit with minimal infrastructure. Temporal/Camunda remain later options."),
            ("Trigger transport", "Required case outbox; secured REST and recovery polling; optional Kafka adapter", "Reliability does not depend on Kafka, while approved event infrastructure can be added without changing the trigger contract."),
            ("Data and metadata", "Optimizer PostgreSQL schema plus institution-scoped read-only EFRM analytical/replica access", "Keeps operational writes inside the optimizer and production EFRM access read-only."),
            ("Historical snapshots", "PostgreSQL-managed manifests; Parquet/object storage only when data scale requires it", "Starts simply while preserving an immutable, scale-ready artifact contract."),
            ("Analytics", "Deterministic SQL plus a Python worker/process using Polars/pandas, NumPy, and SciPy when needed", "Keeps exact calculations reproducible without forcing each module into a service."),
            ("Candidate compile/replay", "RuleCompilerPort and RuleReplayPort; isolated Java/Drools worker only when required", "Reuses verified existing capabilities while preserving runtime and security isolation."),
            ("Optional LLM/RAG", "Provider-neutral Model Gateway; documentation index only after deterministic workflow matures", "Supports local or approved external models without making AI infrastructure a first-release dependency."),
        ],
        widths=[1900, 3500, 3960],
        font_size=8.0,
    )
    # Implementation plan
    add_section_intro(
        doc, "15", "How to implement the Agent in controlled stages",
        "The safest path completes source provenance and a deterministic Recommendation Center before replay, candidate generation, or LLM features."
    )
    add_table(
        doc,
        ["Stage", "Build", "Exit evidence"],
        [
            ("0. P0 provenance gate", "Case outbox; canonical outcome versions; rule execution manifest; immutable metric/policy versions; tenant and retention controls.", "Owners prove historical identity and reliable trigger delivery before replay claims."),
            ("1. Modular foundation", "One service; PostgreSQL jobs/inbox/outbox; read-only EFRM access; lineage, attribution, configuration, and frozen snapshot modules.", "Idempotent transaction/device traces and reproducible evidence."),
            ("2. Deterministic Recommendation Only", "Health, threshold, decay, overlap, workload, leakage, claim validation, and structured recommendation.", "First useful capability without LLM, compiler, or production writes."),
            ("3. Recommendation Center + governance", "Secured reviewer APIs, evidence drill-down, stale/superseded handling, four-eyes review, and governed handoff.", "Real human review works end to end; sample UI state is removed."),
            ("4. Point-in-time replay", "RuleReplayPort; frozen payloads/metrics/reference data; temporal holdout; historical parity diagnostics.", "Owner-approved parity under pinned configuration and runtime."),
            ("5. Candidate + Test", "Bounded numeric threshold/window candidates; compiler adapter; baseline/candidate comparison.", "Safe candidate artifacts outside production with correct claim scope."),
            ("6. Broader scenario templates", "Approved condition mutations only after dependencies and full-context replay are proven.", "Risk/product approval per mutation template."),
            ("7. Optional LLM and RAG", "Bounded explanation/hypothesis tools, masking/DLP, claim validation, fallback, approved documentation retrieval.", "Security, red-team, evidence-fidelity, and provider approval."),
            ("8. Controlled rollout", "Shadow operation, UAT, SLO/load tests, observability, recovery runbooks, and phased institution configuration.", "Business, risk, security, architecture, and operations sign-off."),
        ],
        widths=[1900, 4750, 2710],
        font_size=7.9,
    )
    doc.add_heading("Recommended first production value", level=2)
    add_para(
        doc,
        "After the P0 provenance gate, deliver deterministic Recommendation Only and the real Recommendation Center first. "
        "Candidate + Test can be enabled per institution only after RuleReplayPort parity, compiler exposure, and constraints are approved. "
        "LLM and RAG remain non-gating enhancements after the evidence workflow is mature."
    )
    doc.add_heading("Inputs still needed before low-level production design", level=2)
    add_bullets(doc, [
        "Institution-configurable hard optimization constraints when product/risk owners define them.",
        "Existing infrastructure standards: broker, workflow engine, object store, identity, secrets, deployment, observability, and external-LLM vendor/data/retention approval if external mode will be enabled.",
        "The actual exposure method of the existing JSON-to-DRL compiler.",
        "Available read APIs and the approved read-replica/analytical access pattern.",
        "Governance behavior for preserving immutable rule history when rules are changed or physically deleted.",
    ])

    # Deep technical architecture
    add_part_banner(
        doc,
        "Part II",
        "Detailed Enterprise Architecture and Orchestration",
        "The following sections define component ownership, durable workflow states, data and configuration contracts, "
        "analytics, replay, LLM controls, security, persistence, operations, testing, and delivery gates.",
        accent=BLUE,
        fill=PALE_BLUE,
    )

    add_section_intro(
        doc, "16", "Complete component model and trigger integration",
        "The following are logical components with clear responsibilities, typed inputs and outputs, least-privilege permissions, and an explicit LLM rule. "
        "They do not need to be separate microservices. The first deployment packages control components as modules in one Rule Optimization Service and separates "
        "only resource-heavy analytics or replay workers where operationally necessary."
    )
    add_figure(
        doc,
        trigger_png,
        "Figure 4. Required outbox/inbox controls and supported API/recovery paths converge on one authenticated, idempotent trigger contract; Kafka remains optional.",
        width=6.35,
    )
    doc.add_heading("Control and integration components", level=2)
    add_table(
        doc,
        ["Component", "Responsibility", "Input -> output", "Technique / LLM"],
        [
            ("Existing Case Service", "Own final case state and decision.", "Case update -> authoritative case record.", "Existing backend; no LLM."),
            ("Required Case Outbox + Publisher", "Reliably publish a case-finalized event from the same transaction as the case decision.", "Outbox row -> versioned event.", "Transactional backend; no LLM."),
            ("Event Consumer", "Receive at-least-once messages and acknowledge only after inbox persistence.", "Broker event -> normalized trigger.", "Backend adapter; no LLM."),
            ("Trigger API", "Support clients without event infrastructure.", "Authenticated POST + idempotency key -> trigger identity.", "REST backend; no LLM."),
            ("Reconciliation Poller", "Recover missed triggers using a durable high-water mark and overlap window.", "Cursor -> candidate case decisions.", "Scheduled backend; no LLM."),
            ("Trigger Gateway", "Normalize, authenticate, authorize, validate schema, and assign correlation.", "Any trigger path -> canonical trigger.", "Backend/security logic; no LLM."),
            ("Transactional Inbox", "Persist before acknowledgement and deduplicate retries.", "Canonical trigger -> accepted/duplicate record.", "Database transaction; no LLM."),
            ("Policy + Outcome Resolver", "Resolve effective institution policy and canonical outcome.", "Case decision + timestamp/scope -> policy snapshot + eligibility.", "Deterministic config logic; no LLM."),
            ("Durable Orchestrator", "Advance states, call modules/workers, apply guards, retry, cancel, and record audit.", "Artifact references -> next commands/state.", "Initial PostgreSQL state machine; optional workflow engine later; no LLM authority."),
        ],
        widths=[1850, 3100, 2700, 1710],
        font_size=7.7,
    )
    doc.add_heading("Evidence and analysis components", level=2)
    add_table(
        doc,
        ["Component", "Responsibility", "Main output", "Technique / LLM"],
        [
            ("Case Lineage Resolver", "Route polymorphic alert identity and reconstruct transaction/device lineage.", "CaseRuleLineage artifact.", "Reviewed APIs/SQL; no LLM."),
            ("Attribution Engine", "Reproduce contribution in rule-group and decision context.", "PRIMARY, SUPPORTING, COINCIDENTAL, or UNRESOLVED.", "Deterministic counterfactual logic; no LLM."),
            ("Configuration Resolver", "Load execution-manifest reference plus historical/current effective bundles and hashes.", "ExecutionManifestReference + rule bundles.", "Governed APIs/SQL; no LLM."),
            ("Rule-Run Aggregator", "Group eligible case triggers into a reusable rule-analysis run.", "RuleAnalysisScope.", "Backend policy logic; no LLM."),
            ("Historical Snapshot Builder", "Reconstruct the complete eligible population and freeze it.", "Snapshot manifest + partitions.", "Parameterized SQL/Parquet; no LLM."),
            ("Metric SQL Executor", "Rebuild Metric Definition Custom SQL under point-in-time controls.", "Versioned metric values + lineage.", "Governed SQL runtime; no LLM."),
            ("Data Quality Gate", "Decide which later activities are safe.", "Permission matrix + stop/limit reasons.", "Deterministic validation; no LLM."),
            ("Analytics Engine", "Calculate health, threshold, decay, overlap, volume, and leakage evidence.", "RuleFindingSet.", "SQL + deterministic Python/statistics."),
            ("Optional Model Gateway", "Run bounded hypothesis/explanation prompts with approved tools.", "Structured reasoning artifact.", "Optional LLM; never authoritative facts."),
        ],
        widths=[1900, 3350, 2450, 1660],
        font_size=7.7,
    )
    doc.add_heading("Candidate, simulation, and output components", level=2)
    add_table(
        doc,
        ["Component", "Responsibility", "Main output", "Technique / LLM"],
        [
            ("Candidate Generator", "Create bounded typed changes allowed by policy.", "Candidate intent/JSON stored outside production.", "Deterministic search; optional LLM idea only."),
            ("Candidate Validator", "Check schema, dependencies, allowed mutations, and missing constraints.", "Valid/invalid candidate + reasons.", "Backend validation; no LLM."),
            ("RuleCompilerPort", "Hide whether the existing compiler is a library or service.", "Compiler request/result + artifact hashes.", "Existing formal compiler; no LLM."),
            ("RuleReplayPort + Isolated Runner", "Adapt verified replay APIs or execute isolated parity, current baseline, and candidate simulations.", "Replay/simulation artifacts.", "Java/Drools or verified existing API; no LLM."),
            ("Result Evaluator", "Compare baseline/candidate and enforce claim boundaries.", "Ranked evidence or no-change result.", "Deterministic Python/backend logic."),
            ("Recommendation Composer", "Build canonical recommendation JSON and plain-English view.", "Immutable recommendation.", "Templates; optional LLM wording."),
            ("Claim Validator", "Verify numeric claims, references, allowed language, and limitations.", "Accepted/rejected content.", "Deterministic validation; no arbitrary causal proof."),
            ("Recommendation API/UI", "Expose evidence and review actions.", "Reviewer experience + governance request.", "Normal backend/UI; no LLM decision."),
            ("Governance Adapter", "After human acceptance, optionally create a draft governance work item when the institution permits it.", "Draft request or governance reference/status.", "Integration logic from recorded human action; never activate or directly update a rule."),
        ],
        widths=[1900, 3350, 2450, 1660],
        font_size=7.7,
    )
    doc.add_heading("How the required case outbox and optimizer inbox work", level=2)
    add_numbered(doc, [
        "The Case Service commits the final case decision and a small outbox row in one transaction.",
        "The publisher reads the outbox row and publishes a versioned event through the approved transport. If publishing fails, the row remains for retry.",
        "The consumer receives the event and stores its event identity in the optimizer inbox before acknowledging it.",
        "If the same event arrives again, the inbox returns the existing case-job identity instead of creating a duplicate.",
        "The Outcome Resolver re-reads the authoritative case. A corrected or reopened case cancels or supersedes work tied to the old decision.",
    ])
    add_callout(
        doc,
        "Why this is safer than only making a synchronous API call",
        "Case closure does not depend on the optimizer being online. The outbox prevents a committed case decision from losing its event, "
        "while the inbox makes at-least-once delivery safe. If EFRM has no broker, the same canonical contract can be used through the API and poller.",
        fill=PALE_GREEN,
        accent=GREEN,
    )
    add_callout(
        doc,
        "Current-state evidence to verify before implementation",
        "The external engineering review reports that the current case flow emits CASE_CLOSED but does not yet persist a durable case-finalized outbox or decision version. "
        "Until the Case Management owner confirms the exact code revision, treat this as a reported gap. The target requirement is unchanged: decision and outbox event commit "
        "atomically, and the optimizer inbox deduplicates by an authoritative decision action/audit event identity.",
        fill=PALE_ORANGE,
        accent=ORANGE,
    )

    add_page_break(doc)
    add_section_intro(
        doc, "17", "Durable orchestration, jobs, and the state machine",
        "The orchestrator does not hold a long process in memory. It stores every state and artifact reference so a worker crash, retry, "
        "case correction, or long-running snapshot can be resumed safely."
    )
    add_figure(
        doc,
        state_machine_png,
        "Figure 5. Proposed durable state machine with common phases, two operating-mode branches, human review, and explicit stop states.",
        width=6.25,
    )
    doc.add_heading("Three records that must not be confused", level=2)
    add_table(
        doc,
        ["Record", "Meaning", "Typical identity", "Lifecycle"],
        [
            ("Trigger / inbox record", "One delivered event, API request, or polled decision.", "institution + authoritative event/action ID; otherwise approved stable hash.", "Accepted once; duplicate deliveries point to it."),
            ("Case-analysis job", "Work needed to verify one final case decision and trace its rules.", "trigger_id + case_id + final-decision identity.", "Can be cancelled or superseded if the case changes."),
            ("Rule-analysis run", "A reusable analysis for one current rule/configuration/scope over a configured window.", "institution + rule + current checksum + source/channel/fact + policy/window.", "May aggregate several eligible cases under policy."),
        ],
        widths=[1900, 3000, 2860, 1600],
        font_size=8.0,
    )
    doc.add_heading("Standard stage hand-off contract", level=2)
    add_table(
        doc,
        ["Field", "Purpose"],
        [
            ("command_id / stage_id", "Unique idempotency key for one stage attempt."),
            ("job_id / rule_analysis_run_id", "Links case work and shared rule-level analysis."),
            ("institution_scope", "Mandatory authorization and partition key."),
            ("input_artifact_ids + hashes", "Proves exactly which immutable inputs were used."),
            ("policy/config/code versions", "Makes the stage reproducible and debuggable."),
            ("attempt, lease, heartbeat", "Prevents two workers from owning the same stage and detects abandoned work."),
            ("deadline + retry class", "Separates transient retry, permanent stop, and limited-evidence completion."),
            ("output artifact + checksum", "The only result passed to the next stage."),
            ("audit correlation ID", "Connects event, SQL, analysis, LLM, replay, recommendation, and review logs."),
        ],
        widths=[2900, 6460],
        font_size=8.6,
    )
    doc.add_heading("Transition rules", level=2)
    add_bullets(doc, [
        ("Persist before dispatch. ", "A state change and outgoing command/outbox record are committed together."),
        ("Idempotent workers. ", "Repeating a command with the same stage ID returns the prior result or safely resumes."),
        ("Guard every transition. ", "The next state starts only when required artifacts, hashes, policy, and permission gates exist."),
        ("Retry only transient failures. ", "Network errors, replica lag, and temporary capacity can retry with bounded exponential backoff."),
        ("Stop on evidence defects. ", "Ambiguous lineage, unreproducible metrics, or replay mismatch becomes a named state, not a guessed answer."),
        ("Cancel and supersede. ", "Case corrections and current-rule changes invalidate older work immediately."),
    ])

    add_page_break(doc)
    add_section_intro(
        doc, "18", "Exact case-to-rule lineage and attribution",
        "Lineage answers which stored rule references are connected to the case. Attribution is a separate decision: it asks which rule "
        "actually contributed to the alert or case outcome in the effective rule-group and policy context."
    )
    add_figure(
        doc,
        lineage_png,
        "Figure 6. Correct transaction and device lineage. Alert and match records are siblings under their result; institution scope must be inherited and verified.",
        width=6.35,
    )
    doc.add_heading("Lineage resolution rules", level=2)
    add_bullets(doc, [
        "Start from case_master so institution and final case state are authoritative.",
        "Treat case_alert_mapping as polymorphic; use alert_id, alert_type, and alert_source_table together.",
        "Route to the correct alert table and verify its result relationship.",
        "Read matches through the shared transaction_result_id or device_result_id, not through a guessed alert-to-match join.",
        "Follow the result to transaction/device master and request records to verify institution, source, channel, event time, and required facts.",
        "Preserve every database record ID, business rule/version reference, query hash, and exclusion reason as evidence.",
    ])
    doc.add_heading("Deterministic attribution algorithm", level=2)
    add_numbered(doc, [
        "Reconstruct all matches and signals in the effective rule-group order.",
        "Load rule weights, decision policy, decision upgrades, alert-generation policy, and any group aggregation behavior.",
        "Reproduce the recorded decision with the full context.",
        "Remove or neutralize one matched rule at a time and re-evaluate the decision.",
        "Classify a rule as PRIMARY_CONTRIBUTOR only when its removal changes the qualifying decision in the approved attribution model.",
        "Classify supporting or coincidental matches according to approved policy; if the context cannot be reproduced, return UNRESOLVED.",
    ])
    add_table(
        doc,
        ["Attribution status", "Meaning", "Allowed next action"],
        [
            ("PRIMARY_CONTRIBUTOR", "Reproducible evidence shows the rule was necessary for the qualifying alert/decision.", "May continue if all other gates pass."),
            ("SUPPORTING_CONTRIBUTOR", "The rule materially supported the decision but was not uniquely necessary.", "Analyze under configured overlap/priority policy."),
            ("COINCIDENTAL_MATCH", "The rule matched but did not influence the qualifying decision.", "Do not optimize from this trigger."),
            ("UNRESOLVED", "Rule-group, decision, deleted history, or signal context cannot be reproduced.", "Manual review or limited recommendation only."),
        ],
        widths=[2300, 4600, 2460],
        font_size=8.2,
    )
    add_callout(
        doc,
        "Deleted or retired rule handling",
        "transaction_match/device_match references may point to versions that were deleted and are not enforced by foreign keys. "
        "Candidate generation requires an immutable historical configuration bundle and a current-version relevance check. "
        "The Agent may recommend retirement review, but it cannot retire, delete, or modify the rule.",
        fill=PALE_RED,
        accent=RED,
    )

    add_page_break(doc)
    add_section_intro(
        doc, "19", "Policy, configuration resolution, and rule-run aggregation",
        "Institution and channel differences are configuration, not code branches. A versioned policy snapshot determines eligibility, windows, "
        "mappings, permissions, operating mode, and analysis behavior for the entire run."
    )
    doc.add_heading("Proposed policy-resolution precedence", level=2)
    add_table(
        doc,
        ["Layer", "Examples", "Resolution behavior"],
        [
            ("Platform defaults", "Safe fallbacks, mandatory controls, supported schema versions.", "Always applied; cannot weaken mandatory security."),
            ("Institution policy", "Outcome mapping, operating mode, windows, retention, alert-volume limit, LLM policy.", "Overrides configurable defaults for that institution."),
            ("Source/channel/fact scope", "Mobile transfer versus device behavior; required metrics and maturity windows.", "Applies only when the trigger scope matches."),
            ("Rule/group policy", "Allowed mutation templates, cooldown, aggregation, replay requirements.", "Most specific approved override."),
        ],
        widths=[1900, 4200, 3260],
        font_size=8.4,
    )
    add_para(
        doc,
        "This precedence is proposed and must be approved. Every resolved value records its source policy ID, effective time, and policy hash. "
        "Decision-code meanings may change; historical labels therefore use the mapping effective for the relevant decision time, while final case decision still takes precedence. "
        "The optimizer never compares a client code directly with the literal value FALSE_POSITIVE. It first resolves a governed mapping containing institution_id, "
        "client_decision_code, canonical_outcome, optimization_eligible, effective_from, effective_to, mapping_version, mapping_hash, and approval_reference."
    )
    doc.add_heading("Effective configuration bundle", level=2)
    add_table(
        doc,
        ["Bundle area", "Required contents"],
        [
            ("Rule identity", "rule_master ID/code; rule_version ID/business version; lifecycle/status; created/effective evidence."),
            ("Formal logic", "rule_version.logic JSON; generated DRL; checksum metadata; rule_drl_context; compiler/runtime identity."),
            ("Group context", "rule_group_version_map, ordering, weights, group/source bindings, channel/fact, effective source mappings."),
            ("Decision context", "Immutable decision_policy_version ID/contents/checksum/effective period; rule_decision_upgrade; alert-generation behavior; canonical outcome-mapping version."),
            ("Data dependencies", "rule_required_data by rule_version; rule_metric_dependency by rule_master; source/engine attribute definitions and mappings."),
            ("Metric dependencies", "Immutable metric_definition_version ID, SQL text/checksum, window/context parameters, effective period, owner, approval reference, and bound parameters."),
            ("Policy", "Mode, reasoning policy, analysis windows, maturity, alert-volume constraint, allowed mutations, retention, LLM/RAG permission."),
        ],
        widths=[2200, 7160],
        font_size=8.3,
    )
    add_callout(
        doc,
        "Historical time-travel requirement",
        "The current schema does not provide complete validity intervals for every configuration object. Production must obtain historical truth from "
        "an authoritative history API, audit history, archived bundle, or a new immutable pre-change snapshot. Otherwise the correct state is CONFIGURATION_AMBIGUITY.",
        fill=PALE_RED,
        accent=RED,
    )
    add_callout(
        doc,
        "Immutable production-time execution identity",
        "The preferred historical source is a rule_execution_manifest linked to the evaluation/result. It pins the institution, event/evaluation identity and time, ordered group "
        "membership, exact rule/metric/decision/source/reference-data versions and hashes, compiler/Drools/JDK/application versions, event schema, and outcome-policy version. "
        "Metric SQL and decision policies must never be updated in place after use; every change creates a new effective-dated, approved immutable version. "
        "Without a complete manifest, legacy data is limited to claims that can be independently reconstructed and proven.",
        fill=PALE_RED,
        accent=RED,
    )
    doc.add_heading("Rule-run aggregation", level=2)
    add_para(
        doc,
        "Multiple false-positive cases can point to the same current rule issue. After case-level verification and attribution, a proposed aggregator "
        "can attach them to one rule-analysis run when institution, current configuration checksum, source/channel/fact, policy version, and analysis window match."
    )
    add_bullets(doc, [
        "A configurable cooldown prevents repeated expensive reruns for the same unchanged rule.",
        "A new current configuration checksum creates a new run; it never silently reuses stale evidence.",
        "A case correction removes or supersedes that case's trigger link and recalculates eligibility.",
        "Concurrency limits are applied per institution and workload class.",
        "The false-positive case remains the reason the rule was selected; the full historical population remains the evidence used to judge it.",
    ])

    add_page_break(doc)
    add_section_intro(
        doc, "20", "Historical evidence, point-in-time metrics, and data-quality gates",
        "The Snapshot Builder creates a frozen, reproducible analytical dataset. It does not send changing live rows directly to an LLM or calculate results over only false-positive cases."
    )
    add_figure(
        doc,
        evidence_pipeline_png,
        "Figure 7. Governed historical evidence pipeline from approved EFRM reads to an immutable snapshot and deterministic findings.",
        width=6.35,
    )
    doc.add_heading("Minimum snapshot manifest", level=2)
    add_table(
        doc,
        ["Manifest field", "What it proves"],
        [
            ("snapshot_id, run_id, institution, scope", "Ownership and tenant boundary."),
            ("analysis cutoff + configured windows", "Which event and outcome periods were included."),
            ("execution manifest ID/hash", "Which complete source execution context originally produced the stored result."),
            ("historical/current configuration hashes", "Which rules and policies defined eligibility and the current comparison target."),
            ("frozen event/payload + fact hashes", "That baseline and candidate use identical event-time inputs."),
            ("query/template IDs and hashes", "Which reviewed extraction logic ran."),
            ("metric value/SQL/parameter hashes", "Which point-in-time derived values, metric versions, and parameters were used."),
            ("reference/policy/outcome version hashes", "Which reference data, decision policy, and canonical outcome mapping governed interpretation."),
            ("included/excluded row counts by reason", "How the eligible population was constructed."),
            ("partition/object hashes", "That the frozen files have not changed."),
            ("timezone, event-time, processing-time rules", "How dates, late arrivals, and corrections were handled."),
            ("compiler/runtime/code compatibility identity", "Which compiler, Drools/JDK, application, library, and container versions produced or replayed the snapshot."),
            ("readiness permissions and limitations", "Which analyses, LLM reasoning, candidate generation, and replay are allowed."),
        ],
        widths=[3500, 5860],
        font_size=8.4,
    )
    doc.add_heading("Metric Definition Custom SQL control contract", level=2)
    add_table(
        doc,
        ["Control", "Required behavior"],
        [
            ("Immutable version identity", "Reference metric_definition_version; store approved SQL text, effective period, parameters, owner, approval reference, and calculated checksum."),
            ("Parser/AST gate", "Allow one read-only statement; reject writes, DDL, unsafe functions, dynamic execution, and unapproved schemas."),
            ("Bound parameters", "Require explicit institution/scope/event-time placeholders; never inject predicates into arbitrary SQL text."),
            ("Tenant/time enforcement", "Use secured views or row-level controls plus event-time parameters and tests for future-data leakage."),
            ("Workload limits", "Statement timeout, row/memory/temp limits, plan-cost threshold, isolated pool, cancellation, and audit."),
            ("Reproducibility", "Compare known-result fixtures and record METRIC_NOT_REPRODUCIBLE when historical inputs no longer exist."),
        ],
        widths=[2600, 6760],
        font_size=8.2,
    )
    add_para(
        doc,
        "The current runtime does not use aggregated_metric. Aggregate and derived values come from Metric Definition Custom SQL. "
        "The existing metric_definition schema does not itself provide complete institution versioning or validity history, so production requires immutable "
        "metric_definition_version records. Decision policies require the same version/effective-period rule. Analysis-time snapshots are audit evidence; they do not replace "
        "source-side versions that were pinned when the original evaluation executed."
    )
    doc.add_heading("Data-quality permission matrix", level=2)
    add_table(
        doc,
        ["Gate result", "Analysis", "LLM reasoning", "Candidate", "Replay", "Recommendation wording"],
        [
            ("READY", "Allowed", "Optional", "Mode/policy dependent", "Mode/policy dependent", "Evidence-supported."),
            ("READY_WITH_LIMITS", "Allowed subset", "Only approved aggregate evidence", "Usually blocked or tightly limited", "Only validated scope", "State every limitation."),
            ("INSUFFICIENT_EVIDENCE", "Descriptive only", "Template explanation preferred", "Blocked", "Blocked", "Monitor/review; no change claim."),
            ("LINEAGE/CONFIG/METRIC FAILURE", "Blocked", "Blocked", "Blocked", "Blocked", "Explicit technical stop reason."),
        ],
        widths=[1800, 1450, 1600, 1450, 1400, 1660],
        font_size=7.8,
    )

    add_section_intro(
        doc, "21", "Deterministic analytics: health, decay, thresholds, and leakage",
        "Every metric must define its numerator, denominator, time window, maturity rule, scope, minimum sample, uncertainty method, and version. "
        "The result is repeatable evidence, not an LLM opinion."
    )
    doc.add_heading("Core metric catalogue", level=2)
    add_table(
        doc,
        ["Metric", "Simple formula", "Use and limitation"],
        [
            ("Rule fire rate", "rule-fired eligible events / all rule-eligible events", "Selectivity and drift; depends on reconstructable eligibility."),
            ("Alert conversion rate", "qualifying alerts / rule-fired events", "Shows downstream policy impact; requires alert decision context."),
            ("Internal false-positive rate", "mapped mature false-positive outcomes / all mapped mature outcomes linked to the rule", "Investigation burden; not a population fraud precision estimate."),
            ("Mature internal outcome yield", "mapped mature non-false-positive outcomes / all mapped mature outcomes", "Internal usefulness proxy only; exact decision mapping is institution-specific."),
            ("Repeat burden", "alerts or cases for repeatedly affected entities / total selected alerts or cases", "Shows customer/operations concentration; entity keys must be governed."),
            ("Threshold proximity", "distribution of distance between event value and threshold", "Finds dense bands where a small change may matter."),
            ("Band outcome rate", "mapped outcomes in a value band / mature outcomes in that band", "Supports threshold candidates with denominators and uncertainty."),
            ("Rule overlap", "events matched by target and another rule / target matches", "Shows duplication; ordering and group context matter."),
            ("Unique internal coverage", "mature internal outcomes matched only by target / mature outcomes matched by target", "Estimates distinct control value inside observed EFRM outcomes."),
            ("Alert-volume impact", "candidate alerts - current baseline alerts", "The only confirmed operational-capacity constraint today."),
            ("Candidate internal leakage", "mature internal outcomes retained by current but not candidate / current-retained mature internal outcomes", "Does not measure unknown fraud outside EFRM observation."),
            ("Input-quality rate", "missing/invalid required inputs / eligible events", "Separates rule problems from source-data problems."),
        ],
        widths=[2100, 3150, 4110],
        font_size=7.9,
    )
    add_callout(
        doc,
        "Language discipline",
        "Without an approved independent truth source, do not call a mature non-false-positive EFRM outcome 'confirmed fraud', "
        "and do not present internal false-positive rate as true model precision. The recommendation must state exactly which internal outcome mapping was used.",
        fill=PALE_RED,
        accent=RED,
    )
    doc.add_heading("Decay detection sequence", level=2)
    add_numbered(doc, [
        "Resolve configurable historical baseline and recent windows, outcome-maturity delay, and analysis cutoff.",
        "Check minimum eligible events and minimum mature outcomes in both windows.",
        "Compare fire rate, internal false-positive rate, mature outcome yield, alert conversion, input missingness, and overlap.",
        "Calculate effect size and uncertainty using approved proportion/rate intervals or tests; adjust for repeated segment testing.",
        "Use configurable persistence checks such as consecutive time buckets, EWMA/CUSUM, or approved change-point detection.",
        "Separate population/input drift from rule-quality deterioration using segment and data-quality evidence.",
        "Classify STABLE, MONITOR, or DECAY only when magnitude, confidence, persistence, and data-quality rules all pass.",
    ])
    add_table(
        doc,
        ["Decay decision input", "Configurable policy"],
        [
            ("Windows", "Baseline/recent lengths, calendar alignment, seasonality comparison, and retention."),
            ("Minimum evidence", "Eligible-event and mature-outcome minimums per rule/segment."),
            ("Materiality", "Minimum absolute and relative deterioration before a warning."),
            ("Statistical confidence", "Confidence level/test and multiple-testing approach."),
            ("Persistence", "Number of consecutive buckets or confirmed change-point behavior."),
            ("Exclusions", "Known campaigns, outages, missing inputs, rule changes, and incomplete outcome maturity."),
        ],
        widths=[2500, 6860],
        font_size=8.4,
    )
    doc.add_heading("Threshold optimization and scenario recommendation", level=2)
    add_bullets(doc, [
        ("Generate bounded values. ", "Use approved ranges, step sizes, observed quantiles, and supported mutation templates; never free-form DRL editing."),
        ("Use temporal splits. ", "Generate on a development period, compare on validation, and report an untouched holdout when enough history exists."),
        ("Respect monotonicity and semantics. ", "Reject values that reverse intended rule meaning or violate dependency/schema checks."),
        ("Compare a Pareto set. ", "Show burden reduction against mature internal outcome retention, alert volume, overlap, and uncertainty."),
        ("Prefer no change when evidence is weak. ", "A valid output is MONITOR, REVIEW DATA, or NO CHANGE."),
        ("Do not claim policy compliance yet. ", "Only alert-volume capacity is currently known; missing hard constraints require CANDIDATE_FOR_REVIEW."),
    ])
    doc.add_heading("Risk leakage specification", level=2)
    add_table(
        doc,
        ["Measure", "Definition", "Safe claim"],
        [
            ("Candidate retention loss", "Current-retained mature internal outcomes not retained by candidate / current-retained mature internal outcomes.", "The candidate did not retain X of Y mature mapped outcomes."),
            ("Cross-rule rescue", "Lost target-rule outcomes still captured by another effective rule / target-rule outcomes lost.", "Other rules covered X of Y internally observed outcomes."),
            ("Unique-coverage loss", "Mature outcomes uniquely captured by current target but missed by candidate.", "The proposed change reduces unique internal coverage by the stated amount."),
            ("Delayed coverage", "Outcomes captured only by a later alert/rule after the candidate would stop matching.", "Coverage may be delayed under the simulated path."),
        ],
        widths=[2200, 3900, 3260],
        font_size=8.0,
    )

    add_section_intro(
        doc, "22", "Optional LLM and RAG architecture",
        "The LLM is optional. It can improve hypothesis formation and explanation after deterministic evidence exists, but it has no database credentials, "
        "no production-write tools, and no authority to declare a rule safe or approve a recommendation."
    )
    add_figure(
        doc,
        llm_loop_png,
        "Figure 8. Bounded LLM loop with masking, an allowlisted tool broker, exact deterministic tool results, structured output, and claim validation.",
        width=6.35,
    )
    doc.add_heading("LLM input contract", level=2)
    add_table(
        doc,
        ["Allowed input", "Examples", "Never include"],
        [
            ("Verified aggregate evidence", "Counts with denominators, rates, intervals, cohort labels, drift/threshold findings.", "Raw unrestricted transaction/customer tables."),
            ("Typed rule context", "Masked rule identity, structured logic subset, allowed mutation schema, current/candidate diff.", "Secrets, compiler credentials, production-write endpoints."),
            ("Evidence references", "Snapshot ID, finding IDs, metric definitions, limitations, policy hash.", "Unsupported numbers or hidden data."),
            ("Optional RAG passages", "Approved glossary, SOP, workflow, rule-language, or governance text with version and source.", "Live configuration, live rows, unapproved documents, or raw audit logs."),
        ],
        widths=[2100, 4300, 2960],
        font_size=8.2,
    )
    doc.add_heading("Approved LLM actions", level=2)
    add_bullets(doc, [
        "Explain a verified pattern in plain English.",
        "Propose a bounded hypothesis or supported scenario template.",
        "Request one of a small set of typed follow-up analytics tools.",
        "Draft recommendation wording from approved evidence and limitations.",
        "Summarize why no safe recommendation can be made.",
    ])
    doc.add_heading("Mandatory LLM controls", level=2)
    add_table(
        doc,
        ["Control", "Technical implementation"],
        [
            ("Provider-neutral gateway", "Same internal contract for local hosting or an approved external API."),
            ("Data loss prevention", "Field-level allowlist, masking/tokenization, payload scanner, institution policy, and egress block."),
            ("Prompt/tool injection defense", "Treat retrieved text as data; fixed system policy; typed tools; no executable instructions from documents."),
            ("Bounded execution", "Maximum iterations, tool calls, tokens, latency, cost, concurrency, and circuit breaker."),
            ("Structured output", "JSON Schema with allowed recommendation types, evidence IDs, uncertainty, and limitations."),
            ("Post-generation validation", "Reject unsupported numbers/references, forbidden operations, or ungrounded claims."),
            ("Auditability", "Record model/provider, prompt/tool versions, request ID, latency, token/cost use, and fallback."),
            ("Deterministic fallback", "Template recommendation when unavailable, disabled, timed out, or invalid."),
            ("Incident control", "Per-institution disable switch and provider egress kill switch."),
        ],
        widths=[2500, 6860],
        font_size=8.2,
    )
    doc.add_heading("RAG decision", level=2)
    add_para(
        doc,
        "RAG is optional and should not be a first-release dependency. It is valuable for stable explanatory knowledge: approved rule descriptions, "
        "business glossary, decision meanings, SOPs, compiler/rule-language documentation, and governance policy. Live rule JSON/DRL, case rows, "
        "transaction/device facts, metrics, and simulation results must come from governed services and deterministic artifacts."
    )
    add_bullets(doc, [
        "If enabled, preserve document ACLs, institution scope, effective date, source URI, document version, chunk hash, and citation metadata.",
        "Require approved ingestion, poisoning checks, re-indexing on change, deletion/retention handling, and retrieval-quality tests.",
        "The LLM may quote or summarize retrieved text only when the recommendation can identify the source and version.",
    ])

    add_section_intro(
        doc, "23", "Operating modes and candidate lifecycle",
        "The institution policy chooses the operating mode before candidate execution. The LLM does not choose the mode, and neither mode grants production-rule write access."
    )
    doc.add_heading("Recommendation Only - exact behavior", level=2)
    add_numbered(doc, [
        "Complete trigger, lineage, attribution, configuration, snapshot, data-quality, and deterministic analytics stages.",
        "Optionally run non-executable sensitivity analysis over historical values when it does not require a formal candidate.",
        "Optionally use the LLM to explain the verified pattern and propose a bounded human review action.",
        "Do not create executable candidate JSON, DRL, compiler requests, or candidate backtests.",
        "State clearly that projected candidate alert/case impact was not proven by executable replay.",
        "Publish REVIEW, CONSIDER, MONITOR, NO CHANGE, REVIEW DATA, or REVIEW FOR RETIREMENT language for human governance.",
    ])
    doc.add_heading("Candidate + Test - exact behavior", level=2)
    add_numbered(doc, [
        "Start only after the current-version relevance gate confirms the triggering issue still applies.",
        "Generate a bounded candidate intent from supported mutation templates and policy ranges.",
        "Create structured candidate JSON only in optimizer storage; never write the Rule Engine's production tables.",
        "Validate syntax, types, dependencies, data availability, supported scenario form, alert-volume rule, and known policy constraints.",
        "Call the existing JSON-to-DRL compiler through RuleCompilerPort after its real exposure is confirmed.",
        "Run historical parity, current baseline, and candidate simulation in the isolated replay boundary.",
        "Compare validation/holdout evidence and label the result CANDIDATE_FOR_REVIEW until all hard constraints are approved.",
    ])
    add_table(
        doc,
        ["Candidate type", "First-release support", "Why"],
        [
            ("Numeric threshold", "Recommended", "Bounded, explainable, easy to validate, and compatible with sensitivity curves."),
            ("Time/window value", "Recommended when point-in-time metrics are reproducible", "Useful for new-device/velocity behavior but requires correct event-time state."),
            ("Existing boolean condition enable/disable", "Possible after dependency and policy approval", "Formal but may materially change scenario meaning."),
            ("Add approved existing condition", "Later controlled template", "Requires required-data, compiler, and full-context replay support."),
            ("Free-form new scenario or DRL", "Not supported", "Too difficult to govern, validate, secure, and explain safely."),
        ],
        widths=[2500, 3500, 3360],
        font_size=8.1,
    )
    doc.add_heading("Candidate search and ranking", level=2)
    add_para(
        doc,
        "The deterministic generator can enumerate a bounded grid or use an approved optimization method over allowed values. "
        "Candidates are ranked only after simulation using a Pareto comparison: reduce unnecessary alert burden while retaining mature internal outcomes, "
        "respecting the alert-volume constraint, limiting unique-coverage loss, and reporting uncertainty. A tie-breaker should prefer the smallest explainable change."
    )
    add_callout(
        doc,
        "Current constraint limitation",
        "The full institution-specific hard constraint set is not yet defined. Therefore the Agent cannot claim that a candidate is safe or policy compliant. "
        "It can only present tested evidence and mark the candidate for human review.",
        fill=PALE_ORANGE,
        accent=ORANGE,
    )

    add_section_intro(
        doc, "24", "Replay, backtesting, and current-versus-candidate evaluation",
        "Backtesting is trustworthy only when the isolated runner can first reproduce historical behavior. The comparison must use the same frozen cohort, "
        "event order, point-in-time facts, and rule-group/decision context."
    )
    add_figure(
        doc,
        replay_png,
        "Figure 9. Historical parity replay is distinct from the current baseline and candidate simulation.",
        width=6.35,
    )
    add_callout(
        doc,
        "Existing replay capability is a foundation, not parity proof",
        "The external engineering review reports existing Rule Engine bulk/replay endpoints. After owner verification, RuleReplayPort can adapt those APIs. "
        "Their existence alone does not prove historical parity. Candidate-impact claims remain blocked until frozen payloads, point-in-time metrics, reference data, "
        "decision/outcome policies, compiler/runtime versions, event order/state, and recorded baseline results can all be reproduced within approved tolerance.",
        fill=PALE_ORANGE,
        accent=ORANGE,
    )
    doc.add_heading("Historical parity contract", level=2)
    add_table(
        doc,
        ["Dimension", "Required comparison"],
        [
            ("Configuration", "Rule/group/decision/upgrade/source binding effective for each historical event."),
            ("Runtime", "Compiler, DRL, Drools, JDK, metric code/SQL, and relevant library versions."),
            ("Event semantics", "Event-time ordering, velocity/state windows, late-arrival rule, timezone normalization, and cutoff."),
            ("Recorded result", "Match outcome, signal/reason fields, score/weight contribution, and decision where reproducible."),
            ("Tolerance", "Exact or owner-approved tolerance by field; never silently accept material mismatch."),
            ("Diagnostics", "Mismatch cohort, missing state/input, version difference, ordering difference, and non-reconstructable reason."),
        ],
        widths=[2300, 7060],
        font_size=8.3,
    )
    doc.add_heading("Fair current-versus-candidate experiment", level=2)
    add_bullets(doc, [
        "Freeze one evaluation cohort and one manifest.",
        "Run the current target configuration and candidate against identical events and point-in-time facts.",
        "Preserve the full effective rule group and downstream decision/alert policy, replacing only the selected rule.",
        "Use development, validation, and holdout periods when sufficient data exists; never tune and report only on the same rows.",
        "Store event-level deltas and aggregate metrics with denominators, intervals, and exclusions.",
        "Use deterministic seeds for any approved randomized method.",
    ])
    doc.add_heading("Permitted impact claims", level=2)
    add_table(
        doc,
        ["Replay scope", "Permitted claim", "Not permitted"],
        [
            ("Selected rule only", "Change in selected-rule matches/signals on the frozen cohort.", "Alerts or cases prevented."),
            ("Full rule group + decision policy", "Change in qualifying group decisions/signals.", "Case impact unless alert/case creation is also reproduced."),
            ("Full group + alert pipeline", "Projected alert changes with stated assumptions.", "Real-world operational outcome without rollout evidence."),
            ("Full downstream case logic", "Projected case changes inside the simulated pipeline.", "Confirmed fraud/loss avoided without an independent truth source."),
        ],
        widths=[2400, 3800, 3160],
        font_size=8.1,
    )
    doc.add_heading("Candidate evaluation result", level=2)
    add_table(
        doc,
        ["Decision result", "When used"],
        [
            ("CANDIDATE_FOR_REVIEW", "Evidence is sufficient and the candidate improves at least one objective without a known hard-constraint failure."),
            ("NO_MATERIAL_IMPROVEMENT", "Differences are too small, uncertain, unstable, or fail holdout validation."),
            ("RISK_RETENTION_CONCERN", "Internal outcome retention or unique coverage loss exceeds an approved review boundary."),
            ("ALERT_VOLUME_CONSTRAINT_FAILED", "The candidate violates the configured alert-volume limit."),
            ("POLICY_CONSTRAINTS_INCOMPLETE", "Evidence is available, but unconfirmed hard constraints prevent a safety/compliance claim."),
            ("REPLAY_MISMATCH", "The simulator cannot reproduce approved baseline behavior; candidate impact claims stop."),
        ],
        widths=[3100, 6260],
        font_size=8.5,
    )

    add_section_intro(
        doc, "25", "Recommendation contract, human review, and reviewer experience",
        "The main output is one concise, structured recommendation with drill-down evidence. The UI is an adapter over this contract, so the workflow remains valid "
        "whether the existing Agentic portal is reused or a different approved reviewer surface is selected."
    )
    add_callout(
        doc,
        "Reported Agentic portal foundation - pending verification",
        "The external engineering review reports existing routes for recommendations, approvals, orchestration, explainability, and analytics. These pages may be reused as "
        "the reviewer UI foundation after the owning team verifies them. Production use still requires the Rule Optimization backend, secured institution-scoped APIs, "
        "real evidence drill-down, stale/superseded states, reviewer authorization, four-eyes workflow, immutable audit, and removal of sample/mock state.",
        fill=PALE_ORANGE,
        accent=ORANGE,
    )
    add_figure(
        doc,
        recommendation_ui_png,
        "Figure 10. Illustrative reviewer contract mapped to the reported portal foundation; reuse and production readiness remain subject to owner verification.",
        width=6.35,
    )
    doc.add_heading("Expanded recommendation contract", level=2)
    add_code_block(
        doc,
        '{\n'
        '  "recommendation_id": "REC-...",\n'
        '  "case_job_ids": ["JOB-..."], "rule_analysis_run_id": "RUN-...",\n'
        '  "institution_id": "...", "source_channel_fact": {"source": "...", "channel": "...", "fact": "..."},\n'
        '  "operating_mode": "RECOMMENDATION_ONLY | CANDIDATE_AND_TEST",\n'
        '  "reasoning_mode": "DETERMINISTIC_ONLY | LLM_ASSISTED",\n'
        '  "rule": {"master_id": "...", "historical_version_id": "...", "current_version_id": "...",\n'
        '           "historical_hash": "...", "current_hash": "...", "group_context_hash": "..."},\n'
        '  "source_execution": {"manifest_id": "...", "manifest_hash": "...", "verification": "VERIFIED | LEGACY_LIMITED"},\n'
        '  "attribution": {"status": "PRIMARY_CONTRIBUTOR", "evidence_ids": ["EV-..."]},\n'
        '  "policy": {"policy_hash": "...", "outcome_mapping_version": "...", "constraints_status": "INCOMPLETE"},\n'
        '  "snapshot": {"snapshot_id": "...", "hash": "...", "window": "...", "cutoff": "..."},\n'
        '  "issue": {"type": "...", "summary": "...", "finding_ids": ["F-..."]},\n'
        '  "recommended_action": {"type": "REVIEW_CHANGE", "current": "...", "proposed": "..."},\n'
        '  "candidate": {"created": true, "candidate_id": "...", "diff_hash": "...", "compiler_status": "PASSED"},\n'
        '  "simulation": {"performed": true, "parity": "PASSED", "scope": "FULL_GROUP", "backtest_id": "..."},\n'
        '  "evidence_summary": [{"metric": "...", "value": 0.0, "denominator": 0,\n'
        '                        "uncertainty": "...", "evidence_id": "EV-..."}],\n'
        '  "limitations": ["..."], "evidence_strength": "SUFFICIENT_FOR_REVIEW",\n'
        '  "stale_state": "CURRENT", "action": "HUMAN_REVIEW_REQUIRED",\n'
        '  "audit_correlation_id": "..."\n'
        '}'
    )
    doc.add_heading("What the Claim Validator can and cannot prove", level=2)
    add_table(
        doc,
        ["Can verify automatically", "Cannot prove by itself"],
        [
            ("Every numeric value exists in an evidence artifact with the same denominator and scope.", "That unrestricted causal prose is scientifically true."),
            ("Every evidence ID, snapshot hash, policy/config hash, and replay ID exists and is current.", "That an institution's business risk appetite has been satisfied when constraints are missing."),
            ("Candidate/test wording matches actual mode and replay status.", "That a simulated outcome will occur unchanged in production."),
            ("Forbidden claims and unsupported production actions are absent.", "That mature internal outcomes equal confirmed fraud without an approved truth mapping."),
        ],
        widths=[4680, 4680],
        font_size=8.3,
    )
    doc.add_heading("Human governance actions", level=2)
    add_table(
        doc,
        ["Action", "System behavior"],
        [
            ("View evidence", "Open metric definitions, cohorts, queries/code versions, snapshot manifest, candidate diff, and replay artifacts."),
            ("Request more analysis", "Create a governed follow-up request; do not silently mutate the completed run."),
            ("Reject / no action", "Record reason and reviewer identity; keep immutable recommendation and audit."),
            ("Send to existing governance", "After authorized human acceptance and a fresh stale-state check, create a governance reference or optional draft change work item when institution policy and an approved API permit it. Never activate or update a rule."),
            ("Review for retirement", "Send advice to governance; the Agent cannot retire or delete the rule."),
            ("Case/rule changed", "Mark recommendation STALE or SUPERSEDED and prevent action on old evidence."),
        ],
        widths=[2500, 6860],
        font_size=8.4,
    )
    add_callout(
        doc,
        "Four-eyes principle",
        "The reviewer who accepts the recommendation should not be the only person who can promote a production rule. "
        "The existing governance process should retain role separation, approval evidence, comments/e-signature where required, and rollback ownership.",
        fill=PALE_TEAL,
        accent=TEAL,
    )

    add_section_intro(
        doc, "26", "Optimizer persistence, APIs, events, and provenance",
        "The optimizer needs its own auditable persistence model. Existing EFRM tables remain authoritative and read-only; new jobs, snapshots, evidence, "
        "candidates, simulations, recommendations, and review metadata are stored inside the optimizer boundary."
    )
    doc.add_heading("Proposed optimizer persistence model", level=2)
    add_table(
        doc,
        ["Entity", "Purpose", "Important keys / contents"],
        [
            ("optimization_trigger_inbox", "Deduplicate event/API/poll deliveries.", "institution, source event/action ID or approved stable hash, payload hash, received status."),
            ("optimization_policy_snapshot", "Freeze resolved policy and canonical outcome mapping.", "institution, client code, canonical outcome, optimization eligibility, effective/version/hash, mode, windows, constraints, LLM/RAG permissions."),
            ("case_analysis_job", "Track one final case-decision analysis.", "case ID, decision identity, trigger ID, status, cancellation/supersession."),
            ("rule_analysis_run", "Track shared rule-level analysis.", "institution, current rule/config hash, source/channel/fact, window, state."),
            ("trigger_case_link", "Connect cases/triggers to a rule-analysis run.", "attribution status, eligibility, inclusion/exclusion reason."),
            ("execution_manifest_ref", "Reference the immutable source execution context.", "source manifest ID/hash, institution, evaluation/result identity, availability and verification status."),
            ("configuration_bundle", "Store canonical historical/current configuration.", "execution-manifest link, bundle type/effective time, rule/group/metric/policy/JSON/DRL hashes, compiler/runtime metadata."),
            ("snapshot_manifest", "Describe immutable historical partitions.", "execution/config/payload/metric/reference/policy hashes, scope/window/cutoff, query/code hashes, row counts, partition/object hashes."),
            ("evidence_item", "Index findings, metrics, cohorts, files, and limitations.", "type, content hash, schema version, source lineage, artifact location."),
            ("candidate_artifact", "Store bounded candidate outside production.", "intent, JSON diff, validation, compiler artifact, current-config parent hash."),
            ("replay_run", "Track parity/baseline/candidate execution.", "scope, runtime versions, cohort, result hashes, tolerance, status."),
            ("recommendation", "Store canonical recommendation and readable view.", "mode, status, evidence strength, stale/superseded state, provenance."),
            ("review_decision", "Record human action and any permitted draft-governance request.", "reviewer role, action, reason, timestamp, approval chain, draft/governance reference."),
            ("optimization_audit_event", "Append-only business/technical audit.", "correlation, actor/service, action, before/after refs, timestamp."),
            ("optimization_outbox", "Reliably publish recommendation/status events.", "aggregate ID, event type/version, payload hash, publish state."),
        ],
        widths=[2600, 3100, 3660],
        font_size=7.7,
    )
    doc.add_heading("External and internal API contracts", level=2)
    add_table(
        doc,
        ["Contract", "Purpose", "Key behavior"],
        [
            ("POST /v1/optimization-triggers/case-finalized", "API alternative to event trigger.", "Service auth, institution authorization, versioned body, idempotency key, 202 + existing/new job ID."),
            ("GET /v1/case-analysis-jobs/{id}", "Read case-job state and linked rule runs.", "Institution-scoped, read-only, exposes stop/retry/limit reasons."),
            ("GET /v1/rule-analysis-runs/{id}", "Read stage states and evidence references.", "No raw unrestricted data; stable artifact links and hashes."),
            ("GET /v1/recommendations/{id}", "Retrieve canonical recommendation and reviewer view.", "Role-based field filtering, stale status, provenance."),
            ("POST /v1/recommendations/{id}/review", "Record human review intent.", "Idempotent action, current-state guard, reviewer role, comments/reason."),
            ("RuleConfigurationPort", "Retrieve current/historical bundles.", "Authoritative service preferred; read adapter only when time travel is proven."),
            ("RuleCompilerPort", "Compile typed candidate through unknown existing exposure.", "No direct production deployment method in this interface."),
            ("RuleReplayPort", "Adapt verified bulk/replay APIs or an isolated runner.", "Capability matrix and parity gate determine permitted claims; no production activation method."),
            ("EFRMHistoricalDataPort", "Retrieve approved bulk historical scope.", "API/replica/analytical implementation hidden behind typed contract."),
            ("ModelProviderPort", "Call local or approved external LLM.", "Masked input, structured output, audit, timeout, deterministic fallback."),
        ],
        widths=[3200, 2700, 3460],
        font_size=7.8,
    )
    doc.add_heading("Versioned events", level=2)
    add_table(
        doc,
        ["Event", "Producer -> consumer", "Minimum contract"],
        [
            ("CASE_FINALIZED", "Case integration -> Trigger Gateway", "schema version, event ID, institution/case IDs, finalization/action identity, occurred/created time, correlation; decision code may be included but is re-verified."),
            ("CASE_DECISION_CORRECTED", "Case integration -> Orchestrator", "authoritative correction/reopen identity, prior reference, new state; cancels/supersedes old work."),
            ("RULE_ANALYSIS_STARTED/COMPLETED", "Optimizer -> audit/operations", "run ID, current config hash, stage/status, limitation/error code."),
            ("RECOMMENDATION_READY", "Optimizer outbox -> UI/governance integration", "recommendation ID, institution, rule/current hash, mode, review status; no sensitive evidence payload."),
            ("RECOMMENDATION_SUPERSEDED", "Optimizer -> UI/governance integration", "recommendation ID, reason, replacement/current state."),
            ("REVIEW_DECISION_RECORDED", "UI/governance -> optimizer", "review action ID, reviewer role, decision, comments/reason, governance reference."),
        ],
        widths=[2450, 2600, 4310],
        font_size=8.0,
    )
    add_callout(
        doc,
        "Idempotency identity",
        "The supplied case schema does not confirm a decision_version column. Prefer an authoritative outbox/audit event ID or case event/action ID. "
        "If none exists, define and approve a stable hash over case ID, final decision/approval state and timestamp, institution, and schema version.",
        fill=PALE_ORANGE,
        accent=ORANGE,
    )
    doc.add_heading("End-to-end provenance chain", level=2)
    add_code_block(
        doc,
        "event_id\n"
        "  -> trigger_inbox_id\n"
        "  -> case_analysis_job_id\n"
        "  -> rule_analysis_run_id\n"
        "  -> rule_execution_manifest_id + source checksum\n"
        "  -> metric/decision/outcome/source/reference version hashes\n"
        "  -> historical/current configuration hashes\n"
        "  -> snapshot_id + partition/query/code hashes\n"
        "  -> finding/evidence IDs\n"
        "  -> candidate/compiler/replay IDs (Candidate + Test only)\n"
        "  -> recommendation_id\n"
        "  -> review action + governance reference"
    )

    add_section_intro(
        doc, "27", "Security, privacy, deployment boundaries, and technology status",
        "Security is enforced by architecture, credentials, data contracts, and isolation—not by asking the LLM to behave. "
        "The optimizer must be incapable of modifying production rules even if a component is misused."
    )
    add_figure(
        doc,
        security_png,
        "Figure 11. Mandatory trust zones and permissions, with proposed implementation products kept separate from confirmed EFRM facts.",
        width=6.35,
    )
    doc.add_heading("Mandatory security controls", level=2)
    add_table(
        doc,
        ["Control area", "Required design"],
        [
            ("Service identity", "OAuth2 workload identity and/or mTLS according to EFRM standards; no shared human credentials."),
            ("Authorization", "RBAC/ABAC plus institution, source, channel, fact, artifact, and reviewer-action scope."),
            ("Institution identity invariant", "institution_id is mandatory on every trigger, job, run, configuration/manifest, snapshot, evidence item, candidate, replay, recommendation, review record, query, join, cache key, unique constraint, artifact path, and audit event. Rule or metric code is never used alone."),
            ("Database access", "Read-only EFRM roles; secured views/RLS where available; institution-qualified joins and predicates; optimizer writes only to its own schema/storage."),
            ("Object storage", "Institution/environment prefix isolation, short-lived access, encryption, checksums, retention/legal hold."),
            ("Network", "Private service paths, broker ACLs, egress deny-by-default, isolated compiler/replay workers, no production activation route."),
            ("Sensitive data", "Minimize collection; mask/tokenize identifiers; separate raw-event role from aggregate/recommendation role."),
            ("Secrets and keys", "Approved secrets manager, KMS encryption, rotation, audit, and no secrets in prompts/logs/artifacts."),
            ("LLM egress", "Per-institution approval, DLP, provider no-training/no-retention terms, residency/retention review, kill switch."),
            ("Audit", "Append-only correlated logs for access, queries, artifacts, models, compiler/replay, recommendation, and human review."),
            ("Environment separation", "Separate development/test/UAT/production identities, data, topics/queues, buckets, schemas, and model policies."),
        ],
        widths=[2400, 6960],
        font_size=8.2,
    )
    doc.add_heading("Technology status: do not confuse facts with choices", level=2)
    add_table(
        doc,
        ["Status", "Technology / capability"],
        [
            ("Confirmed existing", "PostgreSQL efrm schema and JSONB; Spring-based Rule Engine family; Drools/DRL; structured rule JSON and generated DRL; formal JSON-to-DRL mechanism; Metric Definition Custom SQL; transaction/device/case lineage structures."),
            ("Initial proposed baseline", "One modular Java/Spring Boot Rule Optimization Service; PostgreSQL job/state/inbox/outbox and optimizer schema; institution-scoped read-only analytical connection; deterministic SQL/Python; RuleConfigurationPort, RuleCompilerPort, and RuleReplayPort; isolated replay only where required."),
            ("Optional scale/evolution choices", "Kafka event adapter; Temporal/Camunda workflow; Parquet + object storage; separately scaled analytics/replay pools; provider-neutral Model Gateway; approved documentation RAG."),
            ("Reviewer-reported, unverified", "Existing bulk/replay endpoints and Agentic portal routes; current CASE_CLOSED behavior; incomplete execution/metric/policy history. Owning teams must verify repository revision and capability."),
            ("Unresolved", "Actual infrastructure products and versions; official read/replay/compiler/governance APIs; production sizing/retention/SLO/RTO/RPO; LLM vendor/hosting/egress approval."),
        ],
        widths=[2100, 7260],
        font_size=8.2,
    )
    doc.add_heading("Product-neutral ports preserve architecture", level=2)
    add_bullets(doc, [
        "EventTransportPort: broker, webhook, or other approved transport.",
        "WorkflowPort: transactional PostgreSQL state machine initially; Temporal or Camunda only when scale/operability evidence justifies it.",
        "ObjectStorePort: S3-compatible or approved enterprise artifact store.",
        "RuleConfigurationPort and EFRMHistoricalDataPort: service API, read replica, or analytical service.",
        "RuleCompilerPort: library, REST, gRPC, build service, or another confirmed compiler interface.",
        "RuleReplayPort: verified existing replay/bulk API or isolated compatible runner behind one parity contract.",
        "ModelProviderPort: local model endpoint or approved external LLM API.",
    ])
    doc.add_heading("Secure delivery and supply chain", level=2)
    add_bullets(doc, [
        "Version-controlled configuration and database migrations with named owners and rollback plans.",
        "CI/CD gates for unit/contract/security tests, dependency and secret scanning, signed artifacts, and SBOM.",
        "Compatibility matrix for trigger schemas, artifact schemas, compiler/Drools/JDK versions, and analytics code.",
        "Shadow/canary rollout per institution; feature flags for LLM, RAG, Recommendation Only, and Candidate + Test.",
        "No production credentials in developer/test environments and no production data copied to lower environments without approved masking.",
    ])

    add_section_intro(
        doc, "28", "Reliability, recovery, observability, and production operations",
        "Enterprise reliability means every stage can be retried or stopped safely, every expensive operation can be resumed, and every recommendation can be traced to its exact inputs."
    )
    doc.add_heading("Failure classification and recovery", level=2)
    add_table(
        doc,
        ["Situation", "Class", "System response"],
        [
            ("Duplicate event/API retry", "Normal duplicate", "Return existing trigger/job identity; do not re-run."),
            ("Broker/API/read-service unavailable", "Transient", "Bounded exponential retry with jitter; alert on age/attempt threshold."),
            ("Replica lag or late data", "Transient or wait", "Delay until freshness policy passes; record cutoff/freshness."),
            ("Worker crash during snapshot/replay", "Recoverable", "Lease expires; resume from checkpointed partition using same stage ID."),
            ("Query timeout/resource limit", "Transient or permanent", "Retry only if policy permits; otherwise limited/failed stage with diagnostics."),
            ("Metric SQL rejected/unreproducible", "Permanent evidence defect", "METRIC_NOT_REPRODUCIBLE; block unsafe downstream work."),
            ("LLM timeout/invalid output", "Optional-service failure", "Circuit breaker and deterministic template fallback."),
            ("Candidate validation/compiler rejection", "Candidate-specific", "Record reason; try next bounded candidate or produce Recommendation Only evidence."),
            ("Historical replay mismatch", "Trust failure", "Stop candidate impact claims; publish REPLAY_MISMATCH diagnostics."),
            ("Case reopened/corrected", "Supersession", "Cancel case job and supersede recommendations tied to old decision."),
            ("Current rule/config changed", "Stale state", "Mark run/recommendation STALE; require new run or explicit refresh."),
            ("Repeated unhandled failure", "Operational", "Dead-letter/manual rescue with immutable context and runbook."),
        ],
        widths=[2800, 1900, 4660],
        font_size=7.9,
    )
    doc.add_heading("Backpressure and scale design", level=2)
    add_bullets(doc, [
        "Partition queues, snapshots, and artifact prefixes by institution and date/workload class.",
        "Apply per-institution concurrency, query, LLM, compiler, and replay quotas so one client cannot exhaust the platform.",
        "Use checkpointed snapshot/replay partitions and ordered processing where velocity/stateful rules require event sequence.",
        "Keep control modules in one deployable service initially; separate analytical/replay workload pools or processes only when isolation, scale, or ownership requires it.",
        "Use bounded retention, archival, and legal-hold policy configurable per institution.",
        "Derive capacity only from future representative production workload tests; never from the supplied sample row counts.",
    ])
    doc.add_heading("Observability model", level=2)
    add_table(
        doc,
        ["Signal", "Examples"],
        [
            ("Workflow", "Jobs/stages by state, stage latency, attempts, cancellations, stale/superseded count, oldest-job age."),
            ("Trigger", "Outbox lag, consumer lag, duplicate rate, inbox errors, poller cursor/freshness, missed-trigger reconciliation."),
            ("Data", "Snapshot throughput/size, excluded rows, missing inputs, replica freshness, metric SQL failures/timeouts."),
            ("Analysis", "Runs by evidence status, decay/threshold module failures, uncertainty/sample gates, no-change rate."),
            ("Replay", "Compiler failures, parity pass rate, mismatch dimensions, event throughput, checkpoint/retry rate."),
            ("LLM", "Enabled institutions, latency, tokens/cost, tool calls, schema/claim rejection, fallback/circuit-breaker rate."),
            ("Security", "Authorization denials, cross-tenant guard failures, unusual query/egress, sensitive-data scanner hits."),
            ("Governance", "Ready/reviewed/rejected/stale recommendations, reviewer cycle time, requests for more analysis."),
        ],
        widths=[2200, 7160],
        font_size=8.3,
    )
    doc.add_heading("SLO, RTO, and RPO approach", level=2)
    add_para(
        doc,
        "Exact numeric targets cannot be confirmed without owner requirements and representative production load. Define separate owner-approved targets for trigger acceptance, "
        "job completion by workload class, recommendation availability, event/data freshness, evidence durability, recovery time, recovery point, and support response. "
        "Recommendation generation is asynchronous and must not affect the live transaction-monitoring SLO."
    )
    add_bullets(doc, [
        "Dashboards must support institution-safe drill-down from event ID to governance reference.",
        "Runbooks must cover trigger backlog, data freshness, Custom SQL rejection, replay mismatch, LLM incident disablement, stale recommendations, and dead-letter rescue.",
        "Backup/restore and disaster-recovery exercises must prove metadata, immutable evidence, and audit restoration before production approval.",
    ])

    add_page_break(doc)
    add_section_intro(
        doc, "29", "Testing strategy and formal acceptance gates",
        "The Agent must be tested as a financial decision-support system. Tests cover data lineage, deterministic mathematics, replay fidelity, tenant isolation, "
        "LLM claim quality, resilience, and the human governance hand-off."
    )
    doc.add_heading("Enterprise test matrix", level=2)
    add_table(
        doc,
        ["Test class", "What must be proven"],
        [
            ("Unit tests", "Policy resolution, mapping precedence, state guards, formulas, uncertainty, candidate validation, claim rules."),
            ("Schema/contract tests", "Backward/forward compatibility for events, APIs, artifact schemas, and ports."),
            ("Transaction lineage fixtures", "Correct polymorphic routing, result/match siblings, master/request scope, deleted-rule behavior."),
            ("Device lineage fixtures", "Equivalent device path, institution cross-check, missing/ambiguous references."),
            ("Attribution golden cases", "Known primary/supporting/coincidental/unresolved results under group/policy context."),
            ("Point-in-time metric tests", "No future-data leakage; correct windows, late arrivals, timezones, SQL checksum, and reproducibility."),
            ("Analytical golden sets", "Known metric numerators/denominators, threshold bands, decay states, overlap, and leakage calculations."),
            ("Compiler/replay parity", "Historical event-level parity under pinned versions, event order, and stateful windows."),
            ("Backtest fairness", "Identical cohort/context, temporal validation/holdout, stable seeds, allowed claim scope."),
            ("Security tests", "Cross-tenant penetration, authorization, SQL abuse, object-store isolation, secrets, egress, audit integrity."),
            ("LLM evaluation", "Evidence citation fidelity, unsupported-claim rejection, prompt/tool injection, sensitive-data leakage, fallback."),
            ("Performance/load", "Trigger bursts, query/snapshot throughput, worker backpressure, large-rule replay, per-tenant fairness."),
            ("Resilience/DR", "Worker crash/resume, broker/API outage, replica lag, dead-letter rescue, backup/restore, regional recovery."),
            ("UAT/governance", "Fraud/risk reviewer understands recommendation, evidence, limitations, stale state, and actions."),
        ],
        widths=[2600, 6760],
        font_size=7.9,
    )
    doc.add_heading("Release acceptance gates", level=2)
    add_table(
        doc,
        ["Gate", "Required evidence"],
        [
            ("Contract approval", "Case, outcome mapping, rule identity, policy, data access, artifact, and recommendation owners sign off."),
            ("Tenant/security approval", "Threat model, roles, read-only proof, isolation tests, external-LLM policy, audit design."),
            ("Lineage/config approval", "Representative transaction/device paths and time-travel configuration are reproducible."),
            ("Data/metric approval", "Snapshot manifest, Custom SQL controls, known-result tests, maturity and sufficiency rules."),
            ("Recommendation Only approval", "Deterministic evidence and human recommendation work end to end without compiler/LLM dependency."),
            ("LLM approval", "Red-team, DLP, claim fidelity, fallback, model/provider policy, and incident disablement pass."),
            ("Replay approval", "Pinned runtime and historical parity meet owner-approved dimensions/tolerances."),
            ("Candidate + Test approval", "Bounded mutations, holdout comparison, claim boundaries, no production-write path."),
            ("Operational approval", "Load/SLO tests, dashboards, alerts, runbooks, backups, restore/DR, and on-call ownership."),
            ("Business UAT", "Fraud/risk/governance teams approve usability, evidence, limitations, and reviewer workflow."),
        ],
        widths=[2700, 6660],
        font_size=8.1,
    )
    add_callout(
        doc,
        "Use of the supplied sample database",
        "The sample dump is suitable for schema discovery, contract fixtures, lineage tests, and demonstration. "
        "Its row counts, inconsistencies, distributions, and measured performance must not be used as production sizing, model-quality, or business-performance evidence.",
        fill=PALE_ORANGE,
        accent=ORANGE,
    )

    add_page_break(doc)
    add_section_intro(
        doc, "30", "Implementation roadmap, ownership, and approval decisions",
        "The recommended sequence completes provenance and a modular deterministic Recommendation Center first, proves point-in-time replay next, "
        "then adds bounded candidates and finally optional LLM/RAG enhancements."
    )
    doc.add_heading("Controlled delivery roadmap", level=2)
    add_table(
        doc,
        ["Stage", "Main deliverables", "Exit gate"],
        [
            ("0. P0 provenance and contracts", "Case outbox; canonical outcome versions; rule execution manifest; immutable metric/policy versions; tenant, retention, read, audit, and governance contracts.", "Owners prove reliable trigger delivery and complete historical identity."),
            ("1. Modular trigger and lineage", "One service; PostgreSQL workflow/inbox/outbox; API/poller adapters; case jobs; corrected transaction/device lineage and corrections.", "Idempotent institution-scoped end-to-end traces."),
            ("2. Configuration and attribution", "Execution-manifest resolver; historical/current bundles; deterministic attribution; current-version gate; rule-run aggregation.", "No unresolved identity/configuration/attribution in approved fixtures."),
            ("3. Snapshot and metric foundation", "Historical data port; eligibility reconstruction; point-in-time Custom SQL; immutable manifests; data-quality permissions.", "Reproducible frozen datasets and explicit stop/limit states."),
            ("4. Deterministic Recommendation Only", "Health/threshold/decay/overlap/workload/leakage evidence; claim validation; recommendation contract.", "First useful capability without LLM, compiler, or production writes."),
            ("5. Recommendation Center and governance", "Secured portal APIs, evidence drill-down, stale/superseded handling, four-eyes review, optional governed draft request.", "Real reviewer workflow and immutable audit; sample UI state removed."),
            ("6. Point-in-time replay platform", "RuleReplayPort; frozen payload/metric/reference/policy inputs; pinned runtime; event ordering; parity diagnostics.", "Owner-approved historical parity pass."),
            ("7. Candidate + Test thresholds/windows", "Bounded candidates, compiler adapter, current baseline, candidate simulation, temporal holdout comparison.", "Safe artifacts outside production; correct claim scope."),
            ("8. Broader scenario templates", "Approved condition mutations and full-context simulation.", "Risk/product approval per template."),
            ("9. Optional LLM and documentation RAG", "Model Gateway, masking/DLP, typed tools/output, claim validation, fallback, approved versioned documents.", "Security, red-team, evidence-fidelity, and provider approval."),
            ("10. Production hardening and rollout", "Observability, SLO/load/DR, runbooks, security testing, shadow/UAT, and phased institution enablement.", "Business, risk, security, architecture, and operations sign-off."),
        ],
        widths=[2300, 4700, 2360],
        font_size=7.7,
    )
    doc.add_heading("Primary ownership map", level=2)
    add_table(
        doc,
        ["Owner", "Accountability"],
        [
            ("Case Management team", "Final case semantics, event/correction contract, case read API, decision audit identity."),
            ("Admin / Rule Engine team", "Rule configuration history, compiler exposure, Drools compatibility, governance integration."),
            ("Data platform / DBA", "Approved read boundary, replica/analytical capacity, Custom SQL controls, retention and backups."),
            ("Fraud / Risk product owners", "Outcome mappings, maturity, metrics, constraints, allowed mutations, interpretation, UAT."),
            ("Security / Privacy", "Identity, isolation, threat model, data minimization, LLM egress/provider, audit and incident controls."),
            ("Architecture / Platform", "Workflow/broker/storage/deployment standards, ports, SLO/RTO/RPO, compatibility."),
            ("Operations / SRE", "Monitoring, runbooks, capacity, recovery, on-call, incident and DR exercises."),
            ("Governance / Compliance", "Reviewer roles, four-eyes approval, evidence retention, retirement/change process."),
        ],
        widths=[2600, 6760],
        font_size=8.2,
    )
    doc.add_heading("Decisions management can approve now", level=2)
    add_bullets(doc, [
        "False-positive final case decisions select possible contributing rules; the full eligible population evaluates them.",
        "Final case decision takes precedence over conflicting alert-level decisions.",
        "The case-side transactional outbox and optimizer inbox are mandatory reliability controls; API and recovery polling remain supported, while Kafka is optional.",
        "Historical replay claims require the Phase 0 execution manifest and immutable metric/decision/outcome versions.",
        "The first deployment is one modular Rule Optimization Service; logical components are not mandatory microservices.",
        "Recommendation Only and Candidate + Test are configurable per institution.",
        "The optimizer is institution- and channel-agnostic and has no production rule-write capability.",
        "Authoritative retrieval, calculations, compilation, and replay are deterministic; LLM reasoning is optional and bounded.",
        "Every output is a recommendation for existing human governance.",
        "Recommendation Only is the recommended first production-value stage.",
    ])
    doc.add_heading("Decisions still required before low-level build approval", level=2)
    add_table(
        doc,
        ["Open decision", "Why it matters"],
        [
            ("Optimization hard constraints beyond alert volume", "Defines which candidates are acceptable and how they are ranked."),
            ("Execution-manifest ownership, historical configuration source, and deletion/retirement policy", "Determines whether exact lineage, immutable source versions, and replay remain possible."),
            ("Existing read APIs versus replica/analytical access", "Defines security, performance, and ownership of data retrieval."),
            ("Compiler exposure and full-context replay capability", "Controls when Candidate + Test can be enabled and which claims are valid."),
            ("EFRM infrastructure standards", "Selects broker, durable workflow, object store, identity, secrets, deployment, and observability products."),
            ("Outcome maturity and decision mapping governance", "Defines labels and how mapping changes apply historically."),
            ("Retention, SLO, RTO, RPO, and representative production sizing", "Defines capacity, durability, and support requirements."),
            ("Local/external LLM approval and data policy", "Defines provider, hosting, egress, retention, residency, cost, and disable controls."),
            ("Reviewer UI and governance integration", "Defines actions, roles, evidence drill-down, handoff, and final status lifecycle."),
            ("Pilot institution/channel", "None is selected; implementation must remain configurable and test with representative synthetic fixtures."),
        ],
        widths=[3400, 5960],
        font_size=8.2,
    )

    # Final summary
    add_section_intro(
        doc, "31", "Complete example trace and final summary",
        "This final trace brings the whole workflow together using the fictional R-204 example."
    )
    add_numbered(doc, [
        "An investigator closes case FP-1042. The institution mapping classifies the final decision as FALSE_POSITIVE.",
        "The Case Service writes the required case-finalized outbox event in the same decision transaction. During migration or recovery, the secured API/poller feeds the same contract. The optimizer inbox deduplicates it.",
        "The lineage resolver follows FP-1042 to alert A-781, its result, and stored rule matches.",
        "Attribution selects R-204 as the primary contributor and excludes a coincidental match.",
        "The configuration resolver verifies the source execution manifest, then loads historical version 5, current version 6, full group context, immutable metric/policy versions, and the canonical outcome mapping. The current relevance gate passes.",
        "The snapshot builder retrieves the full eligible historical mobile-transfer population and freezes it with checksums.",
        "Deterministic analytics finds a persistent false-positive concentration during device days 4-7 for established customers using trusted beneficiaries.",
        "The optional LLM explains the pattern and suggests only policy-allowed hypotheses. It does not change any metric.",
        "In Candidate + Test mode, a structured candidate is compiled and replayed after baseline parity. In Recommendation Only mode, executable candidate work is skipped.",
        "The claim validator produces one recommendation with evidence and limitations. A human reviews it. If authorized, the governance adapter may create a draft work item, while existing testing, approval, activation, rollback, retirement, and deletion controls remain authoritative.",
    ])
    add_callout(
        doc,
        "The trust rule",
        "Exact facts come from governed services, SQL, immutable snapshots, statistics, and Drools replay. "
        "The LLM may reason over and explain those facts, but it never becomes the source of the facts. "
        "Every production decision remains human governed.",
        fill=PALE_TEAL,
        accent=TEAL,
    )

    doc.add_heading("One-sentence architecture explanation for management", level=2)
    add_para(
        doc,
        "When a final case is marked false positive, a controlled workflow traces the responsible rule, rebuilds its exact configuration, analyzes its full historical evidence, optionally tests a safe candidate using the existing rule semantics, and sends an evidence-linked recommendation to a human - without ever modifying production rules."
    )

    doc.add_heading("Source basis and status", level=2)
    add_para(
        doc,
        "Basis: the supplied EFRM platform PRD, database-flow reference, PostgreSQL sample schema/dump extracts, current Rule Optimization discussions, "
        "user-confirmed behavior, and a colleague engineering review. Repository-specific observations from that review are explicitly marked reviewer-reported "
        "and require owner verification because the referenced repository revisions are not available in this workspace. The chatbot PDF influenced teaching style only; "
        "sample counts were not used for performance or sizing. Status: refined enterprise logical architecture and implementation blueprint - not a low-level deployment "
        "specification; unapproved infrastructure and production sizing remain unresolved."
    )

    doc.save(OUTPUT_PATH)
    print(str(OUTPUT_PATH))


if __name__ == "__main__":
    build_document()
