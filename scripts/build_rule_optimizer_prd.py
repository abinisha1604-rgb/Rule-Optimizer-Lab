from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "documents" / "EFRM_Rule_Optimization_Agent_Enterprise_PRD.docx"
ASSET = ROOT / "tmp" / "rule_optimizer_prd_assets"
ASSET.mkdir(parents=True, exist_ok=True)
NAVY, BLUE, TEAL, INK, MUTED = "0B2545", "1F5E86", "167A83", "1F2937", "5B677A"
LIGHT, LIGHT_BLUE, PALE_TEAL, GOLD, WHITE, GRID = "F4F7FA", "EAF2F7", "E8F4F3", "A66A00", "FFFFFF", "CBD5E1"

def rgb(x): return RGBColor.from_string(x)
def set_run(run, name="Calibri", size=10.5, color=INK, bold=None, italic=None):
    run.font.name=name; run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"),name); run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"),name)
    run.font.size=Pt(size); run.font.color.rgb=rgb(color)
    if bold is not None: run.bold=bold
    if italic is not None: run.italic=italic

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn("w:shd"))
    if shd is None: shd=OxmlElement("w:shd"); tcPr.append(shd)
    shd.set(qn("w:fill"),fill)

def margins(cell, top=80, start=120, bottom=80, end=120):
    tcPr=cell._tc.get_or_add_tcPr(); tcMar=tcPr.first_child_found_in("w:tcMar")
    if tcMar is None: tcMar=OxmlElement("w:tcMar"); tcPr.append(tcMar)
    for side,val in (("top",top),("start",start),("bottom",bottom),("end",end)):
        n=tcMar.find(qn("w:"+side))
        if n is None: n=OxmlElement("w:"+side); tcMar.append(n)
        n.set(qn("w:w"),str(val)); n.set(qn("w:type"),"dxa")

def border(cell, color=GRID, size="6"):
    tcPr=cell._tc.get_or_add_tcPr(); b=tcPr.first_child_found_in("w:tcBorders")
    if b is None: b=OxmlElement("w:tcBorders"); tcPr.append(b)
    for e in ("top","left","bottom","right","insideH","insideV"):
        n=b.find(qn("w:"+e))
        if n is None: n=OxmlElement("w:"+e); b.append(n)
        n.set(qn("w:val"),"single"); n.set(qn("w:sz"),size); n.set(qn("w:space"),"0"); n.set(qn("w:color"),color)

def cell_width(cell, w):
    tcW=cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
    if tcW is None: tcW=OxmlElement("w:tcW"); cell._tc.get_or_add_tcPr().append(tcW)
    tcW.set(qn("w:w"),str(w)); tcW.set(qn("w:type"),"dxa")

def table_geometry(t, widths):
    t.alignment=WD_TABLE_ALIGNMENT.LEFT; t.autofit=False
    pr=t._tbl.tblPr
    tw=pr.find(qn("w:tblW"))
    if tw is None: tw=OxmlElement("w:tblW"); pr.append(tw)
    tw.set(qn("w:w"),str(sum(widths))); tw.set(qn("w:type"),"dxa")
    ind=pr.find(qn("w:tblInd"))
    if ind is None: ind=OxmlElement("w:tblInd"); pr.append(ind)
    ind.set(qn("w:w"),"120"); ind.set(qn("w:type"),"dxa")
    grid=t._tbl.tblGrid
    for x in list(grid): grid.remove(x)
    for w in widths:
        g=OxmlElement("w:gridCol"); g.set(qn("w:w"),str(w)); grid.append(g)
    for row in t.rows:
        for i,c in enumerate(row.cells):
            cell_width(c,widths[i]); margins(c); border(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_page_num(p):
    r=p.add_run(); b=OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"),"begin"); ins=OxmlElement("w:instrText"); ins.set(qn("xml:space"),"preserve"); ins.text=" PAGE "; e=OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"),"end")
    r._r.append(b); r._r.append(ins); r._r.append(e); set_run(r,size=8.5,color=MUTED)

def add_numbering(doc):
    n=doc.part.numbering_part.element
    def one(num_id, abs_id, fmt, txt):
        a=OxmlElement("w:abstractNum"); a.set(qn("w:abstractNumId"),str(abs_id))
        ns=OxmlElement("w:nsid"); ns.set(qn("w:val"),f"{abs_id:08X}"); a.append(ns)
        mt=OxmlElement("w:multiLevelType"); mt.set(qn("w:val"),"singleLevel"); a.append(mt)
        l=OxmlElement("w:lvl"); l.set(qn("w:ilvl"),"0")
        s=OxmlElement("w:start"); s.set(qn("w:val"),"1"); l.append(s)
        f=OxmlElement("w:numFmt"); f.set(qn("w:val"),fmt); l.append(f)
        t=OxmlElement("w:lvlText"); t.set(qn("w:val"),txt); l.append(t)
        j=OxmlElement("w:lvlJc"); j.set(qn("w:val"),"left"); l.append(j)
        pp=OxmlElement("w:pPr"); ind=OxmlElement("w:ind"); ind.set(qn("w:left"),"720"); ind.set(qn("w:hanging"),"360"); pp.append(ind); l.append(pp)
        if fmt == "bullet":
            rp=OxmlElement("w:rPr"); rf=OxmlElement("w:rFonts"); rf.set(qn("w:ascii"),"Symbol"); rf.set(qn("w:hAnsi"),"Symbol"); rp.append(rf); l.append(rp)
        a.append(l); n.append(a)
        num=OxmlElement("w:num"); num.set(qn("w:numId"),str(num_id)); aid=OxmlElement("w:abstractNumId"); aid.set(qn("w:val"),str(abs_id)); num.append(aid); n.append(num)
    one(900,900,"bullet","\uf0b7"); one(901,901,"decimal","%1.")

def list_num(p, num_id):
    pp=p._p.get_or_add_pPr(); np=OxmlElement("w:numPr"); il=OxmlElement("w:ilvl"); il.set(qn("w:val"),"0"); ni=OxmlElement("w:numId"); ni.set(qn("w:val"),str(num_id)); np.append(il); np.append(ni); pp.append(np)

def diagram(path, kind):
    W,H=(1800,1100) if kind=="architecture" else ((1800,1320) if kind=="workflow" else (1800,1180))
    im=Image.new("RGB",(W,H),"white"); d=ImageDraw.Draw(im)
    try: F=ImageFont.truetype("arial.ttf",25); S=ImageFont.truetype("arial.ttf",20)
    except: F=S=ImageFont.load_default()
    def box(x,y,w,h,text,fill=LIGHT_BLUE,outline=BLUE):
        d.rounded_rectangle((x,y,x+w,y+h),18,fill="#"+fill,outline="#"+outline,width=4)
        lines=[]
        for line in text.split("\n"):
            words=line.split(); cur=""
            for word in words:
                z=(cur+" "+word).strip()
                if d.textbbox((0,0),z,font=S)[2]>w-36 and cur: lines.append(cur); cur=word
                else: cur=z
            if cur: lines.append(cur)
        total=sum(d.textbbox((0,0),z,font=S)[3] for z in lines)+7*(len(lines)-1); yy=y+(h-total)//2
        for z in lines:
            bb=d.textbbox((0,0),z,font=S); d.text((x+(w-bb[2]+bb[0])//2,yy),z,font=S,fill="#"+NAVY); yy+=bb[3]-bb[1]+7
    def arrow(x1,y1,x2,y2):
        d.line((x1,y1,x2,y2),fill="#"+TEAL,width=5)
        import math
        a=math.atan2(y2-y1,x2-x1); q=18
        d.polygon([(x2,y2),(x2-q*math.cos(a-.5),y2-q*math.sin(a-.5)),(x2-q*math.cos(a+.5),y2-q*math.sin(a+.5))],fill="#"+TEAL)
    if kind=="architecture":
        box(60,70,450,210,"Existing EFRM\nPostgreSQL/read replica",PALE_TEAL,TEAL); box(675,70,450,210,"Predefined SQL tools\nRead-only boundary"); box(1290,70,450,210,"Trigger detector\nPoller/event/API",LIGHT,TEAL)
        box(545,430,710,250,"FastAPI + worker\nLangGraph orchestration\nPython tools and policies"); box(110,870,470,180,"Agentic PostgreSQL\n91 application tables",LIGHT,TEAL); box(665,870,470,180,"langgraph_ckpt\n4 runtime tables"); box(1270,870,470,180,"Artifacts + recommendation\nHuman review API/UI",PALE_TEAL,TEAL)
        arrow(510,175,675,175); arrow(1125,175,1290,175); arrow(1515,280,1050,430); arrow(700,680,345,870); arrow(900,680,900,870); arrow(1110,680,1500,870)
    elif kind=="workflow":
        ns=[("1 Load job",50,55),("2 Verify case",410,55),("3 Get alerts",770,55),("4 Get matches",1130,55),("5 Select rules",1490,55),("6 Resolve config",220,410),("7 Snapshot",590,410),("8 Validate",960,410),("9 Analytics",1330,410),("10 Strategy",410,770),("11 Optional LLM",850,770),("12 Compose/validate",1290,770),("13 Save",850,1080)]
        for text,x,y in ns: box(x,y,280,145,text,PALE_TEAL if "LLM" in text else LIGHT_BLUE,GOLD if "LLM" in text else TEAL)
        for a,b in zip(ns,ns[1:]): arrow(a[1]+280,a[2]+72,b[1],b[2]+72)
        arrow(1630,200,360,410); arrow(500,555,550,770); arrow(870,555,990,770); arrow(1240,555,1430,770); arrow(990,915,990,1080); arrow(1430,915,1120,1080)
    else:
        bs=[("case_master\nfinal decision",60,60),("case_alert_mapping\nalert source",675,60),("alert/result\nengine context",1290,60),("match tables\nrule/version",60,390),("rule_version + groups\nmetric definition",675,390),("case_rule_lineage_item\nqueryable evidence",1290,390),("snapshot_manifest\nartifact_part",60,720),("rule_finding_set\nfinding_metric",675,720),("recommendation\nhuman review",1290,720)]
        for text,x,y in bs: box(x,y,450,170,text,PALE_TEAL if "recommendation" in text or "case_master" in text else LIGHT_BLUE,TEAL if "recommendation" in text or "case_master" in text else BLUE)
        for x1,y1,x2,y2 in [(510,145,675,145),(1125,145,1290,145),(285,230,285,390),(900,230,900,390),(1515,230,1515,390),(510,475,675,475),(1125,475,1290,475),(285,560,285,720),(900,560,900,720),(1515,560,1515,720)]: arrow(x1,y1,x2,y2)
    im.save(path)

diagram(ASSET/"architecture.png","architecture"); diagram(ASSET/"workflow.png","workflow"); diagram(ASSET/"dbflow.png","db")

doc=Document(); sec=doc.sections[0]
sec.top_margin=Inches(.78); sec.bottom_margin=Inches(.78); sec.left_margin=Inches(.9); sec.right_margin=Inches(.9); sec.header_distance=Inches(.35); sec.footer_distance=Inches(.35)
normal=doc.styles["Normal"]; normal.font.name="Calibri"; normal._element.rPr.rFonts.set(qn("w:ascii"),"Calibri"); normal._element.rPr.rFonts.set(qn("w:hAnsi"),"Calibri"); normal.font.size=Pt(10.5); normal.font.color.rgb=rgb(INK); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.1
for name,size,col,be,af in [("Heading 1",16,BLUE,16,8),("Heading 2",13,BLUE,12,6),("Heading 3",11.5,NAVY,8,4)]:
    s=doc.styles[name]; s.font.name="Calibri"; s._element.rPr.rFonts.set(qn("w:ascii"),"Calibri"); s._element.rPr.rFonts.set(qn("w:hAnsi"),"Calibri"); s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=rgb(col); s.paragraph_format.space_before=Pt(be); s.paragraph_format.space_after=Pt(af); s.paragraph_format.keep_with_next=True
add_numbering(doc)
hdr=sec.header.paragraphs[0]; hdr.alignment=WD_ALIGN_PARAGRAPH.LEFT; set_run(hdr.add_run("EFRM AI AGENTIC PLATFORM  |  PRODUCT REQUIREMENTS DOCUMENT"),size=8.5,color=MUTED,bold=True)
ftr=sec.footer.paragraphs[0]; ftr.alignment=WD_ALIGN_PARAGRAPH.RIGHT; set_run(ftr.add_run("Internal architecture baseline  |  Page "),size=8.5,color=MUTED); add_page_num(ftr)

def para(text="",style=None,size=10.5,color=INK,align=None,after=None):
    p=doc.add_paragraph(style=style)
    if align is not None: p.alignment=align
    if after is not None: p.paragraph_format.space_after=Pt(after)
    set_run(p.add_run(text),size=size,color=color); return p
def h1(x): return para(x,"Heading 1")
def h2(x): return para(x,"Heading 2")
def bullet(x):
    p=doc.add_paragraph(); list_num(p,4); p.paragraph_format.left_indent=Inches(.28); p.paragraph_format.first_line_indent=Inches(-.18); p.paragraph_format.space_after=Pt(3); set_run(p.add_run(x),size=10.2); return p
def numbered(x):
    p=doc.add_paragraph(); list_num(p,5); p.paragraph_format.left_indent=Inches(.28); p.paragraph_format.first_line_indent=Inches(-.18); p.paragraph_format.space_after=Pt(4); set_run(p.add_run(x),size=10.2); return p
def callout(label,text,fill=LIGHT_BLUE,edge=BLUE):
    t=doc.add_table(rows=1,cols=1); table_geometry(t,[9360]);
    trPr=t.rows[0]._tr.get_or_add_trPr(); trPr.append(OxmlElement("w:cantSplit"))
    c=t.cell(0,0); shade(c,fill); border(c,edge,"12"); p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); set_run(p.add_run(label+"  "),size=10.5,color=edge,bold=True); set_run(p.add_run(text),size=10.5)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
def add_table(headers, rows, widths, font=8.4):
    t=doc.add_table(rows=1,cols=len(headers)); table_geometry(t,widths)
    for i,x in enumerate(headers):
        c=t.rows[0].cells[i]; shade(c,NAVY); p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); set_run(p.add_run(x),size=font,color=WHITE,bold=True)
    for ridx,row in enumerate(rows):
        tr = t.add_row()
        # Keep each business table row together when a table crosses a page.
        # This avoids splitting a correction/contract row across pages and
        # makes the PRD easier to review in Word and in the exported PDF.
        trPr = tr._tr.get_or_add_trPr()
        trPr.append(OxmlElement("w:cantSplit"))
        cs=tr.cells
        for i,x in enumerate(row):
            c=cs[i]; shade(c,"FFFFFF" if ridx%2==0 else LIGHT); p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); set_run(p.add_run(str(x)),size=font)
    doc.add_paragraph().paragraph_format.space_after=Pt(2); return t
def code_block(x):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.18); p.paragraph_format.line_spacing=1.0; set_run(p.add_run(x),name="Consolas",size=8.7,color=NAVY); return p
def pb(): doc.add_page_break()

# Cover
para("PRODUCT REQUIREMENTS DOCUMENT",size=10,color=TEAL,after=4)
title=para("EFRM Rule Optimization Agent",size=28,color=NAVY,after=6); title.runs[0].bold=True
para("Enterprise scope, workflow, data contracts and implementation baseline",size=15,color=MUTED,after=18)
table=add_table(["Document field","Value"],[
("Document status","Approved architecture baseline for implementation planning"),
("Primary audience","Product owners, fraud operations, architects, engineering, risk and governance teams"),
("Phase 1 authority","Recommendation Only; no production rule write"),
("Agent runtime","Python + FastAPI + LangGraph + predefined SQL/Python tools"),
("Database baseline","Agentic PostgreSQL v1.1: agentic schema + langgraph_ckpt schema"),
("Last updated","27 August 2026")],[2200,7160],9.2)
callout("Executive decision","Approve a separate, institution-agnostic Rule Optimization Agent that analyzes finalized false-positive cases, creates traceable evidence, and sends validated recommendations to human reviewers. Production rule changes remain outside the agent.",PALE_TEAL,TEAL)
h1("1. Executive summary")
para("The Rule Optimization Agent is a background analytical service inside the shared Agentic Platform. It is not part of the existing EFRM Rule Engine and it does not generate alerts. A run starts when a case is finally mapped to the institution’s configured FALSE_POSITIVE outcome. The agent traces that case to alerts, results, matches and exact rule versions, then evaluates only the rules connected to that case.")
para("The agent reads authoritative EFRM data through a read-only database boundary. Predefined, parameterized SQL tools retrieve data; deterministic Python and SQL calculate metrics; LangGraph controls the workflow; and an optional LLM turns validated findings into a clear explanation. The recommendation is stored with evidence, limitations, hashes, versions and audit events for human review.")
callout("Plain-English explanation","The agent finds rules that produced false-positive cases, studies why they may be doing so, and gives a human a defensible recommendation. It never changes a live rule by itself.",LIGHT_BLUE,BLUE)

h1("2. Problem, outcomes and principles")
h2("2.1 Problem statement")
para("Transaction-monitoring rules are binary: a rule matches or does not match. A match may still result in a case that an investigator finally closes as a false positive. Those finalized outcomes are the strongest Phase 1 signal for prioritizing rule review. Today, the evidence is distributed across case, alert, result, match, configuration and outcome records. The agent makes that evidence queryable, reproducible and actionable.")
h2("2.2 Business outcomes")
for x in ["Reduce analyst time spent finding the rules behind false-positive cases.","Provide consistent, institution-configurable health, decay, threshold and overlap analysis.","Give fraud and risk teams a recommendation with evidence instead of an unexplained score.","Preserve a complete audit trail so every claim can be reconstructed later.","Create a safe foundation for future candidate testing without granting production write authority."]: bullet(x)
h2("2.3 Design principles")
for x in ["Institution-agnostic and channel-agnostic: mappings, windows, thresholds, data scopes and policies are configuration.","Evidence first: authoritative EFRM facts and deterministic calculations are the source of truth.","Human approval: recommendations are advisory; existing governance owns production change.","Reproducibility: immutable versions, hashes, cut-off times and source watermarks are recorded.","Least privilege: EFRM access is read-only and the LLM receives minimized structured findings.","Bounded data: snapshots are filtered and partitioned; full transaction populations are never loaded into graph memory."]: bullet(x)

h1("3. Scope and authority")
h2("3.1 In scope for Phase 1")
add_table(["Capability","Phase 1 behavior","Owner"],[
("False-positive intake","Find finalized eligible cases using a scheduled SQL poller; support future API/event intake.","Trigger detector"),
("Case-to-rule lineage","Resolve case → alert → result → match → rule/version with source-aware identifiers.","Predefined SQL tools"),
("Configuration resolution","Freeze historical/current rule, group, metric and policy context.","Configuration resolver"),
("Historical snapshot","Build a bounded, partitioned population at a fixed cutoff.","Snapshot builder"),
("Rule health","Calculate configurable volume, yield, false-positive and stability metrics.","Analytics tools"),
("Rule decay","Compare baseline and recent periods with sample-size and persistence gates.","Decay analytics"),
("Threshold review","Run sensitivity analysis around configured threshold values.","Threshold analytics"),
("Overlap/gap review","Find duplicate signals and supported coverage gaps; avoid unsupported fraud claims.","Overlap analytics"),
("Recommendation","Compose, validate and persist a human-review recommendation.","Recommendation service"),
("Audit/recovery","Persist tool, node, model, transition, retry and evidence history.","Agentic Platform")],[1900,5600,1860],8.2)
h2("3.2 Out of scope for Phase 1")
for x in ["Alert generation or live transaction decisions.","Automatic rule creation, modification, retirement, deletion or deployment.","Executable candidate rule compilation and production Rule Engine write.","True fraud recall or definitive risk-leakage claims without an independent truth source.","Unrestricted dynamic SQL or LLM-generated SQL.","Enterprise sizing conclusions based on the supplied sample dataset."]: bullet(x)
h2("3.3 Future operating mode")
para("The target product supports two institution-configurable modes. Mode A, Recommendation Only, is the only enabled mode in Phase 1. Mode B, Candidate + Test, may later generate a candidate, run it through an approved offline simulator or replay service, compare it with the current rule, and then recommend it. Candidate mode requires approved compiler/simulator access, separate permissions and additional simulation tables; it is not part of this delivery.")

h1("4. Product architecture")
doc.add_picture(str(ASSET/"architecture.png"),width=Inches(6.5)); para("Figure 1. Shared Agentic Platform boundary and Rule Optimizer data path",size=8.8,color=MUTED,align=WD_ALIGN_PARAGRAPH.CENTER,after=8)
h2("4.1 Component responsibilities")
add_table(["Component","Responsibility","Must not do"],[
("Trigger detector","Poll or receive finalized case decisions and create a durable idempotent trigger.","Analyze rules or call the LLM."),
("FastAPI service","Expose start/status/recommendation/review APIs and authenticate the caller.","Run long analytics inside the request thread."),
("Background worker","Lease queued jobs, invoke LangGraph, enforce timeouts and recovery.","Bypass authorization or write EFRM rules."),
("LangGraph","Control node order, fan-out, checkpointing and state transitions.","Be the source of business metrics."),
("Predefined SQL tools","Execute approved parameterized reads against EFRM.","Accept model-created SQL or write operational data."),
("Python analytics","Validate data and calculate findings.","Invent missing facts or approve governance."),
("LLM gateway","Optionally explain validated findings in plain language.","Query databases, calculate metrics or approve actions."),
("Agentic repository","Persist jobs, artifacts, evidence, recommendations, audit and review state.","Become the source of truth for EFRM facts."),
("Human governance","Review, approve, reject or request changes and make production changes.","Delegate production change to the agent.")],[1800,4850,2710],8.0)
h2("4.2 Technology choices")
add_table(["Technology","Use in this product","Decision"],[
("Python","Tool handlers, data contracts, analytics, validators and graph code.","Required"),
("FastAPI","Control-plane APIs and health endpoints.","Required"),
("LangGraph","Durable background workflow and fan-out.","Required"),
("FastMCP","Standard interface where shared tools need MCP exposure.","Optional but supported"),
("PostgreSQL","Agentic application data and LangGraph checkpoints.","Required"),
("pgvector","Shared knowledge retrieval for future policy/document support.","Later; not core"),
("Object storage / analytical file store","Partitioned Parquet snapshots and large artifacts.","Required"),
("SFTP","Secure transport for approved files and exports.","Integration option"),
("Local/API LLM","Explanation only, behind an allow-listed gateway.","Optional")],[1900,5100,2360],8.2)

h1("5. Trigger and job lifecycle")
h2("5.1 Trigger options")
add_table(["Option","How it works","Phase 1 position"],[
("Scheduled SQL poller","Reads new finalized case decisions from an EFRM read replica using a durable high-water mark and overlap window.","Primary"),
("Event/API","EFRM publishes CASE_FINALIZED or calls the trigger endpoint; payload is normalized and stored.","Supported when confirmed"),
("Manual reanalysis","Authorized reviewer starts a new generation for an existing case.","Supported with policy")],[1600,5000,2760],8.5)
h2("5.2 Trigger intake sequence")
for x in ["Read the institution’s approved outcome mapping and identify eligible finalized FALSE_POSITIVE outcomes.","Create a normalized payload with institution, case identity, decision identity, source event identity and occurrence time.","Insert into trigger_inbox with a unique institution/source-event key.","Resolve the active agent_version and trigger_binding.","Freeze configuration_snapshot; create agent_invocation, graph_run and case_analysis_job.","Lease the job to a background worker. Duplicate or invalid triggers are recorded, never silently discarded."]: numbered(x)
h2("5.3 Job states")
code_block("RECEIVED → VERIFYING_OUTCOME → RESOLVING_LINEAGE → ELIGIBLE → ANALYZING → COMPLETED\n                                      └→ INELIGIBLE / LIMITED / FAILED / CANCELLED / SUPERSEDED")
callout("Important","The final case decision wins when an alert-level decision conflicts with it. Decision-code meanings are not hardcoded; outcome_mapping_set and outcome_mapping_entry control interpretation for each institution.",LIGHT,GOLD)

h1("6. LangGraph workflow")
doc.add_picture(str(ASSET/"workflow.png"),width=Inches(6.5)); para("Figure 2. Parent graph and per-rule analysis path",size=8.8,color=MUTED,align=WD_ALIGN_PARAGRAPH.CENTER,after=8)
h2("6.1 Parent graph nodes")
add_table(["Node","Input","Processing","Output / next step"],[
("load_job","job_id","Load job, policy, versions and deadlines.","JobContext → verify_case"),
("verify_case","institution_id, case_id","Predefined SQL confirms final eligible outcome and maturity.","VerifiedCase or stop → get_alerts"),
("get_alerts","case_id","Read case-alert mappings with source/type.","AlertRefs → get_rule_matches"),
("get_rule_matches","AlertRefs","Read result/match lineage, rule code and version.","RuleMatchLineage → select_rules"),
("select_rules","RuleMatchLineage","Python deduplicates and records inclusion/exclusion.","RuleTargets → per-rule fan-out"),
("fanout_rules","RuleTargets","Run isolated rule subgraphs with bounded concurrency.","RuleRun IDs → collect_results"),
("collect_results","RuleRun IDs","Join results, failures and limitations.","Job summary → complete_job"),
("complete_job","Job summary","Update status, publish completion and expose API result.","COMPLETED / LIMITED / FAILED")],[1500,2100,3600,2160],7.9)
h2("6.2 Per-rule subgraph nodes")
add_table(["Node","Tool / component","Input → output","LLM?"],[
("resolve_configuration","rule.get_configuration_bundle","Rule target + match time → bundle and evidence.","No"),
("plan_snapshot","optimizer.plan_snapshot","Rule/config/policy → bounded SnapshotPlan.","No"),
("build_snapshot","optimizer.build_snapshot","SnapshotPlan → manifest + artifact parts.","No"),
("validate_snapshot","snapshot.validate","Snapshot ID → quality gates and readiness.","No"),
("health","analytics.calculate_rule_health","Validated snapshot → health metrics.","No"),
("decay","analytics.detect_rule_decay","Baseline/recent metrics → decay finding.","No"),
("threshold","analytics.analyse_threshold","Snapshot + config → sensitivity finding.","No"),
("overlap","analytics.analyse_overlap_and_leakage","Lineage + snapshot → overlap/gap finding.","No"),
("strategy","strategy.build","Findings + policy → bounded strategy items.","No in Phase 1"),
("explain","llm.generate_explanation","Findings + strategy → explanation artifact.","Optional; here only"),
("compose","recommendation.compose","Findings + strategy + explanation → draft.","No"),
("validate","recommendation.validate","Draft + evidence → valid/invalid/limited.","No"),
("save","Agentic repository","Validated draft → recommendation + outbox.","No")],[1700,2650,3900,1110],7.6)

h1("7. Tool contracts and processing rules")
para("Every tool is registered in tool_definition and tool_version, enabled through agent_tool_binding, and assigned to a stage through workflow_tool_binding. The database stores the contract and handler_key; Python code remains the executable implementation. Each call creates tool_execution, authorization_decision and audit_event records.")
h2("7.1 Data-access tools")
add_table(["Tool","What it does","Input","Output / validation"],[
("case.get_final_decision","Reads the final case decision and applies configured mapping.","institution_id, case_id","VerifiedCase; rejects non-final, changed or unmapped outcomes."),
("case.get_alerts","Links a case to alerts using source-aware identifiers.","institution_id, case_id","AlertRef[]; rejects cross-institution rows and missing source identity."),
("alert.get_rule_matches","Follows alert → result → match; captures rule/version.","AlertRef[]","Lineage rows; marks partial/ambiguous/deleted history."),
("rule.get_configuration_bundle","Resolves historical/current rule, group and metric context.","Rule target, match time, fact type","ConfigurationBundle; never guesses a missing version."),
("optimizer.plan_snapshot","Defines windows, cohorts, fields and partitions.","Rule target, policy, cutoff","SnapshotPlan with row estimate and query hashes."),
("optimizer.build_snapshot","Executes approved SQL in pages/partitions and writes Parquet metadata.","SnapshotPlan","Manifest, artifact parts, counts and hashes.")],[2100,3400,2350,1510],7.7)
h2("7.2 Analytics and output tools")
add_table(["Tool","What it does","Input → output","Technique"],[
("snapshot.validate","Checks readiness, lineage, sample size and reproducibility.","Snapshot → gates/limitations","Deterministic Python/SQL"),
("analytics.calculate_rule_health","Computes volume, false-positive rate, yield, conversion and stability.","Snapshot → metrics","SQL aggregates + Python statistics"),
("analytics.detect_rule_decay","Compares baseline/recent periods with persistence gates.","Metrics → decay status","Deltas, intervals, control/change rules"),
("analytics.analyse_threshold","Studies outcomes and volume around threshold alternatives.","Snapshot + config → finding","Sensitivity bands and capacity checks"),
("analytics.analyse_overlap_and_leakage","Finds duplicate signals and supported gaps without unsupported fraud claims.","Lineage + snapshot → finding","Set overlap and observed-outcome analysis"),
("strategy.build","Selects allowed review directions from policy/evidence.","Findings → strategy","Deterministic policy rules"),
("recommendation.compose","Builds canonical JSON and readable summary.","Strategy + evidence → draft","Python schema composition"),
("recommendation.validate","Checks claims, hashes, freshness, permissions and prohibited actions.","Draft + evidence → result","Python validator"),
("llm.generate_explanation","Explains validated facts in plain English.","Findings → explanation","Allow-listed local/API LLM")],[2300,3500,2300,1260],7.5)

h1("8. Snapshot and analytics design")
h2("8.1 What “eligible transactions” means")
para("The builder does not copy every transaction in the EFRM database for every run. It creates a bounded population defined by the rule, institution, channel/fact scope, historical windows and available lineage. The preferred pattern is database aggregation and partitioned extraction, not one giant application-memory result.")
add_table(["Population","Purpose","Required controls"],[
("Rule-evaluated population","Transactions/results where the selected rule was evaluated or matched.","Point-in-time rule/version context; source watermark."),
("False-positive-linked population","Alerts/cases linked to finalized false-positive outcomes.","Outcome mapping and maturity policy."),
("Comparable non-match population","Comparable records where technically available.","Explicit cohort definition; no invented negatives."),
("Near-threshold band","Records close to current threshold for sensitivity analysis.","Threshold band policy and minimum sample size."),
("Overlap population","Records where selected and other rules fired together.","Source-aware match lineage.")],[2200,4300,2860],8.2)
h2("8.2 Snapshot controls")
for x in ["Fixed analysis_cutoff and UTC timezone.","Window_from/window_to stored in snapshot_manifest.","Institution, rule, fact scope and exclusions included in scope_hash.","Query hashes, metric-definition hashes, extraction code and connector versions persisted.","Partitioned Parquet in approved object storage/analytical file store; only manifest, artifact parts, hashes, counts and evidence in PostgreSQL.","No raw full snapshot in LangGraph checkpoints.","A failed readiness gate blocks actions configured by data_quality_result: analysis, LLM or recommendation."]: bullet(x)
h2("8.3 Rule-health metric families")
add_table(["Family","Examples","Guardrail"],[
("Volume","Match count, alert count, case count, alert rate","Always show denominator and window."),
("Outcome quality","False-positive rate among finalized outcomes, yield, conversion","Do not treat unresolved cases as negatives."),
("Contribution","Share of alerts/cases where the rule contributed","Use match-level lineage and deduplicate."),
("Stability","Period variance, intervals, segment drift","Require minimum sample and sustained pattern."),
("Capacity","Alert-volume change under recommendation","Alert-volume limit is the current capacity constraint."),
("Evidence strength","Completeness, maturity, lineage and version resolution","Separate evidence strength from severity.")],[1900,4300,3160],8.2)
h2("8.4 Decay decision logic")
for x in ["Compare configured baseline and recent windows.","Calculate absolute and relative changes for selected metrics.","Apply minimum sample size and data-quality gates.","Require persistence; one anomalous day is not decay.","Classify NOT_DETECTED, SUSPECTED, CONFIRMED or UNKNOWN.","Carry periods, formulas, thresholds and limitations into evidence."]: numbered(x)

h1("9. LLM boundary and reasoning design")
callout("Single approved LLM location","The LLM is called in the explain node, after snapshot validation, deterministic analytics and strategy selection. It receives structured findings and evidence references, never unrestricted database access.",PALE_TEAL,TEAL)
add_table(["Question","Decision"],[
("Does the LLM retrieve EFRM data?","No. Predefined SQL tools retrieve it."),
("Does it calculate health or decay?","No. Python/SQL calculate official metrics."),
("Does it select the recommendation?","No in Phase 1. strategy.build applies deterministic policy."),
("What does it add?","A readable explanation, supported causes and limitations."),
("What if it fails?","Use a deterministic template and record fallback_used."),
("Can external LLM be used?","Only through an allow-listed gateway with masking, egress policy and structured output."),
("What is stored?","Deployment, prompt, hashes, tokens, latency, validation and artifact references; never secrets or private chain-of-thought.")],[2600,6760],8.4)
h2("9.1 Explanation validation")
for x in ["Claim extraction identifies material statements.","Each claim links to supporting, limiting or contradicting evidence.","Numbers are compared to finding_metric values.","Unsupported fraud, recall or leakage claims are rejected.","A stale snapshot, stale rule hash or failed blocking gate invalidates the recommendation."]: bullet(x)

pb(); h1("10. Recommendation product")
h2("10.1 Recommendation structure")
code_block('{\n  "rule": {"rule_reference": "...", "analysed_version": "..."},\n  "issue": "...",\n  "recommended_action_type": "REVIEW_THRESHOLD",\n  "recommended_change": "...",\n  "reason": "...",\n  "evidence": [{"evidence_id": "...", "supports": "..."}],\n  "expected_impact": {"alert_volume_direction": "DECREASE"},\n  "evidence_strength": "MEDIUM",\n  "limitations": ["..."],\n  "action_code": "HUMAN_REVIEW_REQUIRED"\n}')
h2("10.2 Recommendation decision logic")
add_table(["Condition","Result"],[
("Case is not final FALSE_POSITIVE","Stop as INELIGIBLE; no recommendation."),
("Lineage or historical configuration is ambiguous","Store LIMITED/DEFERRED analysis; do not invent a version."),
("Snapshot readiness is blocked","Stop analytics or recommendation according to gate policy."),
("Evidence is sufficient and policy is met","Create REVIEW_REQUIRED recommendation."),
("Evidence is insufficient","Create a collect-more-evidence or no-change recommendation."),
("Current rule changed after analysis","Mark recommendation STALE; require reanalysis."),
("Human approves","Record review_decision and status history; governance owns implementation."),
("Human rejects/requests changes","Record reason and preserve immutable history.")],[3400,5960],8.2)
h2("10.3 Human-review experience")
for x in ["Reviewers see the issue, proposed action, evidence strength, periods, metrics, limitations and exact rule/version references.","Every claim can open its evidence artifact and source reference.","A reviewer can approve, reject, acknowledge or request changes.","No review action grants the agent permission to write production rules.","Future UI may expose progress through SSE/WebSocket; recommendations remain available through REST."]: bullet(x)

h1("11. Database design and table flow")
doc.add_picture(str(ASSET/"dbflow.png"),width=Inches(6.5)); para("Figure 3. Authoritative EFRM reads and Agentic Platform persistence",size=8.8,color=MUTED,align=WD_ALIGN_PARAGRAPH.CENTER,after=8)
h2("11.1 Physical setup")
add_table(["Area","Tables / ownership","Purpose"],[
("EFRM operational database","Existing EFRM tables; separate read-only role efrm_agent_read_ro","Authoritative cases, alerts, transactions, devices, matches, rules and decisions."),
("Agentic application","agentic schema; 91 application-managed tables","Configuration, jobs, tools, artifacts, evidence, analytics, recommendations, chat and audit."),
("LangGraph runtime","langgraph_ckpt schema; 4 runtime-managed tables","Checkpoint persistence owned by installed LangGraph checkpointer."),
("Large data","Object storage / analytical file store","Partitioned Parquet snapshots and large artifacts."),
("Transfer","SFTP connector where approved","Moves documents/exports; not primary analytical storage.")],[1800,3800,3760],8.1)
h2("11.2 EFRM read flow")
code_block("case_master → case_alert_mapping → transaction_alert/device_alert\n→ transaction_result/device_result → transaction_match/device_match\n→ transaction_master/device_master → rule_master → rule_version\n→ rule groups/bindings → rule_metric_dependency → metric_definition")
para("All access is through predefined, parameterized SQL handlers. The optimizer does not assume a service connection to the Case or Alert service. If later integrations are approved, they implement the same tool contracts without changing the graph.")
h2("11.3 Agentic write flow")
code_block("reconciliation_cursor → trigger_inbox → agent_activation/trigger_binding\n→ configuration_snapshot → agent_invocation → graph_run\n→ case_analysis_job → trigger_case_link/case_rule_lineage_item\n→ rule_analysis_run → configuration_bundle/item → snapshot_manifest/artifact_part\n→ data_quality_result → rule_finding_set/finding_metric\n→ strategy_set/strategy_item → model_execution (optional)\n→ claim_record/claim_evidence_link → recommendation → outbox_event\n→ review_decision/recommendation_status_history")
h2("11.4 Cross-cutting writes")
for x in ["graph_node_execution records every node attempt, lease, retry and output reference.","tool_execution records every planned, authorized, denied, running, successful or failed tool call.","authorization_decision records deterministic capability decisions.","audit_event and graph_transition provide business and security history.","idempotency_record prevents duplicate jobs and outputs.","dead_letter_item holds items requiring manual recovery."]: bullet(x)

h1("12. Updated Agentic DB v1.1 alignment")
para("The updated database is the baseline for this PRD. It resolves the earlier separation, ownership and audit gaps: one efrm_agentic database, agentic application schema, dedicated langgraph_ckpt ownership, configurable tool bindings, outcome mapping, immutable configuration, artifact/evidence chains, recovery state and human review history.")
h2("12.1 Key tables used by the Rule Optimizer")
add_table(["Domain","Tables","How the optimizer uses them"],[
("Routing/configuration","agent_definition, agent_version, agent_activation, workflow_definition, workflow_version, workflow_stage, trigger_binding","Selects active graph and institution configuration."),
("Tools","tool_definition, tool_version, agent_tool_binding, workflow_tool_binding, tool_validation_run, tool_capability_requirement, tool_field_rule","Controls available predefined tools and contracts."),
("Policies","configuration_definition/version/binding/snapshot/source, feature_flag_definition/assignment, outcome_mapping_set/entry","Freezes policy and decision-code meaning."),
("Runtime","trigger_inbox, agent_invocation, graph_run, graph_node_execution, graph_signal, idempotency_record, outbox_event, dead_letter_item, reconciliation_cursor","Makes execution durable, idempotent, recoverable and publishable."),
("Optimizer","case_analysis_job, rule_analysis_run, trigger_case_link, case_rule_lineage_item, configuration_bundle/item, snapshot_manifest, data_quality_result, rule_finding_set, finding_metric, strategy_set/item, recommendation, review_decision, recommendation_status_history","Stores the full analytical lifecycle."),
("Evidence/audit","artifact, artifact_part, evidence_item, evidence_relation, claim_record, claim_evidence_link, audit_event, authorization_decision, tool_execution, model_execution, graph_transition","Connects every claim to evidence and execution history."),
("Checkpoints","langgraph_ckpt.checkpoints, checkpoint_blobs, checkpoint_writes, checkpoint_migrations","Stores graph state only; never the full snapshot.")],[1900,4400,3060],7.6)
pb(); h2("12.2 Corrections required before migration sign-off")
add_table(["Correction","Why it matters","Required action"],[
("Source-aware lineage","Transaction-specific fields do not fully support device/future fact types or duplicate IDs across alert sources.","Add alert_source/alert_type and generic result/fact/match references."),
("Multiple rule versions in one case","A case may contain matches for the same logical rule at different versions.","Include historical version/execution key in uniqueness."),
("Deleted current rule","A deleted current rule may have no current hash.","Make recommendation.current_rule_hash_at_creation nullable for MISSING/DELETED."),
("LLM metadata meaning","Phase 1 strategy selection is deterministic while schema permits LLM-assisted strategy reasoning.","Use DETERMINISTIC_ONLY for strategy; add explanation_mode/model reference."),
("Evidence subject dictionary","STRATEGY_ITEM is used but not listed in examples.","Add it to the controlled subject-type dictionary."),
("Tool seed data","A schema without bindings cannot run the graph.","Seed all Rule Optimizer tools, stages, versions, policies and bindings.")],[2000,4300,3060],7.8)

h1("13. APIs and service contracts")
h2("13.1 Control-plane APIs")
add_table(["Endpoint","Purpose","Authorization"],[
("POST /v1/rule-optimizer/jobs","Start authorized analysis/reanalysis or accept normalized trigger.","Service/admin capability; institution from trusted context."),
("GET /v1/rule-optimizer/jobs/{job_id}","Return status, node progress, limitations and recommendation IDs.","Same-institution read access."),
("GET /v1/rule-optimizer/recommendations/{id}","Return canonical recommendation, evidence summary and stale state.","Reviewer capability."),
("POST /v1/rule-optimizer/recommendations/{id}/review","Append human review decision and governance reference.","Approved reviewer; idempotency key required."),
("POST /v1/rule-optimizer/jobs/{id}/cancel","Request cancellation through graph_signal.","Job owner/admin."),
("GET /health/live and /health/ready","Service and dependency health.","Platform health access.")],[2600,4650,2110],8.2)
h2("13.2 Internal tool contract rules")
for x in ["Input and output are JSON-schema validated before handler execution.","Tool arguments are built by Python graph code; model output cannot set institution_id or bypass scope.","Handlers receive trusted institution, correlation ID, deadline, policy snapshot and read-only connection.","Result artifacts are content-addressed with SHA-256 hashes.","Large results are returned by artifact reference, not embedded in graph state.","Timeouts, retryability, result bytes and classification come from tool_version."]: bullet(x)

pb(); h1("14. Non-functional requirements")
add_table(["Area","Requirement","Acceptance measure"],[
("Tenant isolation","Every institution-owned row has institution_id; RLS and tenant-consistent composite FKs are enforced.","Cross-tenant tests return zero rows and audit DENY."),
("Security","EFRM uses efrm_agent_read_ro; least privilege; no BYPASSRLS/SUPERUSER.","Role/grant review and negative authorization tests pass."),
("Reliability","Jobs are durable, leased, idempotent, checkpointed and recoverable.","Worker failure resumes without duplicate recommendation."),
("Data integrity","Snapshots, configs, findings and recommendations are hash-linked and versioned.","Same cutoff reproduces manifest and metric hashes."),
("Performance","Queries are bounded by institution/window/rule; large data is partitioned/aggregated.","Load tests use later production-like volume assumptions."),
("Freshness","Recommendations become stale when rule hash or source data changes.","Stale check prevents obsolete review."),
("Observability","Metrics, logs, traces and audit cover every node/tool/model call.","Dashboards alert on queue lag, DQ, retries and validation failures."),
("Privacy","LLM egress disabled by default; masking/classification apply.","Restricted raw transaction data does not leave boundary."),
("Retention","Retention is configurable and respects legal hold.","Expiry never deletes held/evidence-linked artifacts.")],[1700,5000,2660],7.8)
h2("14.1 Error and retry policy")
add_table(["Failure","Action","Terminal record"],[
("Transient EFRM/database timeout","Retry with bounded backoff; do not advance cursor until success.","tool_execution + audit_event"),
("Invalid SQL result/schema","Stop node; quarantine tool version; alert owner.","tool_validation_run / dead_letter_item"),
("Missing/ambiguous historical rule","Mark bundle partial/ambiguous and stop unsupported claims.","configuration_bundle/item + limitation"),
("Snapshot too large","Reduce by partition/page limits or stop with capacity limitation.","snapshot_manifest + data_quality_result"),
("LLM timeout/provider failure","Retry only within policy; use deterministic explanation.","model_execution with fallback_used"),
("Recommendation validation failure","Do not publish; retain draft and validation reason.","recommendation status / graph failure"),
("Worker crash/lease expiry","Recover graph from checkpoint and node idempotency key.","graph_run + graph_node_execution"),
("Repeated failure","Move to dead-letter queue for manual recovery.","dead_letter_item")],[2200,4850,2310],7.8)
h2("14.2 Audit minimum")
for x in ["Correlation ID from trigger through recommendation.","Actor identity for poller, service, worker, tool and reviewer.","Agent/workflow/tool/model versions and configuration snapshot hash.","Input/output artifact hashes and source watermarks.","Authorization decisions, retry attempts, errors and final status.","Evidence links for every material recommendation claim."]: bullet(x)

h1("15. Test and acceptance strategy")
h2("15.1 Functional acceptance criteria")
for x in ["A finalized eligible false-positive case creates exactly one initial case_analysis_job.","A duplicate trigger does not create a second job for the same source event.","A case is traced to source-aware alerts, results, matches and rule versions.","Rules unrelated to the triggering case are not analyzed.","The snapshot is bounded, partitioned and reproducible at its cutoff.","Findings contain formulas, periods, denominators and evidence IDs.","The recommendation is rejected if a material claim lacks evidence.","LLM outage still produces a validated deterministic explanation when policy allows.","No execution path writes to EFRM rule configuration in Phase 1.","Human review creates append-only review and status records."]: bullet(x)
h2("15.2 Test layers")
add_table(["Layer","Examples"],[
("Unit","Outcome mapping, scope filters, metric formulas, decay gates, threshold bands, claim validation."),
("Contract","SQL result schemas, tool schemas, EFRM source-aware identifiers and API schemas."),
("Integration","Read replica, Agentic DB/RLS, object storage, checkpointer and optional LLM gateway."),
("Replay","Known lineages, deleted versions, duplicate alerts and corrected decisions."),
("Reliability","Worker crash, lease expiry, retry storm, outbox outage, restore and dead-letter recovery."),
("Security","Cross-tenant reads, prompt injection in stored text, unauthorized review/write and LLM egress."),
("Performance","Partitioned extraction, per-rule fan-out and high-volume audit/tool history.")],[1700,7660],8.2)

pb(); h1("16. Implementation plan")
add_table(["Stage","Deliverables","Exit gate"],[
("0. Contract and schema hardening","Apply v1.1 corrections; seed agents, workflow, tools, policies and mappings.","Migration and tenant/RLS tests pass."),
("1. Platform foundation","FastAPI, worker lease loop, repositories, secrets, read-only EFRM connector and object storage connector.","Health, auth and idempotency tests pass."),
("2. Trigger and lineage","SQL poller, cursor, inbox, outcome verification, alerts, matches and source-aware lineage.","Known cases resolve correctly."),
("3. LangGraph runtime","Parent graph, per-rule fan-out, checkpointer, node audit, retry and cancellation.","Crash/restart/replay tests pass."),
("4. Configuration and snapshot","Bundle resolver, planner/builder, manifest, artifacts and quality gates.","Reproducible bounded snapshot produced."),
("5. Analytics","Health, decay, threshold and overlap tools with versioned formulas/evidence.","Known-result fixtures pass."),
("6. Recommendation","Deterministic strategy, optional explanation, claim validation, API and outbox.","No unsupported claims; review contract passes."),
("7. Hardening","Observability, retention, DR, load, privacy, security review and runbooks.","Production readiness approved.")],[1800,5200,2360],7.8)
h2("16.1 Production approval boundary")
callout("Phase 1 release rule","The service may read EFRM data and write Agentic Platform records only. It may not call a production rule-write API, submit a rule for deployment, retire a rule or alter a rule configuration. A human must use the existing EFRM governance process for any change.",PALE_TEAL,TEAL)
h2("16.2 Open inputs required before production")
for x in ["Institution outcome mappings, finality/maturity rules and approval roles.","Exact semantics of rule_version, rule_group_version and match history identifiers.","Approved EFRM read views/replica, performance limits and source connector contract.","Metric Definition Custom SQL governance and point-in-time reproducibility policy.","Minimum sample-size, health, decay and evidence-strength policies.","Approved local/API LLM provider, classification, masking and retention decisions.","Historical rule retention/deletion contract and handling of deleted configurations.","Approved object storage/analytical file store and legal-hold policy."]: bullet(x)

h1("17. Appendix A — status and ownership reference")
h2("A.1 Who owns each decision?")
add_table(["Decision","Owner"],[
("Is the case finally eligible?","Outcome mapping + deterministic verifier"),
("Which rules are analyzed?","Lineage resolver + deterministic selector"),
("Are data and metrics valid?","Snapshot/data-quality and analytics validators"),
("Which improvement is suggested?","Deterministic strategy policy; LLM only explains"),
("Is recommendation ready?","Recommendation validator"),
("Is a production rule changed?","Authorized human and existing EFRM governance")],[5200,4160],8.7)
h2("A.2 Final management explanation")
callout("One sentence","The Rule Optimization Agent finds finalized false-positive cases, traces them to the responsible rules, uses controlled SQL and deterministic Python analysis to build evidence, optionally uses an LLM to explain that evidence, and sends a validated recommendation to a human without changing any production rule.",LIGHT_BLUE,BLUE)
h2("A.3 Final architecture decision")
para("Use Python, FastAPI and LangGraph in a separate Agentic Platform repository. Use predefined, read-only SQL tools for current EFRM access, PostgreSQL for application data and checkpoints, object storage for large snapshots, and an optional LLM only for explanation. Keep RAG available for future policy/document context, but do not use it as the source of transaction, case or rule truth.")

doc.core_properties.title="EFRM Rule Optimization Agent - Enterprise PRD"; doc.core_properties.subject="Final scope, workflow, data contracts and implementation requirements"; doc.core_properties.author="Agentic Platform Architecture"; doc.core_properties.keywords="EFRM, rule optimization, LangGraph, recommendation, false positive, fintech"
OUT.parent.mkdir(parents=True,exist_ok=True); doc.save(OUT); print(OUT)
