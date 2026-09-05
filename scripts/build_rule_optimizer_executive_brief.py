from pathlib import Path
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "documents" / "EFRM_Rule_Optimization_Agent_Executive_Architecture_Brief.docx"
ASSET_DIR = ROOT / "tmp" / "rule_optimizer_executive_brief"
DIAGRAM = ASSET_DIR / "architecture.png"

NAVY = "16364F"
BLUE = "2E74B5"
TEAL = "168891"
LIGHT_BLUE = "EAF2F8"
LIGHT_TEAL = "E8F4F4"
GREEN = "4C8A4B"
LIGHT_GREEN = "EAF4EA"
ORANGE = "ED8A19"
LIGHT_ORANGE = "FFF2DF"
RED = "C65353"
LIGHT_RED = "FCEAEA"
GRAY = "66727E"
LIGHT_GRAY = "F2F4F7"
GRID = "C9D3DC"
WHITE = "FFFFFF"
BLACK = "1F2D3A"


def font(path_name, size):
    p = Path(r"C:\Windows\Fonts") / path_name
    return ImageFont.truetype(str(p), size=size)


def draw_centered(draw, box, text, fnt, fill, spacing=5):
    x1, y1, x2, y2 = box
    max_chars = max(12, int((x2 - x1) / (fnt.size * 0.58)))
    lines = []
    for part in text.split("\n"):
        lines.extend(wrap(part, width=max_chars) or [""])
    heights = [draw.textbbox((0, 0), line, font=fnt)[3] for line in lines]
    total = sum(heights) + spacing * (len(lines) - 1)
    y = y1 + ((y2 - y1) - total) / 2
    for line, h in zip(lines, heights):
        bbox = draw.textbbox((0, 0), line, font=fnt)
        w = bbox[2] - bbox[0]
        draw.text((x1 + ((x2 - x1) - w) / 2, y), line, font=fnt, fill=fill)
        y += h + spacing


def rounded(draw, box, fill, outline, radius=22, width=4):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def arrow(draw, start, end, color=TEAL, width=6):
    color = "#" + color if not color.startswith("#") else color
    draw.line([start, end], fill=color, width=width)
    x2, y2 = end
    x1, y1 = start
    import math
    angle = math.atan2(y2 - y1, x2 - x1)
    size = 18
    p1 = (x2 - size * math.cos(angle - 0.55), y2 - size * math.sin(angle - 0.55))
    p2 = (x2 - size * math.cos(angle + 0.55), y2 - size * math.sin(angle + 0.55))
    draw.polygon([end, p1, p2], fill=color)


def build_diagram():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (2000, 980), "white")
    d = ImageDraw.Draw(img)
    title = font("seguisb.ttf", 36)
    head = font("seguisb.ttf", 28)
    body = font("segoeui.ttf", 22)
    small = font("segoeui.ttf", 18)

    d.text((70, 38), "Rule Optimization Agent - logical architecture", font=title, fill="#" + NAVY)
    d.text((70, 88), "Existing EFRM remains authoritative. The optimizer is an asynchronous decision-support boundary.", font=body, fill="#" + GRAY)

    # Existing EFRM lane
    d.rounded_rectangle((60, 145, 1940, 330), radius=28, fill="#F3F7FB", outline="#" + BLUE, width=4)
    d.text((95, 165), "EXISTING EFRM", font=head, fill="#" + BLUE)
    boxes = [
        ((100, 220, 430, 300), "Case Service\nFinal case outcome"),
        ((515, 220, 845, 300), "Admin / Configuration\nRules, policies, versions"),
        ((930, 220, 1260, 300), "Read APIs / Replica\nAlerts, matches, history"),
        ((1345, 220, 1675, 300), "Rule Engine + Compiler\nDrools / JSON-to-DRL"),
    ]
    for b, txt in boxes:
        rounded(d, b, "#FFFFFF", "#A8BBD0", radius=16, width=3)
        draw_centered(d, b, txt, body, "#" + NAVY)

    # Optimizer lane
    d.rounded_rectangle((60, 380, 1940, 760), radius=28, fill="#F4FBFB", outline="#" + TEAL, width=4)
    d.text((95, 400), "RULE OPTIMIZATION BOUNDARY", font=head, fill="#" + TEAL)
    y1, y2 = 480, 590
    opt = [
        ((95, y1, 370, y2), "1. Trigger +\nconfirm case"),
        ((440, y1, 715, y2), "2. Trace +\nattribute rules"),
        ((785, y1, 1060, y2), "3-4. Resolve +\nfreeze evidence"),
        ((1130, y1, 1405, y2), "5-6. Analyze +\nform strategies"),
        ((1475, y1, 1905, y2), "7-8. Test or advise\nthen recommend"),
    ]
    fills = [LIGHT_BLUE, LIGHT_BLUE, LIGHT_TEAL, LIGHT_TEAL, LIGHT_GREEN]
    outlines = [BLUE, BLUE, TEAL, TEAL, GREEN]
    for (b, txt), fill, ol in zip(opt, fills, outlines):
        rounded(d, b, "#" + fill, "#" + ol, radius=18, width=4)
        draw_centered(d, b, txt, head, "#" + NAVY)
    for i in range(len(opt) - 1):
        arrow(d, (opt[i][0][2] + 6, 535), (opt[i + 1][0][0] - 8, 535))

    # Source arrows
    arrow(d, (265, 305), (232, 475), BLUE, 5)
    arrow(d, (680, 305), (920, 475), BLUE, 5)
    arrow(d, (1095, 305), (920, 475), BLUE, 5)
    arrow(d, (1510, 305), (1685, 475), BLUE, 5)

    # Mode and optional LLM notes
    rounded(d, (930, 635, 1255, 720), "#" + LIGHT_ORANGE, "#" + ORANGE, radius=16, width=3)
    draw_centered(d, (930, 635, 1255, 720), "Optional LLM\nHypotheses + explanation", body, "#" + NAVY)
    arrow(d, (1255, 678), (1415, 590), ORANGE, 4)
    rounded(d, (1320, 635, 1905, 720), "#" + LIGHT_GREEN, "#" + GREEN, radius=16, width=3)
    draw_centered(d, (1320, 635, 1905, 720), "Configured branch: Candidate + Test OR Recommendation Only", body, "#" + NAVY)

    # Governance lane
    rounded(d, (300, 825, 1700, 930), "#EAF5F6", "#" + TEAL, radius=20, width=4)
    draw_centered(d, (300, 825, 1700, 930), "Human review -> Existing governance -> Any production change\nThe Agent never activates, retires, deletes, or directly modifies a production rule.", head, "#" + NAVY)
    arrow(d, (1690, 760), (1500, 820), TEAL, 5)
    img.save(DIAGRAM, quality=95)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn("w:" + m))
        if node is None:
            node = OxmlElement("w:" + m)
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=GRID, size=5):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn("w:" + edge)
        el = borders.find(tag)
        if el is None:
            el = OxmlElement("w:" + edge)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths)))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "120")
    tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(width))
            tcW.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run(run, size=11, bold=False, color=BLACK, italic=False, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text="", size=11, bold=False, color=BLACK, after=6, before=0,
             align=WD_ALIGN_PARAGRAPH.LEFT, italic=False, keep=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    p.paragraph_format.keep_with_next = keep
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, color=color, italic=italic)
    return p


def add_bullet(doc, label, text, size=10.3, after=3):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.50)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(label + ": ")
    set_run(r, size=size, bold=True, color=NAVY)
    r = p.add_run(text)
    set_run(r, size=size, color=BLACK)
    return p


def set_paragraph_band(p, fill, border):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "22")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), border)
    pBdr.append(left)
    pPr.append(pBdr)


def add_callout(doc, label, text, fill=LIGHT_TEAL, border=TEAL, after=8):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.10)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    set_paragraph_band(p, fill, border)
    r = p.add_run(label + "  ")
    set_run(r, size=10.5, bold=True, color=NAVY)
    r = p.add_run(text)
    set_run(r, size=10.5, color=BLACK)
    return p


def add_step(doc, number, title, bullets, llm_text, output_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    set_paragraph_band(p, LIGHT_BLUE if number not in (6, 7) else LIGHT_ORANGE, BLUE if number not in (6, 7) else ORANGE)
    r = p.add_run(f"STEP {number}  ")
    set_run(r, size=11.2, bold=True, color=TEAL if number not in (6, 7) else ORANGE)
    r = p.add_run(title)
    set_run(r, size=11.2, bold=True, color=NAVY)
    for label, value in bullets:
        add_bullet(doc, label, value, size=9.7, after=1.5)
    add_bullet(doc, "LLM", llm_text, size=9.7, after=1.5)
    add_bullet(doc, "Output", output_text, size=9.7, after=3)


def add_table(doc, headers, rows, widths, font_size=9.2):
    t = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(t, widths)
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        shade_cell(c, NAVY)
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_run(r, size=9.2, bold=True, color=WHITE)
    for idx, row in enumerate(rows):
        cells = t.add_row().cells
        for i, value in enumerate(row):
            shade_cell(cells[i], WHITE if idx % 2 == 0 else LIGHT_GRAY)
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(str(value))
            set_run(r, size=font_size, color=BLACK)
    set_table_geometry(t, widths)
    return t


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run(run, size=8.5, color=GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ):
        s = styles[style_name]
        s.font.name = "Calibri"
        s._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        s._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True


def page_break(doc):
    doc.add_page_break()


def build_doc():
    build_diagram()
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_styles(doc)
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.72)
    sec.bottom_margin = Inches(0.72)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)
    sec.header_distance = Inches(0.32)
    sec.footer_distance = Inches(0.32)

    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = hp.add_run("EFRM RULE OPTIMIZATION AGENT  |  EXECUTIVE ARCHITECTURE BRIEF")
    set_run(r, size=8.2, bold=True, color=GRAY)
    add_page_number(sec.footer.paragraphs[0])

    # PAGE 1
    add_para(doc, "EXECUTIVE ARCHITECTURE BRIEF", size=10.2, bold=True, color=TEAL, after=5)
    add_para(doc, "EFRM Rule Optimization Agent", size=25, bold=True, color=NAVY, after=3)
    add_para(doc, "Clear architecture and step-by-step working", size=13.2, color=GRAY, after=10)
    strip = add_table(
        doc,
        ["TRIGGER", "DECISION SUPPORT", "CONTROL"],
        [("Final false-positive case", "Evidence-based recommendation", "Human governance only")],
        [3120, 3120, 3120],
        font_size=9.4,
    )
    for i, cell in enumerate(strip.rows[1].cells):
        shade_cell(cell, [LIGHT_TEAL, LIGHT_BLUE, LIGHT_GREEN][i])
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(NAVY)
    add_callout(
        doc,
        "In one sentence",
        "When a finalized case is mapped to false positive, the Agent finds the rules that truly contributed, studies their complete historical evidence, optionally tests a safe candidate, and sends one recommendation to a human reviewer.",
        after=6,
    )
    doc.add_picture(str(DIAGRAM), width=Inches(6.42))
    p = doc.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    add_para(doc, "Logical components are modules, not mandatory microservices. Existing EFRM remains the source of truth.", size=8.7, italic=True, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=3)
    add_callout(doc, "Core trust rule", "Deterministic services create the facts. The optional LLM explains approved facts. Humans control every production decision.", fill=LIGHT_GREEN, border=GREEN, after=0)

    # PAGE 2
    page_break(doc)
    doc.add_heading("End-to-end workflow - Steps 1 to 4", level=1)
    add_para(doc, "The false-positive case selects possible rules for investigation. It does not prove that every alert or matched rule is wrong.", size=10.3, color=GRAY, after=3)
    add_step(doc, 1, "Receive and confirm the trigger", [
        ("Input", "Institution, case, decision/action identity, final decision code and event time."),
        ("Processing", "Case outbox/API/poller feeds one authenticated contract; inbox deduplicates; Outcome Resolver re-reads the authoritative case and applies the effective institution mapping."),
        ("Technique", "Spring Boot backend logic, PostgreSQL inbox/job state, service authentication and idempotency."),
    ], "No. Decision-code interpretation must be exact.", "Accepted false-positive case-analysis job, or an explicit ignored/rejected reason.")
    add_step(doc, 2, "Trace the case and attribute contributing rules", [
        ("Input", "Accepted case-analysis job."),
        ("Processing", "Follow case -> alerts -> results -> stored matches; process each alert separately; reproduce group and decision context to classify primary, supporting, coincidental or unresolved rules."),
        ("Technique", "Approved read APIs or reviewed parameterized SQL plus deterministic attribution logic."),
    ], "No. Relationships and contribution must be reproducible.", "Evidence-linked CaseRuleLineage; only primary/policy-approved supporting rules continue.")
    add_step(doc, 3, "Resolve exact historical and current configuration", [
        ("Input", "Attributed rule reference, institution, event time and execution-manifest reference."),
        ("Processing", "Load the production-time execution manifest, historical rule bundle and current bundle; resolve rule/group/policy/metric/reference/compiler versions and compare current relevance."),
        ("Technique", "Configuration Resolver behind RuleConfigurationPort; version IDs and SHA-256 hashes."),
    ], "No. Configuration resolution is an identity and history problem.", "Immutable historical/current bundles, hashes and a relevance or ambiguity result.")
    add_step(doc, 4, "Build a frozen historical evidence snapshot", [
        ("Input", "Resolved rule bundles, windows, required fields, metric versions and institution scope."),
        ("Processing", "Retrieve the full eligible population - fired and non-fired events, mature outcomes and required facts - then freeze it with query/configuration/payload hashes."),
        ("Technique", "Read-only API/replica/analytical store, governed Metric SQL Executor, PostgreSQL metadata and optional Parquet/object storage."),
    ], "No. The LLM never writes or executes SQL.", "Immutable snapshot plus READY, READY_WITH_LIMITS or a named stop status.")

    # PAGE 3
    page_break(doc)
    doc.add_heading("End-to-end workflow - Steps 5 to 8", level=1)
    add_step(doc, 5, "Calculate trusted rule findings", [
        ("Input", "Frozen snapshot, current/historical rule bundles and institution analysis policy."),
        ("Processing", "Calculate fire rate, alert conversion, internal false-positive burden, mature outcome yield, repeat burden, overlap, threshold proximity, data quality, decay and internal leakage indicators."),
        ("Technique", "SQL plus deterministic Python statistics: Polars/pandas, NumPy, SciPy and approved EWMA/CUSUM/change-point methods."),
    ], "No for authoritative metrics.", "RuleFindingSet with denominators, uncertainty, limitations and evidence IDs.")
    add_step(doc, 6, "Form bounded improvement strategies", [
        ("Input", "Verified findings, structured current rule, allowed mutations, ranges and alert-volume constraint."),
        ("Processing", "Deterministic generator enumerates approved threshold/window/condition changes; optional LLM connects verified patterns into bounded hypotheses and plain-English reasoning."),
        ("Technique", "Java strategy module, Python search functions, JSON Schema and provider-neutral Model Gateway."),
    ], "Optional and useful here; it receives only approved evidence and cannot invent metrics or write rule code.", "CandidateStrategySet containing testable candidates or structured human change proposals.")
    add_step(doc, 7, "Execute the institution's configured mode", [
        ("Candidate + Test", "Create candidate JSON only in optimizer storage; validate; compile through RuleCompilerPort; prove historical parity; replay current and candidate on the identical frozen cohort/context."),
        ("Recommendation Only", "Do not create executable candidate JSON/DRL and do not run candidate backtesting; convert verified findings into human review advice with an explicit no-backtest limitation."),
        ("Technique", "Mode Router, Candidate Validator, isolated Java/Drools RuleReplayPort and deterministic result comparison."),
    ], "Optional explanation only; never compilation or execution.", "Mode-specific evidence: simulation comparison or analysis-only proposal.")
    add_step(doc, 8, "Publish one evidence-linked recommendation", [
        ("Input", "Verified findings, candidate/simulation status, policy/configuration versions, evidence IDs and limitations."),
        ("Processing", "Build canonical recommendation JSON first; Claim Validator checks every number, scope, evidence link, allowed statement and production-action restriction."),
        ("Technique", "Spring Boot Recommendation Composer, PostgreSQL evidence metadata, secured API/UI and immutable audit."),
    ], "Optional for summary and the readable 'why'; deterministic template is the fallback.", "REVIEW_REQUIRED recommendation for a human; the workflow does not change production.")

    # PAGE 4
    page_break(doc)
    doc.add_heading("Two operating modes and the role of the LLM", level=1)
    add_table(
        doc,
        ["Decision point", "Mode A - Candidate + Test", "Mode B - Recommendation Only"],
        [
            ("When used", "Compiler and isolated replay access are permitted.", "Client does not permit executable candidate or compiler access."),
            ("Candidate", "Typed JSON stored only in optimizer boundary.", "No executable candidate or DRL."),
            ("Testing", "Parity -> current baseline -> candidate simulation on the same cohort.", "No candidate backtest; deterministic sensitivity may still be shown."),
            ("Output", "Recommendation + candidate diff + simulation evidence.", "Recommendation + verified analysis + explicit no-backtest limitation."),
            ("Production write", "Never.", "Never."),
            ("Human approval", "Required.", "Required."),
        ],
        [1900, 3730, 3730],
        font_size=9.0,
    )
    doc.add_heading("Where the LLM is used", level=2)
    add_table(
        doc,
        ["LLM may help", "LLM must never do"],
        [
            ("Explain verified patterns in simple English.", "Browse raw EFRM tables or hold database credentials."),
            ("Connect approved findings into bounded hypotheses.", "Generate or execute arbitrary SQL."),
            ("Suggest an approved scenario template.", "Calculate authoritative metrics or alter evidence."),
            ("Draft the recommendation's 'why' and limitations.", "Write free-form DRL, compile rules or run replay."),
            ("Request an allow-listed follow-up calculation.", "Approve a candidate or decide a production change."),
        ],
        [4680, 4680],
        font_size=9.2,
    )
    add_callout(doc, "Why this hybrid design is better", "SQL, statistics and Drools are repeatable from the same inputs. The LLM adds language reasoning where it helps, but never becomes the source of truth.", fill=LIGHT_GREEN, border=GREEN, after=5)
    doc.add_heading("Human governance is the final control", level=2)
    for label, text in [
        ("Reviewer", "views evidence, accepts, rejects, requests more analysis or recommends retirement review."),
        ("Optional handoff", "after recorded human acceptance, a Governance Adapter may create a draft work item if an approved API and institution policy permit it."),
        ("Existing EFRM governance", "owns testing, approval, activation, rollback, retirement and deletion."),
        ("Agent boundary", "the Agent never directly modifies, activates, retires or deletes a production rule."),
    ]:
        add_bullet(doc, label, text, size=10.0, after=3)

    # PAGE 5
    page_break(doc)
    doc.add_heading("Technical implementation and approval summary", level=1)
    add_callout(doc, "Initial deployment", "One modular Java/Spring Boot Rule Optimization Service. Logical responsibilities remain separate modules; only resource-heavy analytics or replay workers are isolated when needed.", after=6)
    add_table(
        doc,
        ["Component", "Responsibility", "Primary technique"],
        [
            ("Trigger + Orchestrator", "Authenticate, deduplicate, persist state, retry and resume the eight phases.", "Spring Boot + PostgreSQL inbox/outbox/jobs/state machine"),
            ("Data + Configuration", "Read exact cases, lineage, rules, versions and historical population.", "Typed ports, secured read APIs or parameterized read-only SQL"),
            ("Analytics", "Health, decay, threshold, overlap, volume and internal leakage evidence.", "Deterministic SQL + Python statistics"),
            ("Candidate + Replay", "Validate, compile and compare current/candidate under identical context.", "RuleCompilerPort + isolated Java/Drools RuleReplayPort"),
            ("Optional Model Gateway", "Bounded hypotheses and explanation over approved evidence.", "Local LLM or approved external API; masking and schema output"),
            ("Recommendation Center", "Evidence drill-down, review action and governed handoff.", "Spring REST APIs/UI + immutable PostgreSQL audit"),
        ],
        [2100, 3900, 3360],
        font_size=8.8,
    )
    doc.add_heading("Non-negotiable enterprise controls", level=2)
    controls = [
        "The case-side transactional outbox and optimizer inbox prevent lost or duplicate work; Kafka is optional.",
        "institution_id is mandatory on every event, query, join, job, snapshot, artifact, cache key and audit record.",
        "Historical claims require execution manifests and immutable metric, decision-policy and outcome-mapping versions.",
        "EFRM access is read-only; candidates and simulations remain outside production.",
        "Replay parity must pass before candidate-impact claims are allowed.",
        "Case corrections or current-rule changes cancel, stale or supersede older work.",
    ]
    for text in controls:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.50)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.05
        set_run(p.add_run(text), size=9.5, color=BLACK)
    doc.add_heading("Recommended approval and delivery order", level=2)
    order = [
        "P0 provenance: case outbox, canonical mappings, execution manifests and immutable versions.",
        "Modular trigger, lineage, configuration and frozen-evidence foundation.",
        "Deterministic Recommendation Only plus the real Recommendation Center.",
        "Point-in-time replay parity, then bounded Candidate + Test.",
        "Optional LLM and documentation RAG after the deterministic workflow is mature.",
    ]
    for i, text in enumerate(order, 1):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.50)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.05
        set_run(p.add_run(text), size=9.5, color=BLACK)
    add_callout(doc, "Management decision", "Approve the architecture as an asynchronous, institution/channel-agnostic, evidence-first recommendation system. Enable Candidate + Test only after compiler access, replay parity and client constraints are approved.", fill=LIGHT_ORANGE, border=ORANGE, after=0)

    props = doc.core_properties
    props.title = "EFRM Rule Optimization Agent - Executive Architecture Brief"
    props.subject = "Five-page architecture and workflow summary"
    props.author = "EFRM Product Architecture"
    props.keywords = "EFRM, rule optimization, architecture, workflow, backtesting, LLM"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
